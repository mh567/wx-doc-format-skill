from __future__ import annotations

import re
from typing import Callable
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def normalized_name(name: str | None) -> str:
    return (name or "").casefold().replace(" ", "")


def resolved_style(profile: dict, role: str, fallback: str) -> str:
    return profile.get("resolved_styles", {}).get(role, fallback)


def style_name(obj) -> str:
    try:
        return obj.style.name if obj.style is not None else ""
    except Exception:
        return ""


def set_style(obj, style: str, corrections: list[dict], correction_type: str, text: str = "") -> None:
    current = style_name(obj)
    if normalized_name(current) == normalized_name(style):
        return
    try:
        obj.style = style
        corrections.append({"type": correction_type, "from": current, "to": style, "text": text[:120]})
    except Exception as exc:
        corrections.append({"type": f"{correction_type}_failed", "from": current, "to": style, "error": str(exc), "text": text[:120]})


def style_by_role(profile: dict) -> dict[str, str]:
    return {
        "body": resolved_style(profile, "body", "Normal"),
        "title": resolved_style(profile, "title", "文档标题"),
        "caption": resolved_style(profile, "caption", "Caption"),
        "note": resolved_style(profile, "note", "3.1注-无编号注"),
        "numbered_note": resolved_style(profile, "numbered_note", "3.2注-有编号注"),
        "table_body": resolved_style(profile, "table_body", "表正文"),
        "appendix_title": resolved_style(profile, "appendix_title", "附录标题"),
        "list_letter": resolved_style(profile, "list_letter", "1.1一级列项-编号"),
        "list_dash": resolved_style(profile, "list_dash", "1.2一级列项-无编号"),
        "list_decimal": resolved_style(profile, "list_decimal", "2.1二级列项-有编号"),
        "list_bullet": resolved_style(profile, "list_bullet", "2.2二级列项-无编号"),
    }


def paragraph_role_for_style(profile: dict, paragraph) -> str | None:
    name = normalized_name(style_name(paragraph))
    styles = {role: normalized_name(value) for role, value in style_by_role(profile).items()}
    for role, normalized in styles.items():
        if name == normalized:
            return role
    for level in range(1, 7):
        heading = profile.get("resolved_styles", {}).get(f"heading_{level}", f"Heading {level}")
        if name == normalized_name(heading):
            return f"heading_{level}"
    return None


def apply_paragraph_style_aliases(doc, profile: dict) -> list[dict]:
    corrections = []
    styles = style_by_role(profile)
    aliases = {
        normalized_name("题注"): styles["caption"],
        normalized_name("Caption"): styles["caption"],
        normalized_name("正文"): styles["body"],
        normalized_name("Normal"): styles["body"],
        normalized_name("3.1 注-无编号注"): styles["note"],
        normalized_name("3.1注-无编号注"): styles["note"],
        normalized_name("3.2 注-有编号注"): styles["numbered_note"],
        normalized_name("3.2注-有编号注"): styles["numbered_note"],
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        target = aliases.get(normalized_name(style_name(paragraph)))
        if target:
            set_style(paragraph, target, corrections, "paragraph_style_alias_normalized", text)
    return corrections


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def set_paragraph_lines(paragraph, lines: list[str]) -> None:
    paragraph.text = ""
    for index, line in enumerate(lines):
        if index == 0:
            paragraph.add_run(line)
        else:
            paragraph.add_run().add_break()
            paragraph.add_run(line)


_APPENDIX_MARK_RE = re.compile(r"^附\s*录\s*([A-Z])\s*$")
_APPENDIX_CLASSIFICATION_RE = re.compile(r"^[（(](规范性|资料性)[)）]$")
_APPENDIX_CLASSIFICATION_TITLE_RE = re.compile(r"^[（(](规范性|资料性)[)）]\s*.+$")


def finalize_appendix_structure(doc, profile: dict) -> list[dict]:
    corrections = []
    appendix_style = resolved_style(profile, "appendix_title", "附录标题")
    index = 0
    while index < len(doc.paragraphs):
        paragraph = doc.paragraphs[index]
        text = paragraph.text.strip()
        match = _APPENDIX_MARK_RE.match(text)
        if not match:
            index += 1
            continue
        letter = match.group(1)
        next_one = doc.paragraphs[index + 1] if index + 1 < len(doc.paragraphs) else None
        next_two = doc.paragraphs[index + 2] if index + 2 < len(doc.paragraphs) else None
        next_one_text = next_one.text.strip() if next_one is not None else ""
        next_two_text = next_two.text.strip() if next_two is not None else ""
        # Three-paragraph case: 附录A, (资料性), Title
        if _APPENDIX_CLASSIFICATION_RE.match(next_one_text) and next_two_text:
            set_paragraph_lines(paragraph, ["", next_one_text, next_two_text])
            set_style(paragraph, appendix_style, corrections, "appendix_title_style_normalized", paragraph.text)
            delete_paragraph(next_two)
            delete_paragraph(next_one)
            corrections.append(
                {
                    "type": "appendix_three_paragraph_header_merged",
                    "letter": letter,
                    "classification": next_one_text,
                    "title": next_two_text[:120],
                }
            )
            index += 1
            continue

        # Two-paragraph case: 附录A, (资料性)接口字段与处置命令约束
        if _APPENDIX_CLASSIFICATION_TITLE_RE.match(next_one_text):
            set_paragraph_lines(paragraph, ["", next_one_text])
            set_style(paragraph, appendix_style, corrections, "appendix_title_style_normalized", paragraph.text)
            delete_paragraph(next_one)
            corrections.append(
                {
                    "type": "appendix_two_paragraph_header_merged",
                    "letter": letter,
                    "title": next_one_text[:120],
                }
            )
            index += 1
            continue

        # Standalone appendix mark: clear text, auto-numbering handles display
        paragraph.text = ""
        corrections.append({"type": "appendix_title_prefix_removed", "from": text, "to": ""})
        set_style(paragraph, appendix_style, corrections, "appendix_title_style_normalized", paragraph.text)
        index += 1
    return corrections


def finalize_template_tables(
    doc,
    profile: dict,
    row_height_cm: float,
    row_height_rule: str,
    *,
    row_height_rule_enum,
    cm,
    left_alignment,
    center_alignment,
    set_table_autofit_to_window: Callable,
    looks_like_code_sample_table: Callable,
) -> list[dict]:
    corrections = []
    from text_utils import looks_like_api_example_table as _looks_api
    table_body_style = resolved_style(profile, "table_body", "表正文")
    for table_index, table in enumerate(doc.tables, 1):
        set_table_autofit_to_window(table)
        is_code_sample = looks_like_code_sample_table(table)
        is_api_example = _looks_api(table)
        for row_index, row in enumerate(table.rows, 1):
            target_rule = row_height_rule_enum.AT_LEAST if row_height_rule == "at-least" else row_height_rule_enum.EXACTLY
            if row.height_rule != target_rule:
                row.height_rule = target_rule
                corrections.append({"type": "table_row_height_rule_normalized", "table": table_index, "row": row_index})
            if row.height is None or abs(row.height.cm - row_height_cm) > 0.02:
                row.height = cm(row_height_cm)
                corrections.append({"type": "table_row_height_normalized", "table": table_index, "row": row_index})
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if not text:
                        continue
                    set_style(
                        paragraph,
                        table_body_style,
                        corrections,
                        "table_cell_style_normalized",
                        text,
                    )
                    # Clear run-level formatting so the style definition takes full effect.
                    _KEEP = {qn('w:rStyle'), qn('w:lang'), qn('w:bCs'), qn('w:iCs')}
                    for run in paragraph.runs:
                        rpr = run._element.find(qn('w:rPr'))
                        if rpr is None:
                            continue
                        for child in list(rpr):
                            if child.tag not in _KEEP:
                                rpr.remove(child)
                    paragraph.alignment = left_alignment if (is_code_sample or is_api_example) else center_alignment
    return corrections


def add_table_borders(doc) -> list[dict]:
    """Add full borders (all sides + internal) to every table in the document."""
    corrections = []
    for table_index, table in enumerate(doc.tables, 1):
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None:
            tblPr.remove(existing)

        borders = OxmlElement('w:tblBorders')
        for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            borders.append(border)

        tblPr.append(borders)
        corrections.append({"type": "table_borders_added", "table": table_index})

    return corrections


def normalize_caption_prefixes(doc) -> list[dict]:
    """Ensure every Caption paragraph starts with 表 or 图 prefix before the SEQ field."""
    from docx.oxml.ns import qn as _qn
    from lxml import etree as _etree
    corrections = []
    for paragraph in doc.paragraphs:
        sname = paragraph.style.name if paragraph.style else ''
        if 'Caption' not in sname and 'caption' not in sname:
            continue
        # Check if prefix already exists
        text = paragraph.text
        if text.startswith('表') or text.startswith('图'):
            continue

        # Determine type from SEQ field
        has_table_seq = 'SEQ Table' in _etree.tostring(paragraph._element, encoding='unicode')
        prefix = '表' if has_table_seq else '图'

        # Insert prefix run as the first run (before SEQ field)
        r0 = _etree.Element(_qn('w:r'))
        t0 = _etree.SubElement(r0, _qn('w:t'))
        t0.text = prefix + ' '
        t0.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        # Insert after pPr (if any), before first child
        pPr = paragraph._element.find(_qn('w:pPr'))
        if pPr is not None:
            pPr.addnext(r0)
        else:
            paragraph._element.insert(0, r0)
        corrections.append({'type': 'caption_prefix_added', 'prefix': prefix})

    return corrections


def audit_template_styles(doc, profile: dict) -> dict:
    allowed = {normalized_name(value) for value in profile.get("resolved_styles", {}).values()}
    # TOC styles inserted by the finalizer itself are always allowed.
    allowed |= {normalized_name(n) for n in ("toc 1", "toc 2", "toc 3")}
    used = []
    unexpected = []
    for index, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph.text.strip()
        if not text:
            continue
        name = style_name(paragraph)
        role = paragraph_role_for_style(profile, paragraph)
        record = {"paragraph": index, "style": name, "role": role, "text": text[:120]}
        used.append(record)
        if normalized_name(name) not in allowed:
            unexpected.append(record)
    for table_index, table in enumerate(doc.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if not text:
                        continue
                    name = style_name(paragraph)
                    record = {
                        "table": table_index,
                        "row": row_index,
                        "style": name,
                        "role": paragraph_role_for_style(profile, paragraph),
                        "text": text[:120],
                    }
                    used.append(record)
                    if normalized_name(name) not in allowed:
                        unexpected.append(record)
    return {"used_styles": used, "unexpected_styles": unexpected}


def audit_toc_and_page_layout(doc) -> dict:
    toc_like = []
    for index, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph.text.strip()
        if text in {"目次", "目录", "目  次"}:
            toc_like.append({"paragraph": index, "text": text, "style": style_name(paragraph)})
    sections = []
    for index, section in enumerate(doc.sections, 1):
        sections.append(
            {
                "section": index,
                "header_distance_cm": section.header_distance.cm if section.header_distance is not None else None,
                "footer_distance_cm": section.footer_distance.cm if section.footer_distance is not None else None,
                "has_header_text": bool(section.header.paragraphs and "".join(p.text for p in section.header.paragraphs).strip()),
                "has_footer_text": bool(section.footer.paragraphs and "".join(p.text for p in section.footer.paragraphs).strip()),
            }
        )
    warnings = []
    if toc_like:
        warnings.append(
            {
                "type": "toc_requires_office_update",
                "message": "目录或目次需要由 Word/WPS 更新域、页码和缩进。",
                "count": len(toc_like),
            }
        )
    if len(sections) < 2:
        warnings.append(
            {
                "type": "single_section_page_numbering",
                "message": "当前文档只有一个节，若需要封面/目录/正文不同页码格式，应在 Word/WPS 中复核分节页码。",
            }
        )
    return {"toc_like_paragraphs": toc_like, "sections": sections, "warnings": warnings}
def _style_font_from_profile(profile: dict, style_name: str) -> dict:
    styles = profile.get("styles", {})
    style = styles.get(style_name) or styles.get(profile.get("by_key", {}).get(style_name.casefold().replace(" ", "")))
    if not style:
        return {}
    result = {}
    ea = style.get("font_east_asia")
    ascii_f = style.get("font_ascii")
    h_ansi = style.get("font_h_ansi")
    sz_ht = style.get("font_size_halftones")
    if ea:
        result["east_asia"] = ea
    if ascii_f:
        result["ascii"] = ascii_f
    if h_ansi:
        result["h_ansi"] = h_ansi
    if sz_ht is not None:
        result["size_pt"] = int(sz_ht) / 2.0
    if style.get("bold") is not None:
        result["bold"] = style["bold"]
    return result


_STYLE_FONT_FALLBACK = {
    "文档标题": {"east_asia": "黑体", "ascii": "Times New Roman", "size_pt": 14, "bold": False},
    "heading": {"east_asia": "黑体", "ascii": "黑体", "size_pt": 12, "bold": False},
    "caption": {"east_asia": "黑体", "ascii": "Times New Roman", "size_pt": 12, "bold": False},
    "note": {"east_asia": "宋体", "ascii": "Times New Roman", "size_pt": 10.5, "bold": False},
    "numbered_note": {"east_asia": "宋体", "ascii": "Times New Roman", "size_pt": 10.5, "bold": False},
    "table_body": {"east_asia": "宋体", "ascii": "Times New Roman", "size_pt": 10.5, "bold": False},
    "body": {"east_asia": "宋体", "ascii": "Times New Roman", "size_pt": 12, "bold": None},
    "formula": {"east_asia": "宋体", "ascii": "Times New Roman", "size_pt": 12, "bold": False},
    "appendix_title": {"east_asia": "黑体", "ascii": "Times New Roman", "size_pt": 14, "bold": False},
    "list": {"east_asia": "宋体", "ascii": "Times New Roman", "size_pt": 12, "bold": False},
}


def _resolve_font_for_style(profile: dict, style_name: str) -> dict:
    font = _style_font_from_profile(profile, style_name)
    if font:
        return font
    normalized = (style_name or "").casefold().replace(" ", "")
    if normalized.startswith("heading") or style_name.startswith("Heading"):
        return dict(_STYLE_FONT_FALLBACK["heading"])
    if "附录" in style_name and "标题" in style_name:
        return dict(_STYLE_FONT_FALLBACK["appendix_title"])
    if "注-无编号注" in style_name:
        return dict(_STYLE_FONT_FALLBACK["note"])
    if "注-有编号注" in style_name:
        return dict(_STYLE_FONT_FALLBACK["numbered_note"])
    if "题注" in style_name or style_name == "Caption" or style_name == "caption":
        return dict(_STYLE_FONT_FALLBACK["caption"])
    if "列项" in style_name:
        return dict(_STYLE_FONT_FALLBACK["list"])
    if style_name in ("表正文",):
        return dict(_STYLE_FONT_FALLBACK["table_body"])
    if "文档标题" in style_name:
        return dict(_STYLE_FONT_FALLBACK["文档标题"])
    return dict(_STYLE_FONT_FALLBACK["body"])


def remove_api_example_captions(doc) -> list[dict]:
    """Delete Caption paragraphs that precede API-example (single-cell JSON/HTTP) tables."""
    from text_utils import looks_like_api_example_table as _is_api
    corrections = []
    body = doc.element.body
    to_remove = []
    for i, child in enumerate(body):
        if not child.tag.endswith(('}w:tbl', '}tbl')):
            continue
        from docx.table import Table
        table = Table(child, doc)
        if not _is_api(table):
            continue
        # Look backwards for a Caption paragraph
        for j in range(i - 1, -1, -1):
            prev = body[j]
            if not prev.tag.endswith(('}w:p', '}p')):
                continue
            from docx.text.paragraph import Paragraph
            pp = Paragraph(prev, doc)
            if 'Caption' in (pp.style.name or ''):
                to_remove.append(prev)
            break
    for el in to_remove:
        body.remove(el)
        corrections.append({'type': 'api_example_caption_removed'})
    return corrections


def finalize_run_fonts(doc, profile: dict) -> list[dict]:
    corrections = []
    visited = set()
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name_str = style_name(paragraph)
        if not style_name_str:
            continue
        font_rule = _resolve_font_for_style(profile, style_name_str)
        for run in paragraph.runs:
            run_id = id(run)
            if run_id in visited:
                continue
            visited.add(run_id)
            changed = []
            ea = font_rule.get("east_asia")
            ascii_f = font_rule.get("ascii")
            sz_pt = font_rule.get("size_pt")
            bold_val = font_rule.get("bold")
            if ea or ascii_f:
                try:
                    run.font.name = ascii_f or "Times New Roman"
                except Exception:
                    pass
                rpr = run._element.get_or_add_rPr()
                existing_rpr_fonts = rpr.find(qn('w:rFonts'))
                if existing_rpr_fonts is None:
                    existing_rpr_fonts = OxmlElement("w:rFonts")
                    rpr.append(existing_rpr_fonts)
                if ea:
                    existing_rpr_fonts.set(qn('w:eastAsia'), ea)
                if ascii_f:
                    existing_rpr_fonts.set(qn('w:ascii'), ascii_f)
                    existing_rpr_fonts.set(qn('w:hAnsi'), ascii_f)
                changed.append("font")
            if sz_pt is not None:
                try:
                    run.font.size = Pt(sz_pt)
                    changed.append("size")
                except Exception:
                    pass
            if bold_val is not None:
                try:
                    run.bold = bold_val
                    changed.append("bold")
                except Exception:
                    pass
            if changed:
                corrections.append({
                    "type": "run_font_normalized",
                    "style": style_name_str,
                    "changes": changed,
                    "font_rule": font_rule,
                    "text": text[:80],
                })
    # Table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if not text:
                        continue
                    style_name_str = style_name(paragraph) or "表正文"
                    font_rule = _resolve_font_for_style(profile, style_name_str) if style_name_str != "表正文" else _STYLE_FONT_FALLBACK["table_body"]
                    for run in paragraph.runs:
                        run_id = id(run)
                        if run_id in visited:
                            continue
                        visited.add(run_id)
                        changed = []
                        ea = font_rule.get("east_asia")
                        ascii_f = font_rule.get("ascii")
                        sz_pt = font_rule.get("size_pt")
                        bold_val = font_rule.get("bold")
                        if ea or ascii_f:
                            try:
                                run.font.name = ascii_f or "Times New Roman"
                            except Exception:
                                pass
                            rpr = run._element.get_or_add_rPr()
                            existing_rpr_fonts = rpr.find(qn('w:rFonts'))
                            if existing_rpr_fonts is None:
                                existing_rpr_fonts = OxmlElement("w:rFonts")
                                rpr.append(existing_rpr_fonts)
                            if ea:
                                existing_rpr_fonts.set(qn('w:eastAsia'), ea)
                            if ascii_f:
                                existing_rpr_fonts.set(qn('w:ascii'), ascii_f)
                                existing_rpr_fonts.set(qn('w:hAnsi'), ascii_f)
                            changed.append("font")
                        if sz_pt is not None:
                            try:
                                run.font.size = Pt(sz_pt)
                                changed.append("size")
                            except Exception:
                                pass
                        if bold_val is not None:
                            try:
                                run.bold = bold_val
                                changed.append("bold")
                            except Exception:
                                pass
                        if changed:
                            corrections.append({
                                "type": "run_font_normalized",
                                "style": style_name_str,
                                "changes": changed,
                                "font_rule": font_rule,
                                "text": text[:80],
                            })
    return corrections




def insert_table_of_contents(doc, profile: dict) -> dict:
    """Insert a TOC field with "目  次" heading and page break at the start."""
    from docx.oxml import OxmlElement as _Oxml
    from docx.oxml.ns import qn as _qn

    max_level = 0
    for level in range(1, 10):
        if profile.get("resolved_styles", {}).get(f"heading_{level}"):
            max_level = level
    if max_level == 0:
        max_level = 3

    body = doc.element.body

    # --- "目  次" title (first) — 黑体, Normal style, centered ---
    p_title = doc.add_paragraph()
    try:
        p_title.style = doc.styles['Normal']
    except Exception:
        pass
    p_title.alignment = 1  # center
    r_t = _Oxml('w:r')
    rpr_t = _Oxml('w:rPr')
    rfonts_t = _Oxml('w:rFonts')
    rfonts_t.set(_qn('w:eastAsia'), '黑体')
    rfonts_t.set(_qn('w:ascii'), '黑体')
    rfonts_t.set(_qn('w:hAnsi'), '黑体')
    rpr_t.append(rfonts_t)
    r_t.append(rpr_t)
    t_t = _Oxml('w:t')
    t_t.text = '目  次'
    t_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_t.append(t_t)
    p_title._element.append(r_t)
    body.insert(0, p_title._element)

    # --- TOC field (second) ---
    p_toc = doc.add_paragraph()
    try:
        p_toc.style = doc.styles['toc 1']
    except Exception:
        pass

    # begin
    r_begin = _Oxml('w:r')
    fc_begin = _Oxml('w:fldChar')
    fc_begin.set(_qn('w:fldCharType'), 'begin')
    r_begin.append(fc_begin)
    p_toc._element.append(r_begin)

    # instruction
    r_instr = _Oxml('w:r')
    instr = _Oxml('w:instrText')
    instr.text = f'TOC \\o "1-{max_level}" \\h \\z'
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r_instr.append(instr)
    p_toc._element.append(r_instr)

    # separator
    r_sep = _Oxml('w:r')
    fc_sep = _Oxml('w:fldChar')
    fc_sep.set(_qn('w:fldCharType'), 'separate')
    r_sep.append(fc_sep)
    p_toc._element.append(r_sep)

    # display text
    r_text = _Oxml('w:r')
    t_text = _Oxml('w:t')
    t_text.text = '目录'
    r_text.append(t_text)
    p_toc._element.append(r_text)

    # end
    r_end = _Oxml('w:r')
    fc_end = _Oxml('w:fldChar')
    fc_end.set(_qn('w:fldCharType'), 'end')
    r_end.append(fc_end)
    p_toc._element.append(r_end)

    body.insert(1, p_toc._element)

    # --- Page break in the last run of the TOC paragraph ---
    r_pb = _Oxml('w:r')
    br = _Oxml('w:br')
    br.set(_qn('w:type'), 'page')
    r_pb.append(br)
    p_toc._element.append(r_pb)

    return {"toc_inserted": True, "levels": max_level}


def apply_template_finalizer(
    doc,
    profile: dict | None,
    row_height_cm: float,
    row_height_rule: str,
    *,
    row_height_rule_enum,
    cm,
    left_alignment,
    center_alignment,
    set_table_autofit_to_window: Callable,
    looks_like_code_sample_table: Callable,
) -> dict:
    if not profile:
        return {"enabled": False, "corrections": [], "style_audit": {}, "layout_audit": {}}
    corrections = []
    corrections.append(insert_table_of_contents(doc, profile))
    corrections.extend(finalize_appendix_structure(doc, profile))
    corrections.extend(apply_paragraph_style_aliases(doc, profile))
    corrections.extend(normalize_caption_prefixes(doc))
    corrections.extend(remove_api_example_captions(doc))
    corrections.extend(finalize_run_fonts(doc, profile))
    corrections.extend(
        finalize_template_tables(
            doc,
            profile,
            row_height_cm,
            row_height_rule,
            row_height_rule_enum=row_height_rule_enum,
            cm=cm,
            left_alignment=left_alignment,
            center_alignment=center_alignment,
            set_table_autofit_to_window=set_table_autofit_to_window,
            looks_like_code_sample_table=looks_like_code_sample_table,
        )
    )
    corrections.extend(add_table_borders(doc))
    return {
        "enabled": True,
        "corrections": corrections,
        "style_audit": audit_template_styles(doc, profile),
        "layout_audit": audit_toc_and_page_layout(doc),
    }
