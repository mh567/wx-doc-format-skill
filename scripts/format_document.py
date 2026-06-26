#!/usr/bin/env python3
from __future__ import annotations

import argparse
from copy import deepcopy
import json
import os
from pathlib import Path
import re
import sys
import zipfile

DOCX_IMPORT_ERROR = None
try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.shared import Cm, Pt, RGBColor
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except Exception as exc:
    DOCX_IMPORT_ERROR = exc


def maybe_reexec_with_skill_venv() -> None:
    if DOCX_IMPORT_ERROR is None:
        return
    if os.environ.get("WX_DOC_FORMAT_NO_REEXEC") == "1":
        return
    script_path = Path(__file__).resolve()
    skill_dir = script_path.parent.parent
    candidates = []
    env_venv = os.environ.get("WX_DOC_FORMAT_VENV")
    if env_venv:
        candidates.append(Path(env_venv) / "bin" / "python")
    candidates.append(skill_dir / ".venv" / "bin" / "python")
    current = Path(sys.executable).resolve()
    for candidate in candidates:
        if candidate.exists() and candidate.resolve() != current:
            env = os.environ.copy()
            env["WX_DOC_FORMAT_NO_REEXEC"] = "1"
            os.execve(str(candidate), [str(candidate), str(script_path), *sys.argv[1:]], env)


STYLE_BY_MD_LEVEL = {
    1: "文档标题",
    2: "Heading 1",
    3: "Heading 2",
    4: "Heading 3",
    5: "Heading 4",
    6: "Heading 6",
}

HEADING_PATTERNS = [
    (re.compile(r"^第[一二三四五六七八九十百千万0-9]+章\s*[:：、.\s]?\s*(.+)$"), 1),
    (re.compile(r"^第[一二三四五六七八九十百千万0-9]+节\s*[:：、.\s]?\s*(.+)$"), 2),
    (re.compile(r"^\d+\.\d+\.\d+\.\d+\.\d+\s+\S+"), 5),
    (re.compile(r"^\d+\.\d+\.\d+\.\d+\s+\S+"), 4),
    (re.compile(r"^\d+\.\d+\.\d+\s+\S+"), 3),
    (re.compile(r"^\d+\.\d+\s+\S+"), 2),
    (re.compile(r"^\d+\s+\S+"), 1),
    (re.compile(r"^[一二三四五六七八九十]+[、.．]\s*\S+"), 1),
    (re.compile(r"^（[一二三四五六七八九十]+）\s*\S+"), 3),
    (re.compile(r"^\([一二三四五六七八九十]+\)\s*\S+"), 3),
]

LIST_PATTERNS = [
    re.compile(r"^[a-zA-Z]\)\s*\S+"),
    re.compile(r"^\d+\)\s*\S+"),
    re.compile(r"^[（(]\d+[）)]\s*\S+"),
    re.compile(r"^[·•]\s*\S+"),
    re.compile(r"^[\-\uff0d\u2014]{1,2}\s*\S+"),
]

DEFAULT_TABLE_ROW_HEIGHT_CM = 0.69
DEFAULT_TABLE_ROW_HEIGHT_RULE = "at-least"


def skill_version() -> str:
    version_path = Path(__file__).resolve().parents[1] / "VERSION"
    if version_path.exists():
        return version_path.read_text(encoding="utf-8").strip()
    return "unknown"


def iter_blocks(doc: Document):
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def paragraph_has_graphics(paragraph: Paragraph) -> bool:
    return bool(
        paragraph._p.findall(".//" + qn("w:drawing"))
        or paragraph._p.findall(".//" + qn("w:pict"))
        or paragraph._p.findall(".//" + qn("w:object"))
    )


def set_run_fonts(run, east_asia="宋体", ascii_font="Times New Roman", size_pt=12, bold=None) -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)


def set_style_fonts(style, east_asia="宋体", ascii_font="Times New Roman", size_pt=12, bold=None) -> None:
    style.font.name = ascii_font
    style.font.size = Pt(size_pt)
    style.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)


def ensure_paragraph_style(doc: Document, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def set_spacing_xml(target, line: int | None = 300, line_rule: str = "auto", **attrs) -> None:
    p_pr = target._element.get_or_add_pPr() if hasattr(target, "_element") else target._p.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    for attr in [
        "before",
        "beforeLines",
        "beforeAutospacing",
        "after",
        "afterLines",
        "afterAutospacing",
        "line",
        "lineRule",
    ]:
        key = qn(f"w:{attr}")
        if key in spacing.attrib:
            del spacing.attrib[key]
    if line is not None:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), line_rule)
    for attr, value in attrs.items():
        if value is not None:
            spacing.set(qn(f"w:{attr}"), str(value))


def set_standard_spacing(target) -> None:
    set_spacing_xml(
        target,
        line=300,
        line_rule="auto",
        before=0,
        beforeLines=0,
        beforeAutospacing=0,
        after=0,
        afterLines=0,
        afterAutospacing=0,
    )


def set_heading_spacing(target, style_name: str | None) -> None:
    if style_name == "Heading 1":
        set_spacing_xml(
            target,
            line=300,
            line_rule="auto",
            before=50,
            beforeLines=50,
            beforeAutospacing=0,
            after=50,
            afterLines=50,
            afterAutospacing=0,
        )
    else:
        set_spacing_xml(
            target,
            line=300,
            line_rule="auto",
            beforeLines=0,
            beforeAutospacing=0,
            afterLines=0,
            afterAutospacing=0,
        )


def set_caption_spacing(target) -> None:
    set_spacing_xml(
        target,
        line=240,
        line_rule="auto",
        before=50,
        beforeLines=50,
        beforeAutospacing=0,
        after=50,
        afterLines=50,
        afterAutospacing=0,
    )


def set_table_body_spacing(target) -> None:
    set_spacing_xml(
        target,
        line=0,
        line_rule="atLeast",
        before=0,
        beforeLines=0,
        beforeAutospacing=0,
        after=0,
        afterLines=0,
        afterAutospacing=0,
    )


def set_note_spacing(target) -> None:
    set_spacing_xml(
        target,
        line=300,
        line_rule="auto",
        before=448,
        beforeAutospacing=0,
        after=0,
        afterLines=0,
        afterAutospacing=0,
    )


def set_numbered_note_spacing(target) -> None:
    set_spacing_xml(
        target,
        line=300,
        line_rule="auto",
        before=448,
        beforeAutospacing=0,
        after=0,
        afterLines=0,
        afterAutospacing=0,
    )


def set_formula_spacing(target) -> None:
    set_standard_spacing(target)


def set_toc_spacing(target) -> None:
    set_standard_spacing(target)


def ensure_fallback_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_fonts(normal, east_asia="宋体", size_pt=12, bold=False)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.13)
    set_standard_spacing(normal)

    title = ensure_paragraph_style(doc, "文档标题")
    set_style_fonts(title, east_asia="黑体", size_pt=14, bold=False)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.left_indent = Pt(0)
    title.paragraph_format.first_line_indent = Pt(0)
    set_standard_spacing(title)

    note = ensure_paragraph_style(doc, "3.1注-无编号注")
    set_style_fonts(note, east_asia="宋体", size_pt=10.5, bold=False)
    note.paragraph_format.left_indent = Cm(1.53)
    note.paragraph_format.first_line_indent = Cm(-0.74)
    set_note_spacing(note)

    numbered_note = ensure_paragraph_style(doc, "3.2注-有编号注")
    set_style_fonts(numbered_note, east_asia="宋体", size_pt=10.5, bold=False)
    numbered_note.paragraph_format.left_indent = Cm(1.81)
    numbered_note.paragraph_format.first_line_indent = Cm(-0.93)
    set_numbered_note_spacing(numbered_note)

    formula = ensure_paragraph_style(doc, "公式")
    set_style_fonts(formula, east_asia="宋体", size_pt=12, bold=False)
    formula.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula.paragraph_format.first_line_indent = Pt(0)
    set_formula_spacing(formula)

    table_body = ensure_paragraph_style(doc, "表正文")
    set_style_fonts(table_body, east_asia="宋体", size_pt=10.5, bold=False)
    table_body.paragraph_format.first_line_indent = Pt(0)
    set_table_body_spacing(table_body)

    caption = ensure_paragraph_style(doc, "Caption")
    set_style_fonts(caption, east_asia="黑体", size_pt=12, bold=False)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_caption_spacing(caption)

    heading_indents = {
        "Heading 1": 0.762,
        "Heading 2": 1.0142361111111111,
        "Heading 3": 1.27,
        "Heading 4": 1.524,
        "Heading 5": 1.778,
        "Heading 6": 2.032,
    }
    for name, indent_cm in heading_indents.items():
        style = ensure_paragraph_style(doc, name)
        set_style_fonts(style, east_asia="黑体", size_pt=12, bold=False)
        style.paragraph_format.left_indent = Cm(indent_cm)
        style.paragraph_format.first_line_indent = Cm(-indent_cm)
        set_heading_spacing(style, name)

    appendix_title = ensure_paragraph_style(doc, "附录标题")
    set_style_fonts(appendix_title, east_asia="黑体", size_pt=14, bold=False)
    appendix_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    appendix_title.paragraph_format.left_indent = Pt(0)
    appendix_title.paragraph_format.first_line_indent = Pt(0)
    set_standard_spacing(appendix_title)

    for name, indent_cm in [
        ("附录一级标题", 0.762),
        ("附录二级标题", 1.0142361111111111),
        ("附录三级标题", 1.27),
    ]:
        style = ensure_paragraph_style(doc, name)
        set_style_fonts(style, east_asia="黑体", size_pt=12, bold=False)
        style.paragraph_format.left_indent = Cm(indent_cm)
        style.paragraph_format.first_line_indent = Cm(-indent_cm)
        set_heading_spacing(style, "Heading 1" if name == "附录一级标题" else "Heading 2")

    for level, indent_cm in [(1, 0), (2, 0.85), (3, 1.69), (4, 2.54)]:
        style = ensure_paragraph_style(doc, f"TOC {level}")
        set_style_fonts(style, east_asia="宋体", size_pt=12, bold=False)
        style.paragraph_format.left_indent = Cm(indent_cm)
        style.paragraph_format.first_line_indent = Pt(0)
        set_toc_spacing(style)

    for name, left_cm, hanging_cm in [
        ("1.1一级列项-编号", 1.6474722222222222, -0.8008055555555555),
        ("1.2一级列项-无编号", 1.760361111111111, -0.8766527777777777),
        ("2.1二级列项-有编号", 2.4429861111111113, -0.7496527777777777),
        ("2.2二级列项-无编号", 2.573513888888889, -0.8801805555555555),
    ]:
        style = ensure_paragraph_style(doc, name)
        set_style_fonts(style, east_asia="宋体", size_pt=12, bold=False)
        style.paragraph_format.left_indent = Cm(left_cm)
        style.paragraph_format.first_line_indent = Cm(hanging_cm)
        set_standard_spacing(style)


def apply_page_setup(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)


def strip_manual_number(text: str) -> str:
    text = re.sub(r"^\d+(?:\.\d+)*\s+", "", text, count=1)
    text = re.sub(r"^第[一二三四五六七八九十百千万0-9]+[章节]\s*[:：、.\s]?", "", text, count=1)
    text = re.sub(r"^[一二三四五六七八九十]+[、.．]\s*", "", text, count=1)
    text = re.sub(r"^[（(][一二三四五六七八九十]+[）)]\s*", "", text, count=1)
    return text


def heading_level_from_text(text: str) -> int | None:
    if is_date_like_text(text) or is_caption_text(text) or is_toc_title(text):
        return None
    for pattern, level in HEADING_PATTERNS:
        if pattern.match(text):
            return level
    return None


def heading_style_for_level(level: int) -> str:
    return f"Heading {max(1, min(level, 5))}"


def strip_heading_marker(text: str) -> str:
    text = re.sub(r"^\d+(?:\.\d+)*\s+", "", text, count=1)
    text = re.sub(r"^第[一二三四五六七八九十百千万0-9]+[章节]\s*[:：、.\s]?", "", text, count=1)
    text = re.sub(r"^[一二三四五六七八九十]+[、.．]\s*", "", text, count=1)
    text = re.sub(r"^[（(][一二三四五六七八九十0-9]+[）)]\s*", "", text, count=1)
    return text


def looks_like_list_item(text: str) -> bool:
    return any(pattern.match(text) for pattern in LIST_PATTERNS)


def list_style_for_text(text: str) -> str:
    if re.match(r"^\d+\)\s*\S+|^[（(]\d+[）)]\s*\S+", text):
        return "2.1二级列项-有编号"
    if re.match(r"^[·•]\s*\S+", text):
        return "2.2二级列项-无编号"
    if re.match(r"^[\-\uff0d\u2014]{1,2}\s*\S+", text):
        return "1.2一级列项-无编号"
    return "1.1一级列项-编号"


def list_kind_for_text(text: str) -> str:
    if re.match(r"^\d+\)\s*\S+|^[（(]\d+[）)]\s*\S+", text):
        return "decimal"
    if re.match(r"^[·•]\s*\S+", text):
        return "bullet2"
    if re.match(r"^[\-\uff0d\u2014]{1,2}\s*\S+", text):
        return "dash"
    return "letter"


def is_formula_text(text: str) -> bool:
    return bool(re.match(r"^\(?\d+\)?$", text) or re.search(r"公式\s*\(?\d+\)?", text))


def is_appendix_title(text: str) -> bool:
    return bool(re.match(r"^（?(资料性|规范性)）?$", text) or re.match(r"^附\s*录\s*[A-ZＡ-Ｚ]?", text))


def is_date_like_text(text: str) -> bool:
    return bool(re.match(r"^\d{4}\s*年(?:\s*\d{1,2}\s*月)?(?:\s*\d{1,2}\s*日)?$", text))


def is_toc_title(text: str) -> bool:
    compact = re.sub(r"\s+", "", text)
    return compact in {"目录", "目次"}


def is_caption_text(text: str) -> bool:
    return bool(re.match(r"^[图表]\s*\d+(?:[-－.]\d+)*[\s　]+.+", text))


def is_front_matter_text(paragraph: Paragraph, text: str) -> bool:
    return bool(
        is_date_like_text(text)
        or is_toc_title(text)
        or (paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER and len(text) <= 80)
    )


def existing_heading_number(text: str) -> bool:
    return bool(
        re.match(r"^\d+(?:\.\d+)*\s+\S+", text)
        or re.match(r"^第[一二三四五六七八九十百千万0-9]+[章节]\s*[:：、.\s]?\S+", text)
        or re.match(r"^[一二三四五六七八九十]+[、.．]\s*\S+", text)
        or re.match(r"^[（(][一二三四五六七八九十0-9]+[）)]\s*\S+", text)
    )


def paragraph_num_info(paragraph: Paragraph) -> tuple[int | None, int | None]:
    p_pr = paragraph._p.pPr
    num_pr = p_pr.numPr if p_pr is not None and p_pr.numPr is not None else None
    if num_pr is None and paragraph.style is not None:
        style_p_pr = paragraph.style._element.pPr
        num_pr = style_p_pr.numPr if style_p_pr is not None and style_p_pr.numPr is not None else None
    if num_pr is None:
        return None, None
    ilvl = int(num_pr.ilvl.val) if num_pr.ilvl is not None and num_pr.ilvl.val is not None else None
    num_id = int(num_pr.numId.val) if num_pr.numId is not None and num_pr.numId.val is not None else None
    return ilvl, num_id


def paragraph_numbering_descriptor(paragraph: Paragraph) -> tuple[str | None, str | None]:
    ilvl, num_id = paragraph_num_info(paragraph)
    if ilvl is None or num_id is None:
        return None, None
    numbering = paragraph.part.numbering_part.element
    abstract_id = None
    for num in numbering.findall(qn("w:num")):
        if num.get(qn("w:numId")) != str(num_id):
            continue
        abstract_ref = num.find(qn("w:abstractNumId"))
        if abstract_ref is not None:
            abstract_id = abstract_ref.get(qn("w:val"))
        break
    if abstract_id is None:
        return None, None
    for abstract_num in numbering.findall(qn("w:abstractNum")):
        if abstract_num.get(qn("w:abstractNumId")) != abstract_id:
            continue
        for level in abstract_num.findall(qn("w:lvl")):
            if level.get(qn("w:ilvl")) != str(ilvl):
                continue
            num_fmt = level.find(qn("w:numFmt"))
            lvl_text = level.find(qn("w:lvlText"))
            return (
                num_fmt.get(qn("w:val")) if num_fmt is not None else None,
                lvl_text.get(qn("w:val")) if lvl_text is not None else None,
            )
    return None, None


def source_numbering_heading_level(paragraph: Paragraph) -> int | None:
    if not looks_like_visual_heading(paragraph):
        return None
    num_fmt, lvl_text = paragraph_numbering_descriptor(paragraph)
    if not num_fmt or not lvl_text:
        return None
    if num_fmt in {"chineseCounting", "chineseCountingThousand", "ideographDigital"} and "、" in lvl_text:
        return 1
    if num_fmt == "decimal" and re.fullmatch(r"%1[.．]", lvl_text):
        return 2
    return None


def heading_level_from_style(style_name: str) -> int | None:
    match = re.match(r"Heading\s+([1-6])$", style_name)
    if match:
        return int(match.group(1))
    return None


def resolved_heading_level(style: str | None, num_level: int | None, text: str) -> int:
    return heading_level_from_style(style or "") or heading_level_from_text(text) or ((num_level or 0) + 1)


def heading_number_source(num_id: int | None, num_level: int | None, text: str) -> str:
    if num_id is not None:
        return "docx-numbering"
    if existing_heading_number(text):
        return "docx-text"
    if num_level is not None:
        return "docx-style-numbering"
    return "docx-heading-style"


def strip_list_marker(text: str) -> str:
    return re.sub(r"^([a-zA-Z]\)|\d+\)|[（(]\d+[）)]|[·•]|[\-\uff0d\u2014]{1,2})\s*", "", text, count=1)


def list_level_from_text(text: str, fallback: int | None = None) -> int:
    if re.match(r"^\d+\)\s*\S+|^[（(]\d+[）)]\s*\S+", text):
        return 1
    return int(fallback or 0)


def next_numbering_id(numbering, tag_name: str, attr_name: str) -> int:
    values = []
    for element in numbering.findall(qn(tag_name)):
        raw = element.get(qn(attr_name))
        if raw is not None and raw.isdigit():
            values.append(int(raw))
    return (max(values) + 1) if values else 1


def append_text_child(parent, tag: str, attr: str, value: str):
    child = OxmlElement(tag)
    child.set(qn(attr), value)
    parent.append(child)
    return child


def append_level(parent, ilvl: int, num_fmt: str, lvl_text: str, left_twips: int, hanging_twips: int, style_id: str | None = None) -> None:
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), str(ilvl))
    append_text_child(level, "w:start", "w:val", "1")
    append_text_child(level, "w:numFmt", "w:val", num_fmt)
    append_text_child(level, "w:lvlText", "w:val", lvl_text)
    append_text_child(level, "w:lvlJc", "w:val", "left")
    if style_id:
        append_text_child(level, "w:pStyle", "w:val", style_id)
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left_twips))
    ind.set(qn("w:hanging"), str(hanging_twips))
    p_pr.append(ind)
    level.append(p_pr)
    parent.append(level)


def append_num(numbering, abstract_num_id: int, num_id: int, start_override: int | None = None) -> None:
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_num_id))
    num.append(abstract_ref)
    if start_override is not None:
        lvl_override = OxmlElement("w:lvlOverride")
        lvl_override.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:startOverride")
        start.set(qn("w:val"), str(start_override))
        lvl_override.append(start)
        num.append(lvl_override)
    numbering.append(num)


def new_num_for_abstract(doc: Document, abstract_num_id: int) -> int:
    numbering = doc.part.numbering_part.element
    num_id = next_numbering_id(numbering, "w:num", "w:numId")
    append_num(numbering, abstract_num_id, num_id, start_override=1)
    return num_id


def ensure_auto_numbering(doc: Document) -> dict:
    numbering = doc.part.numbering_part.element
    abstract_id = next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
    num_id = next_numbering_id(numbering, "w:num", "w:numId")

    heading_abs = OxmlElement("w:abstractNum")
    heading_abs.set(qn("w:abstractNumId"), str(abstract_id))
    append_text_child(heading_abs, "w:multiLevelType", "w:val", "multilevel")
    for ilvl, left in enumerate([432, 575, 720, 864, 1008, 1151]):
        tokens = ".".join(f"%{index}" for index in range(1, ilvl + 2))
        append_level(heading_abs, ilvl, "decimal", f"{tokens} ", left, left, f"Heading{ilvl + 1}")
    numbering.append(heading_abs)
    append_num(numbering, abstract_id, num_id)
    heading_num_id = num_id

    letter_abs_id = abstract_id + 1
    letter_num_id = num_id + 1
    letter_abs = OxmlElement("w:abstractNum")
    letter_abs.set(qn("w:abstractNumId"), str(letter_abs_id))
    append_text_child(letter_abs, "w:multiLevelType", "w:val", "singleLevel")
    append_level(letter_abs, 0, "lowerLetter", "%1)", 934, 454, None)
    numbering.append(letter_abs)
    append_num(numbering, letter_abs_id, letter_num_id)

    decimal_abs_id = abstract_id + 2
    decimal_num_id = num_id + 2
    decimal_abs = OxmlElement("w:abstractNum")
    decimal_abs.set(qn("w:abstractNumId"), str(decimal_abs_id))
    append_text_child(decimal_abs, "w:multiLevelType", "w:val", "singleLevel")
    append_level(decimal_abs, 0, "decimal", "%1)", 1385, 425, None)
    numbering.append(decimal_abs)
    append_num(numbering, decimal_abs_id, decimal_num_id)

    return {
        "heading": heading_num_id,
        "list_letter": letter_num_id,
        "list_decimal": decimal_num_id,
        "list_letter_abstract": letter_abs_id,
        "list_decimal_abstract": decimal_abs_id,
    }


def apply_numbering(paragraph: Paragraph, num_id: int, ilvl: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def paragraph_direct_format_score(paragraph: Paragraph) -> tuple[bool, float | None]:
    text = paragraph.text.strip()
    if not text:
        return False, None
    bold_count = 0
    size_values = []
    for run in paragraph.runs:
        if run.bold:
            bold_count += 1
        if run.font.size is not None:
            size_values.append(run.font.size.pt)
    mostly_bold = bool(paragraph.runs) and bold_count >= max(1, len(paragraph.runs) // 2)
    max_size = max(size_values) if size_values else None
    return mostly_bold, max_size


def looks_like_visual_heading(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    if not text or len(text) > 40:
        return False
    if is_date_like_text(text) or is_caption_text(text) or is_toc_title(text):
        return False
    if text.endswith(("。", "；", ";")):
        return False
    mostly_bold, max_size = paragraph_direct_format_score(paragraph)
    centered = paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    return mostly_bold or centered or (max_size is not None and max_size >= 14)


def clean_note_prefix(text: str) -> str:
    text = re.sub(r"^\*\*(备注|编写提示)：\*\*\s*", "", text)
    text = re.sub(r"^(备注|编写提示)：\s*", "", text)
    return text


def normalize_paragraph(paragraph: Paragraph, role: str, style_name: str | None = None) -> None:
    text = paragraph.text.strip()
    if not text:
        return
    if role == "title":
        paragraph.style = "文档标题"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        set_standard_spacing(paragraph)
        for run in paragraph.runs:
            set_run_fonts(run, east_asia="黑体", size_pt=14, bold=False)
    elif role == "heading":
        if style_name:
            paragraph.style = style_name
        set_heading_spacing(paragraph, style_name)
        for run in paragraph.runs:
            set_run_fonts(run, east_asia="黑体", size_pt=12, bold=False)
    elif role == "note":
        paragraph.style = "3.1注-无编号注"
        paragraph.paragraph_format.left_indent = Cm(1.53)
        paragraph.paragraph_format.first_line_indent = Cm(-0.74)
        set_note_spacing(paragraph)
        for run in paragraph.runs:
            set_run_fonts(run, east_asia="宋体", size_pt=10.5, bold=False)
    elif role == "numbered_note":
        paragraph.style = "3.2注-有编号注"
        paragraph.paragraph_format.left_indent = Cm(1.81)
        paragraph.paragraph_format.first_line_indent = Cm(-0.93)
        set_numbered_note_spacing(paragraph)
        for run in paragraph.runs:
            set_run_fonts(run, east_asia="宋体", size_pt=10.5, bold=False)
    elif role == "formula":
        paragraph.style = "公式"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        set_formula_spacing(paragraph)
        for run in paragraph.runs:
            set_run_fonts(run, east_asia="宋体", size_pt=12, bold=False)
    elif role == "appendix_title":
        paragraph.style = "附录标题"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        set_standard_spacing(paragraph)
        for run in paragraph.runs:
            set_run_fonts(run, east_asia="黑体", size_pt=14, bold=False)
    elif role == "caption":
        paragraph.style = "Caption"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_caption_spacing(paragraph)
        for run in paragraph.runs:
            set_run_fonts(run, east_asia="黑体", size_pt=12, bold=False)
    elif role == "list":
        if style_name:
            paragraph.style = style_name
        set_standard_spacing(paragraph)
        for run in paragraph.runs:
            set_run_fonts(run, east_asia="宋体", size_pt=12, bold=False)
    else:
        paragraph.style = "Normal"
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(1.13)
        set_standard_spacing(paragraph)
        for run in paragraph.runs:
            set_run_fonts(run, east_asia="宋体", size_pt=12, bold=None)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_autofit_to_window(table: Table) -> None:
    table.autofit = True
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "autofit")
    for cell in table._tbl.findall(".//" + qn("w:tc")):
        tc_pr = cell.find(qn("w:tcPr"))
        if tc_pr is None:
            continue
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is not None:
            tc_w.set(qn("w:type"), "auto")
            tc_w.set(qn("w:w"), "0")
        no_wrap = tc_pr.find(qn("w:noWrap"))
        if no_wrap is not None:
            tc_pr.remove(no_wrap)


def normalize_table(table: Table, row_height_cm: float, row_height_rule: str) -> None:
    table.style = "Table Grid"
    set_table_autofit_to_window(table)
    for ri, row in enumerate(table.rows):
        row.height = Cm(row_height_cm)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY if row_height_rule == "exact" else WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri == 0:
                set_cell_shading(cell, "DDEBF7")
            for paragraph in cell.paragraphs:
                paragraph.style = "表正文"
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = Pt(0)
                set_table_body_spacing(paragraph)
                for run in paragraph.runs:
                    rpr = run._element.rPr
                    if rpr is not None:
                        run._element.remove(rpr)


def new_report() -> dict:
    return {
        "skill_version": skill_version(),
        "inferred_headings": [],
        "suspect_visual_headings": [],
        "inferred_lists": [],
        "automatic_numbers": [],
        "ambiguous_short_paragraphs": [],
        "content_warnings": [],
        "tables_processed": 0,
        "semantic_object_splits": [],
        "mixed_text_graphic_paragraphs_split": [],
        "graphic_paragraphs_preserved": 0,
        "media_relationships_preserved": 0,
        "non_text_objects": {},
        "risk_warnings": [],
        "audit": {},
    }


def paragraph_num_id(paragraph: Paragraph) -> str | None:
    p_pr = paragraph._p.pPr
    num_pr = p_pr.numPr if p_pr is not None and p_pr.numPr is not None else None
    if num_pr is None or num_pr.numId is None or num_pr.numId.val is None:
        return None
    return str(num_pr.numId.val)


def num_has_start_override(doc: Document, num_id: str) -> bool:
    numbering = doc.part.numbering_part.element
    for num in numbering.findall(qn("w:num")):
        if num.get(qn("w:numId")) != num_id:
            continue
        start = num.find(".//" + qn("w:startOverride"))
        return start is not None and start.get(qn("w:val")) == "1"
    return False


def audit_document(doc: Document, row_height_cm: float, row_height_rule: str) -> dict:
    audit = {
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "heading_sequence": [],
        "list_restart_groups": [],
        "table_paragraphs_not_table_body": [],
        "table_rows_bad_height": [],
        "table_cells_may_clip": [],
        "markdown_residue": [],
        "heading_paragraphs_without_numbering": [],
        "ordered_list_nums_without_restart": [],
    }
    seen_list_num_ids: set[str] = set()
    for idx, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph.text.strip()
        if not text:
            continue
        if "**" in text or re.match(r"^#{1,6}\s+", text):
            audit["markdown_residue"].append({"paragraph": idx, "text": text[:120]})
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name.startswith("Heading"):
            heading_level = heading_level_from_style(style_name)
            _, num_id = paragraph_num_info(paragraph)
            audit["heading_sequence"].append(
                {"paragraph": idx, "level": heading_level, "num_id": num_id, "text": text[:120]}
            )
            if num_id is None:
                audit["heading_paragraphs_without_numbering"].append(
                    {"paragraph": idx, "style": style_name, "text": text[:120]}
                )
        if "列项-编号" in style_name:
            num_id_str = paragraph_num_id(paragraph)
            if num_id_str is not None and num_id_str not in seen_list_num_ids:
                seen_list_num_ids.add(num_id_str)
                has_restart = num_has_start_override(doc, num_id_str)
                audit["list_restart_groups"].append(
                    {
                        "paragraph": idx,
                        "style": style_name,
                        "num_id": num_id_str,
                        "restart_at_one": has_restart,
                        "text": text[:120],
                    }
                )
                if not has_restart:
                    audit["ordered_list_nums_without_restart"].append(
                        {"paragraph": idx, "style": style_name, "num_id": num_id_str, "text": text[:120]}
                    )
    for table_idx, table in enumerate(doc.tables, 1):
        for row_idx, row in enumerate(table.rows, 1):
            if row.height is None or abs(row.height.cm - row_height_cm) > 0.02:
                audit["table_rows_bad_height"].append({"table": table_idx, "row": row_idx})
            for cell in row.cells:
                cell_text = "\n".join(paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip())
                if row_height_rule == "exact" and len(cell_text) > 45:
                    audit["table_cells_may_clip"].append(
                        {"table": table_idx, "row": row_idx, "text": cell_text[:120]}
                    )
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip() and paragraph.style.name != "表正文":
                        audit["table_paragraphs_not_table_body"].append(
                            {"table": table_idx, "row": row_idx, "style": paragraph.style.name, "text": paragraph.text[:80]}
                        )
    return audit


def collect_content_warnings(doc: Document) -> list[dict]:
    texts = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    joined = "\n".join(texts)
    warnings = []
    if any(text in {"目次", "目录"} for text in texts):
        warnings.append({"type": "toc", "message": "发现目录或目次，应确认目录层级缩进、页码和域更新状态。"})
    if "引用文件" in joined or "依据文件" in joined:
        warnings.append({"type": "references", "message": "发现引用文件或依据文件，应人工核对排序、标准号空格、正文引用对应关系。"})
    if "术语" in joined or "缩略语" in joined:
        warnings.append({"type": "terms", "message": "发现术语或缩略语章节，应人工核对术语定义和缩略语排序。"})
    if "公式" in joined or any(re.match(r"^\(?\d+\)?$", text) for text in texts):
        warnings.append({"type": "formula", "message": "发现公式或公式编号，应人工核对公式居中、编号右对齐和全文连续编号。"})
    if any(text.startswith("附录") or re.match(r"^（?(资料性|规范性)）?$", text) for text in texts):
        warnings.append({"type": "appendix", "message": "发现附录内容，应人工核对附录标题、附录编号、附录目录显示和页码。"})
    if any("图" in text and "注" in text for text in texts) or any("脚注" in text for text in texts):
        warnings.append({"type": "figure_table_notes", "message": "发现图注、表注或脚注相关内容，应人工核对其位置和注样式。"})
    return warnings


def scan_non_text_objects(src: Path) -> dict:
    counts = {
        "media_files": 0,
        "drawings": 0,
        "legacy_pictures": 0,
        "text_boxes": 0,
        "equations": 0,
        "fields": 0,
        "footnote_refs": 0,
        "comment_ranges": 0,
        "tracked_insertions": 0,
        "tracked_deletions": 0,
        "headers": 0,
        "footers": 0,
    }
    if src.suffix.lower() != ".docx":
        return counts
    try:
        with zipfile.ZipFile(src) as zf:
            names = zf.namelist()
            xml_names = [name for name in names if name.startswith("word/") and name.endswith(".xml")]
            counts["media_files"] = len([name for name in names if name.startswith("word/media/") and not name.endswith("/")])
            counts["headers"] = len([name for name in names if name.startswith("word/header") and name.endswith(".xml")])
            counts["footers"] = len([name for name in names if name.startswith("word/footer") and name.endswith(".xml")])
            for name in xml_names:
                data = zf.read(name).decode("utf-8", errors="ignore")
                counts["drawings"] += data.count("<w:drawing")
                counts["legacy_pictures"] += data.count("<w:pict")
                counts["text_boxes"] += data.count("<w:txbxContent")
                counts["equations"] += data.count("<m:oMath")
                counts["fields"] += data.count("<w:fldChar")
                counts["footnote_refs"] += data.count("<w:footnoteReference")
                counts["comment_ranges"] += data.count("<w:commentRangeStart")
                counts["tracked_insertions"] += len(re.findall(r"<w:ins(?:\s|>)", data))
                counts["tracked_deletions"] += len(re.findall(r"<w:del(?:\s|>)", data))
    except zipfile.BadZipFile:
        pass
    return counts


def add_risk_warnings(report: dict, row_height_rule: str) -> None:
    non_text = report.get("non_text_objects", {})
    risky_objects = {key: value for key, value in non_text.items() if value}
    source_media = int(non_text.get("media_files") or 0)
    preserved_media = int(report.get("media_relationships_preserved") or 0)
    report["media_preservation_ratio"] = (preserved_media / source_media) if source_media else 1.0
    if source_media > preserved_media:
        report["risk_warnings"].append(
            {
                "type": "media_not_fully_preserved",
                "message": "Source document contains media files that were not all preserved in the output.",
                "source_media_files": source_media,
                "preserved_media_relationships": preserved_media,
            }
        )
    if risky_objects:
        report["risk_warnings"].append(
            {
                "type": "non_text_objects",
                "message": "Source document contains objects that may not be fully rebuilt by text normalization.",
                "objects": risky_objects,
            }
        )
    clipped_cells = report.get("audit", {}).get("table_cells_may_clip", [])
    if row_height_rule == "exact" and clipped_cells:
        report["risk_warnings"].append(
            {
                "type": "table_row_height",
                "message": "Fixed table row height may clip long cell text. Render and review affected tables.",
                "count": len(clipped_cells),
            }
        )
    list_without_restart = report.get("audit", {}).get("ordered_list_nums_without_restart", [])
    if list_without_restart:
        report["risk_warnings"].append(
            {
                "type": "ordered_list_restart",
                "message": "Some ordered list numbering instances do not explicitly restart at 1.",
                "count": len(list_without_restart),
            }
        )


def write_markdown_report(report: dict, path: Path) -> None:
    lines = ["# WX 文档格式转换报告", ""]
    audit = report.get("audit", {})
    lines.extend([
        "## 概览",
        f"- Skill 版本：{report.get('skill_version', 'unknown')}",
        f"- 段落数：{audit.get('paragraph_count', 0)}",
        f"- 表格数：{audit.get('table_count', 0)}",
        f"- 已处理表格数：{report.get('tables_processed', 0)}",
        f"- 已拆分语义对象混合段落数：{len(report.get('semantic_object_splits', []))}",
        f"- 已保留图片段落数：{report.get('graphic_paragraphs_preserved', 0)}",
        f"- 已保留媒体关系数：{report.get('media_relationships_preserved', 0)}",
        f"- 媒体保留比例：{report.get('media_preservation_ratio', 1.0):.2f}",
        "",
    ])
    for title, key in [
        ("推断标题", "inferred_headings"),
        ("疑似视觉标题", "suspect_visual_headings"),
        ("推断列项", "inferred_lists"),
        ("自动编号", "automatic_numbers"),
        ("已拆分语义对象混合段落", "semantic_object_splits"),
        ("模糊短段落", "ambiguous_short_paragraphs"),
    ]:
        items = report.get(key, [])
        lines.append(f"## {title}")
        if not items:
            lines.append("- 无")
        else:
            for item in items[:50]:
                lines.append(f"- {item}")
        lines.append("")
    lines.append("## 非文本对象")
    for key, value in report.get("non_text_objects", {}).items():
        lines.append(f"- {key}：{value}")
    lines.append("")
    lines.append("## 风险提示")
    risk_warnings = report.get("risk_warnings", [])
    if not risk_warnings:
        lines.append("- 无")
    else:
        for warning in risk_warnings:
            lines.append(f"- {warning}")
    lines.append("")
    lines.append("## 内容型复核提示")
    content_warnings = report.get("content_warnings", [])
    if not content_warnings:
        lines.append("- 无")
    else:
        for warning in content_warnings:
            lines.append(f"- {warning}")
    lines.append("")
    lines.append("## 审计问题")
    problem_keys = [
        "table_paragraphs_not_table_body",
        "table_rows_bad_height",
        "table_cells_may_clip",
        "markdown_residue",
        "heading_paragraphs_without_numbering",
        "ordered_list_nums_without_restart",
    ]
    has_problem = False
    for key in problem_keys:
        values = audit.get(key, [])
        if values:
            has_problem = True
            lines.append(f"### {key}")
            for value in values[:50]:
                lines.append(f"- {value}")
            lines.append("")
    if not has_problem:
        lines.append("- 未发现结构化审计问题")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def add_paragraph(doc: Document, text: str, style: str | None = None, role: str = "body") -> Paragraph:
    paragraph = doc.add_paragraph(text, style=style)
    normalize_paragraph(paragraph, role, style)
    return paragraph


def is_md_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    line = lines[index].strip()
    sep = lines[index + 1].strip()
    return line.startswith("|") and line.endswith("|") and re.match(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", sep)


def parse_md_table(lines: list[str], index: int) -> tuple[list[list[str]], int]:
    rows = []
    current = index
    while current < len(lines):
        line = lines[current].strip()
        if not (line.startswith("|") and line.endswith("|")):
            break
        if current == index + 1:
            current += 1
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        rows.append(cells)
        current += 1
    return rows, current


def add_table_from_rows(doc: Document, rows: list[list[str]], row_height_cm: float, row_height_rule: str) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            table.rows[r_idx].cells[c_idx].text = row[c_idx] if c_idx < len(row) else ""
    normalize_table(table, row_height_cm, row_height_rule)


def convert_md(src: Path, doc: Document, report: dict, row_height_cm: float, row_height_rule: str, numbering_ids: dict) -> None:
    lines = src.read_text(encoding="utf-8").splitlines()
    index = 0
    active_list_nums: dict[int, int] = {}
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()
        if not line:
            index += 1
            continue
        if is_md_table_start(lines, index):
            rows, next_index = parse_md_table(lines, index)
            add_table_from_rows(doc, rows, row_height_cm, row_height_rule)
            report["tables_processed"] += 1
            index = next_index
            continue
        match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if match:
            level = len(match.group(1))
            text = match.group(2).strip()
            role = "title" if level == 1 else "heading"
            paragraph = add_paragraph(doc, strip_heading_marker(text), STYLE_BY_MD_LEVEL.get(level, "Heading 5"), role)
            if role == "heading":
                apply_numbering(paragraph, numbering_ids["heading"], min(level - 2, 5))
                report["automatic_numbers"].append({"type": "heading", "text": strip_heading_marker(text), "source": "md-heading"})
            active_list_nums = {}
            index += 1
            continue
        inferred_level = heading_level_from_text(line)
        if inferred_level is not None:
            clean_text = strip_heading_marker(line)
            paragraph = add_paragraph(doc, clean_text, heading_style_for_level(inferred_level), "heading")
            apply_numbering(paragraph, numbering_ids["heading"], inferred_level - 1)
            report["inferred_headings"].append({"text": line, "level": inferred_level, "source": "md-text"})
            report["automatic_numbers"].append({"type": "heading", "text": clean_text, "source": "md-text"})
            active_list_nums = {}
            index += 1
            continue
        role = "note" if line.startswith(("**备注：**", "**编写提示：**", "备注：", "编写提示：")) else "body"
        if line.startswith(("注1：", "注2：", "注3：", "注4：", "注5：")):
            role = "numbered_note"
        elif is_formula_text(line):
            role = "formula"
        elif is_appendix_title(line):
            role = "appendix_title"
        if role == "body" and looks_like_list_item(line):
            kind = list_kind_for_text(line)
            list_level = 1 if kind in {"decimal", "bullet2"} else 0
            if kind == "dash":
                style_name = "1.2一级列项-无编号"
            elif kind == "bullet2":
                style_name = "2.2二级列项-无编号"
            else:
                style_name = "2.1二级列项-有编号" if list_level else "1.1一级列项-编号"
            paragraph = add_paragraph(doc, strip_list_marker(line), style_name, "list")
            if kind in {"letter", "decimal"} and list_level not in active_list_nums:
                abstract_key = "list_decimal_abstract" if list_level else "list_letter_abstract"
                active_list_nums[list_level] = new_num_for_abstract(doc, numbering_ids[abstract_key])
            if kind in {"letter", "decimal"}:
                apply_numbering(paragraph, active_list_nums[list_level], 0)
            report["inferred_lists"].append({"text": line, "source": "md-text"})
            if kind in {"letter", "decimal"}:
                report["automatic_numbers"].append({"type": "list", "text": strip_list_marker(line), "source": "md-text"})
        else:
            add_paragraph(doc, clean_note_prefix(line), role=role)
            active_list_nums = {}
        index += 1


def infer_docx_role(paragraph: Paragraph, strict_normalize: bool, report: dict) -> tuple[str, str | None, str]:
    text = paragraph.text.strip()
    style_name = paragraph.style.name if paragraph.style is not None else ""
    if is_caption_text(text):
        return text, "Caption", "caption"
    if is_date_like_text(text) or is_toc_title(text):
        return text, None, "body"
    if style_name == "文档标题":
        return text, "文档标题", "title"
    if style_name.startswith("Heading"):
        return text, style_name, "heading"
    if style_name == "List Paragraph":
        report["inferred_lists"].append({"text": text, "source": "docx-style"})
        return text, "1.1一级列项-编号", "list"
    inferred_level = heading_level_from_text(text)
    if inferred_level is not None:
        report["inferred_headings"].append({"text": text, "level": inferred_level, "source": "docx-text"})
        return text, heading_style_for_level(inferred_level), "heading"
    if strict_normalize:
        numbering_level = source_numbering_heading_level(paragraph)
        if numbering_level is not None:
            report["inferred_headings"].append({"text": text, "level": numbering_level, "source": "docx-numbering-format"})
            return text, heading_style_for_level(numbering_level), "heading"
    if strict_normalize and looks_like_visual_heading(paragraph):
        report["suspect_visual_headings"].append({"text": text, "assigned_level": 2})
        return text, "Heading 2", "heading"
    if text.startswith(("备注：", "编写提示：", "【备注提示】", "【编写样例】")):
        return clean_note_prefix(text), None, "note"
    if re.match(r"^注\s*\d+[：:]", text):
        return text, None, "numbered_note"
    if is_formula_text(text):
        return text, "公式", "formula"
    if is_appendix_title(text):
        return text, "附录标题", "appendix_title"
    if looks_like_list_item(text):
        report["inferred_lists"].append({"text": text, "source": "docx-text"})
        return text, list_style_for_text(text), "list"
    if strict_normalize and len(text) <= 30 and not text.endswith(("。", "；", ";", "，", ",")):
        report["ambiguous_short_paragraphs"].append(text)
    return text, None, "body"


def append_table_clone(doc: Document, table: Table) -> Table:
    body = doc.element.body
    sect_pr = body[-1] if len(body) and body[-1].tag == qn("w:sectPr") else None
    new_el = deepcopy(table._tbl)
    if sect_pr is not None:
        body.insert(len(body) - 1, new_el)
    else:
        body.append(new_el)
    return Table(new_el, doc)


def clone_related_media_rels(src_paragraph: Paragraph, dst_doc: Document, new_el) -> int:
    remapped: dict[str, str] = {}
    rel_attrs = [qn("r:embed"), qn("r:link"), qn("r:id")]
    media_count = 0
    for element in new_el.iter():
        for attr in rel_attrs:
            old_rid = element.get(attr)
            if old_rid is None or old_rid in remapped:
                continue
            rel = src_paragraph.part.rels.get(old_rid)
            if rel is None:
                continue
            if rel.is_external:
                new_rid = dst_doc.part.relate_to(rel.target_ref, rel.reltype, is_external=True)
            else:
                new_rid = dst_doc.part.relate_to(rel.target_part, rel.reltype)
                if rel.reltype == RT.IMAGE:
                    media_count += 1
            remapped[old_rid] = new_rid
        for attr in rel_attrs:
            old_rid = element.get(attr)
            if old_rid in remapped:
                element.set(attr, remapped[old_rid])
    return media_count


def append_paragraph_clone(doc: Document, paragraph: Paragraph) -> tuple[Paragraph, int]:
    body = doc.element.body
    sect_pr = body[-1] if len(body) and body[-1].tag == qn("w:sectPr") else None
    new_el = deepcopy(paragraph._p)
    media_count = clone_related_media_rels(paragraph, doc, new_el)
    if sect_pr is not None:
        body.insert(len(body) - 1, new_el)
    else:
        body.append(new_el)
    new_paragraph = Paragraph(new_el, doc)
    if paragraph_has_graphics(new_paragraph):
        normalize_graphics_paragraph(new_paragraph)
    return new_paragraph, media_count


def normalize_graphics_paragraph(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.pPr
    if p_pr is not None:
        paragraph._p.remove(p_pr)
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1


def element_has_graphics(element) -> bool:
    return bool(
        element.findall(".//" + qn("w:drawing"))
        or element.findall(".//" + qn("w:pict"))
        or element.findall(".//" + qn("w:object"))
    )


def append_graphics_only_paragraph_clone(doc: Document, paragraph: Paragraph) -> tuple[Paragraph, int]:
    body = doc.element.body
    sect_pr = body[-1] if len(body) and body[-1].tag == qn("w:sectPr") else None
    new_el = deepcopy(paragraph._p)
    for child in list(new_el):
        if child.tag in {qn("w:r"), qn("w:hyperlink")} and not element_has_graphics(child):
            new_el.remove(child)
    media_count = clone_related_media_rels(paragraph, doc, new_el)
    if sect_pr is not None:
        body.insert(len(body) - 1, new_el)
    else:
        body.append(new_el)
    new_paragraph = Paragraph(new_el, doc)
    normalize_graphics_paragraph(new_paragraph)
    return new_paragraph, media_count


def emit_normalized_docx_text(
    doc: Document,
    text: str,
    style: str | None,
    role: str,
    num_level: int | None,
    num_id: int | None,
    report: dict,
    numbering_ids: dict,
    active_list_nums: dict[int, int],
) -> dict[int, int]:
    if role == "heading":
        heading_level = resolved_heading_level(style, num_level, text)
        number_source = heading_number_source(num_id, num_level, text)
        clean_text = strip_heading_marker(text)
        paragraph = add_paragraph(doc, clean_text, style, role)
        apply_numbering(paragraph, numbering_ids["heading"], heading_level - 1)
        report["automatic_numbers"].append(
            {"type": "heading", "text": clean_text, "level": heading_level, "source": number_source}
        )
        return {}

    if role == "list":
        kind = list_kind_for_text(text)
        list_level = 1 if kind in {"decimal", "bullet2"} else int(num_level or 0)
        clean_text = strip_list_marker(text)
        if kind == "dash":
            style = "1.2一级列项-无编号"
        elif kind == "bullet2":
            style = "2.2二级列项-无编号"
        else:
            style = "2.1二级列项-有编号" if list_level else "1.1一级列项-编号"
        paragraph = add_paragraph(doc, clean_text, style, role)
        if kind in {"letter", "decimal"} and list_level not in active_list_nums:
            abstract_key = "list_decimal_abstract" if list_level else "list_letter_abstract"
            active_list_nums[list_level] = new_num_for_abstract(doc, numbering_ids[abstract_key])
        if kind in {"letter", "decimal"}:
            apply_numbering(paragraph, active_list_nums[list_level], 0)
            report["automatic_numbers"].append(
                {"type": "list", "text": clean_text, "source": "docx-numbering" if num_id is not None else "docx-text"}
            )
        return active_list_nums

    add_paragraph(doc, text, style, role)
    return {}


def convert_docx(src: Path, doc: Document, row_height_cm: float, row_height_rule: str, strict_normalize: bool, report: dict, numbering_ids: dict) -> None:
    src_doc = Document(src)
    seen_content = False
    structural_started = False
    active_list_nums: dict[int, int] = {}
    for block in iter_blocks(src_doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if paragraph_has_graphics(block):
                if text:
                    num_level, num_id = paragraph_num_info(block)
                    inferred_text, style, role = infer_docx_role(block, strict_normalize, report)
                    active_list_nums = emit_normalized_docx_text(
                        doc, inferred_text, style, role, num_level, num_id, report, numbering_ids, active_list_nums
                    )
                    _, media_count = append_graphics_only_paragraph_clone(doc, block)
                    split_record = {"text": inferred_text, "role": role}
                    if role == "heading":
                        split_record["level"] = resolved_heading_level(style, num_level, inferred_text)
                    report["semantic_object_splits"].append(split_record)
                    report["mixed_text_graphic_paragraphs_split"].append(split_record)
                    report["graphic_paragraphs_preserved"] += 1
                    report["media_relationships_preserved"] += media_count
                    active_list_nums = {}
                    seen_content = True
                    continue
                _, media_count = append_paragraph_clone(doc, block)
                report["graphic_paragraphs_preserved"] += 1
                report["media_relationships_preserved"] += media_count
                active_list_nums = {}
                seen_content = True
                continue
            if not text:
                continue
            if strict_normalize and not structural_started and seen_content and is_front_matter_text(block, text):
                add_paragraph(doc, text, role="body")
                active_list_nums = {}
                continue
            if strict_normalize and not seen_content and looks_like_visual_heading(block):
                report["suspect_visual_headings"].append({"text": text, "assigned_level": "title"})
                add_paragraph(doc, text, "文档标题", "title")
                seen_content = True
                continue
            num_level, num_id = paragraph_num_info(block)
            text, style, role = infer_docx_role(block, strict_normalize, report)
            active_list_nums = emit_normalized_docx_text(
                doc, text, style, role, num_level, num_id, report, numbering_ids, active_list_nums
            )
            if role == "heading":
                structural_started = True
            seen_content = True
        else:
            new_table = append_table_clone(doc, block)
            normalize_table(new_table, row_height_cm, row_height_rule)
            report["tables_processed"] += 1
            active_list_nums = {}
            seen_content = True


def main() -> None:
    parser = argparse.ArgumentParser(description="Apply built-in WX Word formatting to MD or DOCX.")
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--strict-normalize", action=argparse.BooleanOptionalAction, default=True, help="Infer headings and lists from text and direct formatting, then rebuild normalized output.")
    parser.add_argument("--report", type=Path, default=None, help="Optional JSON report path for inferred structures and audit results.")
    parser.add_argument("--report-md", type=Path, default=None, help="Optional Markdown report path for human-readable audit results.")
    parser.add_argument("--table-row-height-cm", type=float, default=DEFAULT_TABLE_ROW_HEIGHT_CM)
    parser.add_argument("--table-row-height-rule", choices=["exact", "at-least"], default=DEFAULT_TABLE_ROW_HEIGHT_RULE)
    parser.add_argument("--fail-on-risk", action="store_true", help="Exit with an error if conversion risk warnings are detected.")
    args = parser.parse_args()

    maybe_reexec_with_skill_venv()

    if DOCX_IMPORT_ERROR is not None:
        if args.input.suffix.lower() == ".docx":
            script_dir = Path(__file__).resolve().parent
            if str(script_dir) not in sys.path:
                sys.path.insert(0, str(script_dir))
            from format_docx_ooxml import convert_docx_ooxml

            convert_docx_ooxml(args.input, args.output, args.report, args.report_md)
            print(args.output)
            return
        raise SystemExit(
            "python-docx or lxml could not be imported, and the input is not DOCX.\n"
            "Run scripts/bootstrap_macos_lxml.sh to prepare an isolated best-effect environment.\n"
            "You can also set WX_DOC_FORMAT_VENV to a venv that contains python-docx and lxml.\n"
            "Markdown input requires python-docx and lxml to create a new Word document.\n"
            f"Original error: {DOCX_IMPORT_ERROR}"
        )

    out_doc = Document()
    ensure_fallback_styles(out_doc)
    numbering_ids = ensure_auto_numbering(out_doc)
    apply_page_setup(out_doc)
    report = new_report()
    report["non_text_objects"] = scan_non_text_objects(args.input)

    suffix = args.input.suffix.lower()
    if suffix in {".md", ".markdown"}:
        convert_md(args.input, out_doc, report, args.table_row_height_cm, args.table_row_height_rule, numbering_ids)
    elif suffix == ".docx":
        convert_docx(args.input, out_doc, args.table_row_height_cm, args.table_row_height_rule, args.strict_normalize, report, numbering_ids)
    else:
        raise SystemExit(f"Unsupported input type: {args.input.suffix}")

    for table in out_doc.tables:
        normalize_table(table, args.table_row_height_cm, args.table_row_height_rule)

    report["audit"] = audit_document(out_doc, args.table_row_height_cm, args.table_row_height_rule)
    report["content_warnings"] = collect_content_warnings(out_doc)
    add_risk_warnings(report, args.table_row_height_rule)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    out_doc.save(args.output)
    if args.report is not None:
        args.report.parent.mkdir(parents=True, exist_ok=True)
        args.report.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    if args.report_md is not None:
        write_markdown_report(report, args.report_md)
    if args.fail_on_risk and report["risk_warnings"]:
        raise SystemExit("Conversion completed with risk warnings. Review the report before delivery.")
    print(args.output)


if __name__ == "__main__":
    main()
