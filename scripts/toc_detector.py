"""Detect and select source DOCX table-of-contents regions.

The detector runs before semantic AST construction.  It never mutates the
source document.  Selected regions are represented by stable source block
positions and are consumed by both the AST parser and the direct renderer.
"""

from __future__ import annotations

import re
from typing import Any

from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


_TOC_TITLES = {"目录", "目次", "目  次", "目  录"}
_NUMBERED_ENTRY_RE = re.compile(r"^\s*\d+(?:\.\d+)*[\s、.．]+\S+")
_TRAILING_PAGE_RE = re.compile(r"[\t.·… ]+\d{1,4}\s*$")
_LEADING_NUMBER_RE = re.compile(r"^\s*\d+(?:\.\d+)*[\s、.．]+")


def _style_name(paragraph: Paragraph) -> str:
    try:
        return paragraph.style.name or ""
    except Exception:
        return ""


def _has_page_boundary(paragraph: Paragraph) -> bool:
    element = paragraph._p
    for br in element.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    if next(element.iter(qn("w:lastRenderedPageBreak")), None) is not None:
        return True
    if next(element.iter(qn("w:sectPr")), None) is not None:
        return True
    return False


def _field_instruction(paragraph: Paragraph) -> str:
    values = [node.text or "" for node in paragraph._p.iter(qn("w:instrText"))]
    for node in paragraph._p.iter(qn("w:fldSimple")):
        values.append(node.get(qn("w:instr"), ""))
    return " ".join(values).strip()


def _normalize_entry_text(text: str) -> str:
    value = _TRAILING_PAGE_RE.sub("", text.strip())
    value = _LEADING_NUMBER_RE.sub("", value)
    return re.sub(r"\s+", "", value)


def _record_source_blocks(doc) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    position = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, doc)
            text = paragraph.text.strip()
            style = _style_name(paragraph)
            instruction = _field_instruction(paragraph)
            records.append({
                "position": position,
                "kind": "paragraph",
                "text": text,
                "style": style,
                "is_heading_style": style.lower().startswith("heading"),
                "is_toc_style": style.lower().startswith("toc"),
                "is_toc_title": text in _TOC_TITLES,
                "has_page_boundary": _has_page_boundary(paragraph),
                "has_toc_field": bool(re.search(r"\bTOC\b", instruction, re.I)),
                "has_pageref": bool(re.search(r"\bPAGEREF\b", instruction, re.I)),
                "has_hyperlink": next(child.iter(qn("w:hyperlink")), None) is not None,
            })
            position += 1
        elif child.tag == qn("w:tbl"):
            table = Table(child, doc)
            items = [
                paragraph.text.strip()
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
                if paragraph.text.strip()
            ]
            text = " ".join(items)
            records.append({
                "position": position,
                "kind": "table",
                "text": text,
                "style": "",
                "is_heading_style": False,
                "is_toc_style": False,
                "is_toc_title": False,
                "has_page_boundary": False,
                "has_toc_field": False,
                "has_pageref": False,
                "has_hyperlink": False,
                "items": items,
            })
            position += 1
    return records


def _looks_like_entry(record: dict[str, Any]) -> bool:
    text = record.get("text", "")
    if not text:
        return False
    if record.get("is_toc_style") or record.get("has_pageref"):
        return True
    if _TRAILING_PAGE_RE.search(text):
        return True
    if len(text) <= 100 and _NUMBERED_ENTRY_RE.match(text):
        return True
    return False


def _entry_texts(record: dict[str, Any]) -> list[str]:
    if record.get("kind") != "table":
        return [record.get("text", "")] if _looks_like_entry(record) else []
    result: list[str] = []
    for text in record.get("items", []):
        probe = dict(record, kind="paragraph", text=text)
        if _looks_like_entry(probe):
            result.append(text)
    return result


def _later_heading_texts(records: list[dict[str, Any]], start_index: int) -> set[str]:
    result: set[str] = set()
    for record in records[start_index:]:
        if record.get("is_heading_style"):
            normalized = _normalize_entry_text(record.get("text", ""))
            if normalized:
                result.add(normalized)
    return result


def _candidate_from_start(
    records: list[dict[str, Any]],
    start_index: int,
    candidate_number: int,
) -> dict[str, Any] | None:
    start = records[start_index]
    scan_limit = min(len(records), start_index + 100)
    boundary_index: int | None = start_index if start.get("has_page_boundary") else None

    if boundary_index is None:
        for index in range(start_index + 1, scan_limit):
            if records[index].get("has_page_boundary"):
                boundary_index = index
                break

    provisional_end = boundary_index if boundary_index is not None else scan_limit - 1
    entry_scan_start = start_index if start.get("is_toc_style") else start_index + 1
    entry_records = [
        record for record in records[entry_scan_start:provisional_end + 1]
        if _entry_texts(record)
    ]
    entry_texts = [text for record in entry_records for text in _entry_texts(record)]

    normalized_entries = {
        _normalize_entry_text(text)
        for text in entry_texts
        if _normalize_entry_text(text)
    }
    later_headings = _later_heading_texts(records, provisional_end + 1)
    repeated_entries = sorted(normalized_entries & later_headings)

    if boundary_index is None:
        for index in range(start_index + 1, scan_limit):
            record = records[index]
            normalized = _normalize_entry_text(record.get("text", ""))
            if record.get("is_heading_style") and normalized in normalized_entries:
                provisional_end = index - 1
                break

    evidence: list[str] = []
    score = 0.0
    if start.get("is_toc_title"):
        evidence.append("toc_title")
        score += 0.25
    if start.get("has_toc_field"):
        evidence.append("toc_field")
        score += 0.65
    if start.get("is_toc_style"):
        evidence.append("toc_style")
        score += 0.35
    if start_index <= min(120, max(12, len(records) // 3)):
        evidence.append("front_matter_position")
        score += 0.10
    if len(entry_texts) >= 2:
        evidence.append("entry_cluster")
        score += 0.20
    if len(entry_texts) >= 5:
        evidence.append("large_entry_cluster")
        score += 0.10
    if boundary_index is not None:
        evidence.append("page_boundary")
        score += 0.20
    if len(repeated_entries) >= 2:
        evidence.append("entries_repeat_as_headings")
        score += 0.25
    if any(record.get("has_pageref") or record.get("has_hyperlink") for record in entry_records):
        evidence.append("field_or_hyperlink_entries")
        score += 0.20

    score = min(round(score, 3), 1.0)
    if score < 0.40:
        return None

    end_position = records[provisional_end]["position"]
    preview = [text[:120] for text in entry_texts[:12]]
    return {
        "id": f"toc_candidate_{candidate_number}",
        "block_type": "unknown",
        "role": "toc_candidate",
        "text": start.get("text", "") or "Word TOC field",
        "candidate_id": f"toc_candidate_{candidate_number}",
        "start_source_position": start["position"],
        "end_source_position": end_position,
        "confidence": score,
        "evidence": evidence,
        "entry_count": len(entry_texts),
        "entry_preview": preview,
        "repeated_heading_count": len(repeated_entries),
        "selected": False,
    }


def _report_source_toc(context: dict, report: dict, *, method: str) -> None:
    meta = context.get("source_context", {})
    candidates = context.get("document", {}).get("blocks", [])
    selected = [candidate for candidate in candidates if candidate.get("selected")]
    report["source_toc"] = {
        "status": meta.get("toc_status", "no_candidate"),
        "method": method,
        "candidates": [
            {key: value for key, value in candidate.items() if key != "block_type"}
            for candidate in candidates
        ],
        "selected_candidate_id": selected[0].get("candidate_id") if selected else None,
        "excluded_source_positions": list(meta.get("excluded_source_positions", [])),
    }


def finalize_toc_selection(context: dict, report: dict, *, method: str) -> set[int]:
    candidates = context.get("document", {}).get("blocks", [])
    selected = [candidate for candidate in candidates if candidate.get("selected")]
    if len(selected) > 1:
        selected.sort(key=lambda candidate: candidate.get("confidence", 0), reverse=True)
        for candidate in selected[1:]:
            candidate["selected"] = False
        selected = selected[:1]

    excluded: set[int] = set()
    if selected:
        candidate = selected[0]
        start = int(candidate["start_source_position"])
        end = int(candidate["end_source_position"])
        excluded.update(range(start, end + 1))
        context["source_context"]["toc_status"] = "detected"
    context["source_context"]["excluded_source_positions"] = sorted(excluded)
    _report_source_toc(context, report, method=method)
    return excluded


def detect_toc_regions(doc, report: dict) -> dict:
    """Return a source context model containing TOC region candidates."""
    records = _record_source_blocks(doc)
    starts: list[int] = []
    for index, record in enumerate(records):
        if record.get("is_toc_title") or record.get("has_toc_field"):
            starts.append(index)
        elif record.get("is_toc_style") and not starts:
            starts.append(index)

    candidates: list[dict[str, Any]] = []
    for number, start_index in enumerate(starts, start=1):
        candidate = _candidate_from_start(records, start_index, number)
        if candidate is not None:
            candidates.append(candidate)

    candidates.sort(key=lambda candidate: candidate.get("confidence", 0), reverse=True)
    status = "no_candidate"
    if candidates:
        best = candidates[0]
        runner_up = candidates[1] if len(candidates) > 1 else None
        margin = best["confidence"] - runner_up["confidence"] if runner_up else 1.0
        if best["confidence"] >= 0.75 and margin >= 0.15:
            best["selected"] = True
            status = "detected"
        else:
            status = "ambiguous"

    context = {
        "schema_version": "1.0",
        "document": {"blocks": candidates},
        "source_context": {
            "kind": "docx_toc_regions",
            "toc_status": status,
            "excluded_source_positions": [],
            "source_block_count": len(records),
        },
    }
    finalize_toc_selection(context, report, method="rules")
    return context


def selected_source_positions(context: dict) -> set[int]:
    return set(context.get("source_context", {}).get("excluded_source_positions", []))


def audit_toc_replacement(doc, context: dict | None) -> dict:
    """Audit the canonical output TOC and residual source-directory entries."""
    toc_field_count = 0
    for paragraph in doc.paragraphs:
        if re.search(r"\bTOC\b", _field_instruction(paragraph), re.I):
            toc_field_count += 1

    selected = []
    if context:
        selected = [
            candidate
            for candidate in context.get("document", {}).get("blocks", [])
            if candidate.get("selected")
        ]

    duplicate_entries: list[dict[str, Any]] = []
    if selected:
        for entry in selected[0].get("entry_preview", []):
            normalized = _normalize_entry_text(entry)
            matches = [
                index for index, paragraph in enumerate(doc.paragraphs)
                if _normalize_entry_text(paragraph.text) == normalized
            ]
            if len(matches) > 1:
                duplicate_entries.append({
                    "text": entry,
                    "paragraphs": matches,
                })

    source_toc_title_residue = [
        index for index, paragraph in enumerate(doc.paragraphs[2:], start=2)
        if paragraph.text.strip() in _TOC_TITLES
    ]
    return {
        "toc_field_count": toc_field_count,
        "has_single_canonical_toc": toc_field_count == 1,
        "duplicate_source_entries": duplicate_entries,
        "source_toc_title_residue": source_toc_title_residue,
        "passed": (
            toc_field_count == 1
            and not duplicate_entries
            and not source_toc_title_residue
        ),
    }
