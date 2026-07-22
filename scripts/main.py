from __future__ import annotations

import argparse
from copy import deepcopy
import json
import sys
from pathlib import Path

DOCX_IMPORT_ERROR = None
try:
    from docx import Document
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm
except Exception as exc:
    DOCX_IMPORT_ERROR = exc

from text_utils import (
    scan_non_text_objects,
    looks_like_code_sample_table,
    set_table_autofit_to_window,
)
from document_model import summarize_document_model, validate_document_model, compare_document_models
from docx_pipeline import parse_docx_to_model
from md_pipeline import parse_md_to_model
from model_normalization import normalize_document_model, summarize_source_document_model
from word_model_renderer import (
    render_document_model,
    style_from_profile,
)
from docx_render import render_docx_direct
from audit import audit_document, collect_content_warnings
from reporting import new_report, add_risk_warnings, write_markdown_report
from template_finalizer import apply_template_finalizer
from template_profile import load_template_profile
from note_semantics import audit_note_preservation
from appendix_semantics import audit_appendix_preservation
from toc_detector import (
    audit_toc_replacement,
    detect_toc_regions,
    finalize_toc_selection,
    selected_source_positions,
)
from list_detector import analyze_docx_lists, audit_list_preservation
from front_matter import (
    analyze_front_matter,
    audit_output_structure,
    front_matter_source_positions,
    inject_document_title,
)
from table_semantics import audit_model_table_semantics
from caption_placement import audit_model_caption_placement, audit_rendered_caption_placement
from review_loop import (
    accepted_candidate,
    build_review_packet,
    collect_audit_findings,
    unknown_pattern_packet,
)

from llm_enhancer import (
    enhance_document_model,
    compute_suspicion_score,
    normalize_mode,
    should_enhance,
)

from llm_file_protocol import (
    collect_phase_requests,
    replay_phase_responses,
    generate_run_id,
    compute_source_sha256,
    build_run_info,
    write_requests_and_run,
    read_run_info,
    read_requests,
    read_responses,
    ProtocolError,
    verify_source_snapshot,
)


SKILL_VERSION = "0.7.0"
DEFAULT_TABLE_ROW_HEIGHT_CM = 0.69
DEFAULT_TABLE_ROW_HEIGHT_RULE = "at-least"


def skill_version() -> str:
    return SKILL_VERSION


def _load_version() -> str:
    try:
        v = Path(__file__).resolve().parent.parent / "VERSION"
        return v.read_text().strip()
    except Exception:
        return SKILL_VERSION


def maybe_reexec_with_skill_venv() -> None:
    if DOCX_IMPORT_ERROR is not None:
        return
    import os, sys
    venv = os.environ.get("WX_DOC_FORMAT_VENV", "")
    if not venv:
        return
    # Already running inside the target venv — nothing to do.
    if sys.prefix == venv or os.path.realpath(sys.prefix) == os.path.realpath(venv):
        return
    venv_python = Path(venv) / "bin" / "python3"
    if venv_python.exists():
        me = Path(__file__).resolve()
        os.execv(str(venv_python), [str(venv_python), str(me)] + sys.argv[1:])


def clear_document_body(doc) -> None:
    from copy import deepcopy
    from docx.oxml.ns import qn
    body = doc.element.body
    sect_pr = None
    if len(body) and body[-1].tag == qn("w:sectPr"):
        sect_pr = deepcopy(body[-1])
    for child in list(body):
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def merge_template_numbering_ids(template_profile: dict | None, fallback_ids: dict) -> dict:
    if template_profile is None:
        return fallback_ids
    return template_profile.get("numbering_ids", {})


def ensure_fallback_style_setup(doc) -> dict:
    from fallback_styles import ensure_fallback_styles, ensure_auto_numbering
    ensure_fallback_styles(doc)
    return ensure_auto_numbering(doc)


def build_document_model_from_output_wrapper(doc, source_path: Path, report: dict) -> dict:
    from text_utils import build_document_model_from_output
    return build_document_model_from_output(doc, source_path, report)


def audit_document_wrapper(
    doc,
    row_height_cm: float,
    row_height_rule: str,
    *,
    template_profile: dict | None = None,
    table_roles: list[str] | None = None,
) -> dict:
    from text_utils import heading_level_from_style as _hls, strip_heading_marker as _shm, paragraph_num_info as _pni
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    return audit_document(
        doc, row_height_cm, row_height_rule,
        heading_level_from_style=_hls,
        paragraph_direct_num_info=_pni,
        existing_heading_number=lambda t: bool(_shm(t) != t),
        looks_like_code_sample_table=looks_like_code_sample_table,
        qn=qn,
        center_alignment=WD_ALIGN_PARAGRAPH.CENTER,
        template_profile=template_profile,
        table_roles=table_roles,
    )


def _fresh_output_document(args, template_profile):
    if args.template is not None:
        doc = _open_template_renamed_media(args.template)
        clear_document_body(doc)
    else:
        doc = Document()
        ensure_fallback_style_setup(doc)
    return doc


def _cleanup_template_copy(doc) -> None:
    path = getattr(doc, "_tmpl_path", None)
    if not path:
        return
    try:
        import os
        os.unlink(path)
    except Exception:
        pass


def _finalize_and_audit_output(
    out_doc,
    report: dict,
    *,
    args,
    template_profile,
    normalized_model: dict | None,
    source_model: dict | None,
    models: dict,
) -> None:
    report["template_finalizer"] = apply_template_finalizer(
        out_doc, template_profile,
        args.table_row_height_cm, args.table_row_height_rule,
        row_height_rule_enum=WD_ROW_HEIGHT_RULE,
        cm=Cm,
        left_alignment=WD_ALIGN_PARAGRAPH.LEFT,
        center_alignment=WD_ALIGN_PARAGRAPH.CENTER,
        set_table_autofit_to_window=set_table_autofit_to_window,
        looks_like_code_sample_table=looks_like_code_sample_table,
        table_roles=[
            block.get("table_type", "data")
            for block in (normalized_model or {}).get("document", {}).get("blocks", [])
            if block.get("block_type") == "table"
        ],
    )
    suffix = args.input.suffix.lower()
    report["toc_replacement_audit"] = audit_toc_replacement(
        out_doc,
        models.get("toc_context") if suffix == ".docx" else None,
    )
    report["output_structure_audit"] = audit_output_structure(out_doc, template_profile)
    report["caption_placement_audit"] = audit_rendered_caption_placement(out_doc)

    rendered_model = build_document_model_from_output_wrapper(out_doc, args.input, report)
    report["rendered_document_model_summary"] = report.get("document_model_summary", {})
    report["rendered_document_model_issues"] = report.get("document_model_issues", [])
    if normalized_model is not None:
        report["document_model_summary"] = summarize_document_model(normalized_model)
        report["document_model_issues"] = validate_document_model(normalized_model)
    report["document_model_diff"] = compare_document_models(
        normalized_model or source_model, rendered_model,
    )
    final_table_roles = [
        block.get("table_type", "data")
        for block in (normalized_model or {}).get("document", {}).get("blocks", [])
        if block.get("block_type") == "table"
    ]
    report["audit"] = audit_document_wrapper(
        out_doc,
        args.table_row_height_cm,
        args.table_row_height_rule,
        template_profile=template_profile,
        table_roles=final_table_roles,
    )
    report["list_preservation_audit"] = audit_list_preservation(
        out_doc,
        normalized_model,
        report.get("source_lists"),
        template_profile,
    )
    report["note_preservation_audit"] = audit_note_preservation(
        out_doc, normalized_model, template_profile,
    )
    report["appendix_preservation_audit"] = audit_appendix_preservation(
        out_doc, normalized_model, template_profile,
    )
    report["content_warnings"] = collect_content_warnings(out_doc)
    report["risk_warnings"] = []
    add_risk_warnings(report, args.table_row_height_rule)


def _render_review_candidate(
    args,
    model: dict,
    report: dict,
    *,
    template_profile,
    numbering_ids: dict,
    models: dict,
):
    out_doc = _fresh_output_document(args, template_profile)
    suffix = args.input.suffix.lower()
    if suffix in {".md", ".markdown"}:
        render_document_model(
            model, out_doc, report,
            args.table_row_height_cm, args.table_row_height_rule, numbering_ids,
            template_profile=report.get("template_profile"),
        )
    else:
        src_doc = Document(args.input)
        render_docx_direct(
            src_doc, out_doc, report,
            args.table_row_height_cm, args.table_row_height_rule, numbering_ids,
            template_profile=report.get("template_profile"),
            strict_normalize=args.strict_normalize,
            role_overrides=_extract_role_overrides_from_model(model, report) or None,
            heading_level_overrides=_extract_heading_level_overrides_from_model(model, report),
            table_type_overrides=_extract_table_type_overrides_from_model(model, report),
            model=model,
            excluded_source_positions=models.get("excluded_source_positions", set()),
        )
    return out_doc


def _llm_call_from_command(command: str):
    """Build an LLM callable from a shell command template.

    The command receives the prompt via stdin by default.  If the command
    contains the ``{prompt_file}`` placeholder, the prompt is written to a
    temporary file and the placeholder is replaced with its quoted path.

    Returns a ``(prompt: str) -> str`` callable, or ``None`` if *command*
    is empty.
    """
    import os as _os
    import shlex
    import subprocess
    import tempfile

    command = command.strip()
    if not command:
        return None

    def _llm(prompt: str) -> str:
        if "{prompt_file}" not in command:
            r = subprocess.run(
                command,
                input=prompt,
                text=True,
                shell=True,
                capture_output=True,
                timeout=60,
            )
            if r.returncode != 0:
                raise RuntimeError((r.stderr or r.stdout or "").strip())
            return r.stdout or ""

        prompt_path = None
        try:
            with tempfile.NamedTemporaryFile("w", encoding="utf-8", suffix=".prompt.txt", delete=False) as f:
                f.write(prompt)
                prompt_path = f.name
            r = subprocess.run(
                command.replace("{prompt_file}", shlex.quote(prompt_path)),
                text=True,
                shell=True,
                capture_output=True,
                timeout=60,
            )
            if r.returncode != 0:
                raise RuntimeError((r.stderr or r.stdout or "").strip())
            return r.stdout or ""
        finally:
            if prompt_path:
                try:
                    _os.unlink(prompt_path)
                except OSError:
                    pass

    return _llm


def _resolve_llm_call(args):
    """Auto-detect LLM backend for semantic enhancement.

    Returns a callable that accepts a prompt string and returns the LLM
    response, or None if no LLM backend is available.
    """
    import os as _os

    # 1. Explicit LLM command (--llm-command / LLM_COMMAND)
    command = getattr(args, "llm_command", None) if args is not None else None
    command = command or _os.environ.get("LLM_COMMAND")
    command_call = _llm_call_from_command(command or "")
    if command_call is not None:
        return command_call

    # 2. Anthropic API
    key = _os.environ.get("ANTHROPIC_API_KEY")
    if key:
        import anthropic
        client = anthropic.Anthropic(api_key=key)
        model = _os.environ.get("LLM_MODEL", "claude-sonnet-4-20250514")

        def _llm(prompt: str) -> str:
            msg = client.messages.create(
                model=model,
                max_tokens=1024,
                messages=[{"role": "user", "content": prompt}],
            )
            return msg.content[0].text

        return _llm

    # 3. OpenAI API
    key = _os.environ.get("OPENAI_API_KEY")
    if key:
        import openai
        client = openai.OpenAI(api_key=key, base_url=_os.environ.get("OPENAI_BASE_URL"))
        model = _os.environ.get("LLM_MODEL", "gpt-4o")

        def _llm(prompt: str) -> str:
            resp = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0,
            )
            return resp.choices[0].message.content or ""

        return _llm

    return None


def convert_md(src: Path, doc, report: dict, row_height_cm: float, row_height_rule: str, numbering_ids: dict, args=None, file_protocol_ctx=None) -> dict:
    source_model = parse_md_to_model(src, report, skill_version=skill_version)
    summarize_source_document_model(report, source_model)

    # ── LLM Enhancement ──
    enhance_mode = getattr(args, "llm_enhance", "off") if args is not None else "off"
    enhance_mode = normalize_mode(enhance_mode)
    llm_call = _resolve_llm_call(args) if enhance_mode != "off" else None
    llm_hint = getattr(args, "llm_hint", None) if args is not None else None

    # Phase A: block role review (before normalization)
    if should_enhance(report, "A", enhance_mode):
        source_model = _run_phase_enhancement(
            source_model, report, phase="A",
            llm_call=llm_call, hint=llm_hint,
            file_protocol_ctx=file_protocol_ctx,
        )

    model = normalize_document_model(source_model, report)

    # Phase B: caption text generation via LLM (after normalization)
    if should_enhance(report, "B", enhance_mode):
        model = _run_phase_enhancement(
            model, report, phase="B",
            llm_call=llm_call, hint=llm_hint,
            file_protocol_ctx=file_protocol_ctx,
        )

    report["table_semantics_audit"] = audit_model_table_semantics(model)
    report["caption_placement_model_audit"] = audit_model_caption_placement(model)

    render_document_model(
        model, doc, report, row_height_cm, row_height_rule, numbering_ids,
        template_profile=report.get("template_profile"),
    )
    return {"source": source_model, "normalized": model}


def convert_docx(src: Path, dst_doc, row_height_cm: float, row_height_rule: str, strict_normalize: bool, report: dict, numbering_ids: dict, args=None, file_protocol_ctx=None) -> dict:
    from docx.oxml.ns import qn as _qn
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    from docx_pipeline import infer_docx_role

    src_doc = Document(src)
    toc_context = detect_toc_regions(src_doc, report)
    enhance_mode = normalize_mode(getattr(args, "llm_enhance", "off")) if args is not None else "off"
    llm_call = _resolve_llm_call(args) if enhance_mode != "off" else None
    llm_hint = getattr(args, "llm_hint", None) if args is not None else None

    if should_enhance(report, "S", enhance_mode):
        toc_context = _run_phase_enhancement(
            toc_context, report, phase="S",
            llm_call=llm_call, hint=llm_hint,
            file_protocol_ctx=file_protocol_ctx,
        )
    finalize_toc_selection(toc_context, report, method=(
        "llm" if any(
            item.get("operation") == "exclude_toc_region"
            for item in report.get("llm_enhancer", {}).get("applied", [])
        ) else "rules"
    ))
    front_matter_context = analyze_front_matter(src_doc, toc_context, src, report)
    excluded_source_positions = (
        selected_source_positions(toc_context)
        | front_matter_source_positions(front_matter_context)
    )
    numbering_context = analyze_docx_lists(
        src_doc,
        report,
        excluded_source_positions=excluded_source_positions,
    )
    source_model = parse_docx_to_model(
        src, src_doc, strict_normalize, 0,
        skill_version=skill_version,
        new_report=lambda: {},
        iter_blocks=lambda d: (
            Paragraph(child, d) if child.tag == _qn("w:p")
            else Table(child, d)
            for child in d.element.body.iterchildren()
            if child.tag in (_qn("w:p"), _qn("w:tbl"))
        ),
        paragraph_class=Paragraph,
        infer_docx_role=infer_docx_role,
        looks_like_code_sample_table=looks_like_code_sample_table,
        caption_pattern=None,
        excluded_source_positions=excluded_source_positions,
        numbering_context=numbering_context,
    )
    report["parse_report"] = source_model.get("parse_report", {})
    source_model = inject_document_title(source_model, front_matter_context)
    summarize_source_document_model(report, source_model)

    # ── LLM Enhancement ──
    # Phase A: block role review (before normalization)
    applied_start = len(report.get("llm_enhancer", {}).get("applied", []))
    if should_enhance(report, "A", enhance_mode):
        source_model = _run_phase_enhancement(
            source_model, report, phase="A",
            llm_call=llm_call, hint=llm_hint,
            file_protocol_ctx=file_protocol_ctx,
        )

    model = normalize_document_model(source_model, report)

    # Phase B: caption text generation via LLM (after normalization)
    if should_enhance(report, "B", enhance_mode):
        model = _run_phase_enhancement(
            model, report, phase="B",
            llm_call=llm_call, hint=llm_hint,
            file_protocol_ctx=file_protocol_ctx,
        )

    report["table_semantics_audit"] = audit_model_table_semantics(model)
    report["caption_placement_model_audit"] = audit_model_caption_placement(model)

    # role_overrides for direct rendering path — extracted from Phase A applied decisions
    phase_a_role_overrides: dict[int, str] = {}
    if should_enhance(report, "A", enhance_mode):
        phase_a_role_overrides = _extract_role_overrides_from_model(
            model, report, applied_start=applied_start,
        )
    role_overrides = phase_a_role_overrides or None

    # Build heading_level_overrides and table_type_overrides from the
    # enhanced model.  The index position tracks model block order which
    # matches source-document iteration order for typical documents
    # without empty-paragraph gaps.
    heading_level_overrides = _extract_heading_level_overrides_from_model(model, report)
    table_type_overrides = _extract_table_type_overrides_from_model(model, report)

    render_docx_direct(
        src_doc, dst_doc, report, row_height_cm, row_height_rule, numbering_ids,
        template_profile=report.get("template_profile"),
        strict_normalize=strict_normalize,
        role_overrides=role_overrides,
        heading_level_overrides=heading_level_overrides,
        table_type_overrides=table_type_overrides,
        model=model,
        excluded_source_positions=excluded_source_positions,
    )
    return {
        "source": source_model,
        "normalized": model,
        "toc_context": toc_context,
        "front_matter_context": front_matter_context,
        "excluded_source_positions": excluded_source_positions,
    }


def _extract_heading_level_overrides_from_model(model: dict, report: dict) -> dict[int, int]:
    """Build heading_level_overrides from the enhanced model.

    Maps model-block position → heading level for heading blocks that
    Phase B adjusted (adjust_level / retype-to-heading).  The model-block
    position is a best-effort match for *rerender_docx_direct*'s
    ``_para_idx``; it works when model blocks align 1:1 with source
    paragraphs/tables (the common case for well-formed documents).

    Returns an empty dict when there are no Phase B heading changes.
    """
    # Collect block IDs changed by Phase B heading operations.
    changed_heading_ids: set[str] = set()
    for dec in report.get("llm_enhancer", {}).get("applied", []):
        op = dec.get("operation", "")
        bid = dec.get("block_id", "")
        if op == "adjust_level":
            changed_heading_ids.add(bid)
        elif op == "retype":
            to_type = dec.get("to", {}).get("block_type", "")
            if to_type == "heading":
                changed_heading_ids.add(bid)

    if not changed_heading_ids:
        return {}

    overrides: dict[int, int] = {}
    idx = 0
    for block in model.get("document", {}).get("blocks", []):
        btype = block.get("block_type")
        if btype in ("heading", "body", "list_item", "caption",
                      "table", "image", "appendix", "unknown"):
            if btype == "heading" and block.get("id") in changed_heading_ids:
                overrides[idx] = block.get("level", 1)
            idx += 1

    return overrides


def _extract_table_type_overrides_from_model(model: dict, report: dict) -> dict[int, str]:
    """Build table_type_overrides from the enhanced model.

    Maps model-block position → table type string for tables that
    were changed via set_table_type.  See heading_level_overrides
    for index-mapping caveats.

    Returns an empty dict when there are no table-type changes.
    """
    changed_table_ids: set[str] = set()
    for dec in report.get("llm_enhancer", {}).get("applied", []):
        op = dec.get("operation", "")
        if op == "set_table_type":
            bid = dec.get("block_id", "")
            if bid:
                changed_table_ids.add(bid)

    if not changed_table_ids:
        return {}

    overrides: dict[int, str] = {}
    idx = 0
    for block in model.get("document", {}).get("blocks", []):
        btype = block.get("block_type")
        if btype in ("heading", "body", "list_item", "caption",
                      "table", "image", "appendix", "unknown"):
            if btype == "table" and block.get("id") in changed_table_ids:
                overrides[idx] = block.get("table_type", "data")
            idx += 1

    return overrides


def _extract_role_overrides_from_model(
    model: dict,
    report: dict,
    *,
    applied_start: int = 0,
) -> dict[int, str]:
    """从 Phase A 的 applied decisions 提取段落级别的 role_overrides。

    只提取 retype 操作中与原始类型不同的映射，返回 ``{模型块索引: 新角色}``
    用于 ``render_docx_direct`` 的段落实例化。
    """
    changed_roles: dict[str, str] = {}
    applied = report.get("llm_enhancer", {}).get("applied", [])[applied_start:]

    for dec in applied:
        if dec.get("operation") != "retype":
            continue
        to_type = dec.get("to", {}).get("block_type", "")
        from_type = dec.get("from", {}).get("block_type", "")
        if to_type in {"heading", "body", "list_item", "caption"} and to_type != from_type:
            bid = dec.get("block_id", "")
            if bid:
                changed_roles[bid] = to_type

    if not changed_roles:
        return {}

    overrides: dict[int, str] = {}
    idx = 0
    for block in model.get("document", {}).get("blocks", []):
        btype = block.get("block_type")
        if btype in {"heading", "body", "list_item", "caption",
                      "table", "image", "appendix", "unknown"}:
            bid = block.get("id")
            if bid in changed_roles:
                overrides[idx] = changed_roles[bid]
            idx += 1

    return overrides


def _open_template_renamed_media(template_path):
    """Open template, renaming its media files to _tmpl suffix to avoid conflicts."""
    import zipfile, io, tempfile, os, re
    buf = io.BytesIO()
    with zipfile.ZipFile(template_path, 'r') as z:
        with zipfile.ZipFile(buf, 'w') as out:
            for item in z.infolist():
                data = z.read(item.filename)
                if item.filename.startswith('word/media/') and not item.filename.endswith('/'):
                    dot_idx = item.filename.rfind('.')
                    if dot_idx > 0:
                        new_name = item.filename[:dot_idx] + '_tmpl' + item.filename[dot_idx:]
                    else:
                        new_name = item.filename + '_tmpl'
                    out.writestr(new_name, data)
                elif item.filename == 'word/_rels/document.xml.rels':
                    rels_text = data.decode('utf-8')
                    rels_text = re.sub(
                        r'Target="media/([^"]+?)\.(png|jpg|jpeg|gif|bmp|emf|wmf)"',
                        r'Target="media/\1_tmpl.\2"',
                        rels_text
                    )
                    out.writestr(item, rels_text.encode('utf-8'))
                else:
                    out.writestr(item, data)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    tmp.write(buf.getvalue())
    tmp.close()
    from docx import Document
    doc = Document(tmp.name)
    doc._tmpl_path = tmp.name
    return doc


def _run_phase_enhancement(model, report, *, phase, llm_call, hint, file_protocol_ctx=None):
    """Run enhancement for a phase — normal, generate, or resume.

    * Normal mode: calls ``enhance_document_model`` with the real LLM callable.
    * Generate mode: collects prompts into ``file_protocol_ctx["requests"]``,
      returns model unchanged.
    * Resume mode: replays saved responses through validate + apply.
    """
    if file_protocol_ctx is None:
        return enhance_document_model(
            model, report, phase=phase, llm_call=llm_call, hint=hint,
        )

    mode = file_protocol_ctx.get("mode")

    if mode == "generate":
        reqs = collect_phase_requests(model, report, phase=phase, hint=hint)
        file_protocol_ctx.setdefault("requests", []).extend(reqs)
        return model  # no LLM call in generate mode

    if mode == "resume":
        from llm_enhancer import _resolve_legacy_phase

        request_phase = _resolve_legacy_phase(phase)
        phase_requests = [
            request for request in file_protocol_ctx.get("requests", [])
            if request.get("phase") == request_phase
        ]
        if not phase_requests:
            reqs = collect_phase_requests(model, report, phase=phase, hint=hint)
            if reqs:
                file_protocol_ctx.setdefault("requests", []).extend(reqs)
                file_protocol_ctx.setdefault("new_requests", []).extend(reqs)
            return model
        return replay_phase_responses(
            model, report, phase=phase,
            requests=file_protocol_ctx.get("requests", []),
            responses=file_protocol_ctx.get("responses", []),
            hint=hint,
        )

    return model


def _write_protocol_requests(args, requests: list[dict], *, request_stage: str) -> dict:
    run_id = generate_run_id()
    source_sha256 = compute_source_sha256(args.input)
    run_info = build_run_info(
        run_id=run_id,
        source_path=str(args.input),
        source_sha256=source_sha256,
        args={
            "input": str(args.input),
            "output": str(args.output),
            "template": str(args.template) if args.template else None,
            "report": str(args.report) if args.report else None,
            "report_md": str(args.report_md) if args.report_md else None,
            "ast": str(args.ast) if args.ast else None,
            "source_ast": str(args.source_ast) if args.source_ast else None,
            "fail_on_risk": bool(args.fail_on_risk),
            "llm_enhance": args.llm_enhance,
            "llm_hint": args.llm_hint,
            "strict_normalize": args.strict_normalize,
            "table_row_height_cm": args.table_row_height_cm,
            "table_row_height_rule": args.table_row_height_rule,
        },
        work_dir=str(args.generate_requests),
    )
    run_info["request_stage"] = request_stage
    write_requests_and_run(requests, run_info, args.generate_requests)
    return run_info


def _run_generate_requests(args, report) -> None:
    """Phase 1: parse source, build prompts, write files, stop."""
    suffix = args.input.suffix.lower()
    llm_hint = getattr(args, "llm_hint", None)
    all_requests: list[dict] = []
    fp_ctx: dict = {"mode": "generate", "requests": all_requests}
    enhance_mode = normalize_mode(getattr(args, "llm_enhance", "off"))
    toc_context = None
    excluded_source_positions: set[int] = set()

    # ── Parse source ─────────────────────────────────────────────────
    if suffix in {".md", ".markdown"}:
        source_model = parse_md_to_model(
            args.input, report, skill_version=skill_version,
        )
    elif suffix == ".docx":
        from docx import Document as _Document
        from docx.oxml.ns import qn as _qn
        from docx.text.paragraph import Paragraph
        from docx.table import Table
        from docx_pipeline import infer_docx_role

        src_doc = _Document(args.input)
        toc_context = detect_toc_regions(src_doc, report)
        if should_enhance(report, "S", enhance_mode):
            _run_phase_enhancement(
                toc_context, report, phase="S",
                llm_call=None, hint=llm_hint,
                file_protocol_ctx=fp_ctx,
            )
        if all_requests:
            _write_protocol_requests(args, all_requests, request_stage="source")
            print(f"Generated {len(all_requests)} source-stage LLM request(s) in {args.generate_requests.resolve()}")
            print("Run with --resume <run.json> after processing llm_responses.jsonl.")
            return
        front_matter_context = analyze_front_matter(
            src_doc, toc_context, args.input, report,
        )
        excluded_source_positions = (
            selected_source_positions(toc_context)
            | front_matter_source_positions(front_matter_context)
        )
        numbering_context = analyze_docx_lists(
            src_doc,
            report,
            excluded_source_positions=excluded_source_positions,
        )
        source_model = parse_docx_to_model(
            args.input, src_doc, True, 0,
            skill_version=skill_version,
            new_report=lambda: {},
            iter_blocks=lambda d: (
                Paragraph(child, d) if child.tag == _qn("w:p")
                else Table(child, d)
                for child in d.element.body.iterchildren()
                if child.tag in (_qn("w:p"), _qn("w:tbl"))
            ),
            paragraph_class=Paragraph,
            infer_docx_role=infer_docx_role,
            looks_like_code_sample_table=looks_like_code_sample_table,
            caption_pattern=None,
            excluded_source_positions=excluded_source_positions,
            numbering_context=numbering_context,
        )
        source_model = inject_document_title(source_model, front_matter_context)
    else:
        raise SystemExit(f"Unsupported input type: {args.input.suffix}")

    summarize_source_document_model(report, source_model)

    # ── Collect Phase A requests ─────────────────────────────────────
    if should_enhance(report, "A", enhance_mode):
        _run_phase_enhancement(
            source_model, report, phase="A",
            llm_call=None, hint=llm_hint,
            file_protocol_ctx=fp_ctx,
        )

    # ── Normalize ────────────────────────────────────────────────────
    model = normalize_document_model(source_model, report)

    # ── Collect Phase B requests ─────────────────────────────────────
    if should_enhance(report, "B", enhance_mode):
        _run_phase_enhancement(
            model, report, phase="B",
            llm_call=None, hint=llm_hint,
            file_protocol_ctx=fp_ctx,
        )

    # ── Write requests and run.json ──────────────────────────────────
    _write_protocol_requests(args, fp_ctx["requests"], request_stage="ast")

    print(f"Generated {len(fp_ctx['requests'])} LLM request(s) in {args.generate_requests.resolve()}")
    print("Run with --resume <run.json> after processing llm_responses.jsonl.")


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert MD or DOCX to template-formatted DOCX.")
    parser.add_argument("--input", type=Path, help="Source file (optional when --resume is used)")
    parser.add_argument("--output", type=Path, help="Output .docx path (optional when --resume is used)")
    parser.add_argument("--template", type=Path, default=None, help="DOCX template driving styles and numbering.")
    parser.add_argument("--strict-normalize", action=argparse.BooleanOptionalAction, default=True)
    parser.add_argument("--report", type=Path, default=None)
    parser.add_argument("--report-md", type=Path, default=None)
    parser.add_argument("--ast", type=Path, default=None, help="Normalized document model (Step 2 output).")
    parser.add_argument("--source-ast", type=Path, default=None, help="Parsed source document model (Step 1 output).")
    parser.add_argument("--table-row-height-cm", type=float, default=DEFAULT_TABLE_ROW_HEIGHT_CM)
    parser.add_argument("--table-row-height-rule", choices=["exact", "at-least"], default=DEFAULT_TABLE_ROW_HEIGHT_RULE)
    parser.add_argument("--fail-on-risk", action="store_true")
    parser.add_argument(
        "--llm-enhance",
        choices=[
            "off", "auto",
            "s", "a", "b", "c", "sa", "sb", "sc", "ab", "ac", "bc",
            "sab", "sac", "sbc", "abc", "sabc",
            "force-s", "force-a", "force-b", "force-c", "force-sa", "force-sb", "force-sc",
            "force-ab", "force-ac", "force-bc", "force-sab", "force-sac", "force-sbc",
            "force-abc", "force-sabc",
            "toc_region_review", "list_detect", "caption_gen", "document_review", "all",
        ],
        default="off",
        help=(
            "LLM enhancement level (default: off). "
            "New names: toc_region_review / list_detect / caption_gen / document_review / all. "
            "Compact phase combinations remain supported."
        ),
    )
    parser.add_argument(
        "--llm-hint",
        type=str,
        default=None,
        help="Natural-language hint injected into LLM enhancement prompts"
    )
    parser.add_argument(
        "--llm-command",
        type=str,
        default=None,
        help=(
            "Shell command used for LLM enhancement. The prompt is sent to stdin "
            "and stdout is used as the response. Use {prompt_file} if the command "
            "expects a prompt file path."
        ),
    )
    parser.add_argument(
        "--generate-requests",
        type=Path,
        default=None,
        help=(
            "Generate LLM request files (llm_requests.jsonl + run.json) in DIR "
            "and stop.  The agent processes the requests and writes "
            "llm_responses.jsonl, then --resume continues."
        ),
    )
    parser.add_argument(
        "--resume",
        type=Path,
        default=None,
        help=(
            "Resume from a run.json file produced by --generate-requests. "
            "Reads llm_responses.jsonl, validates integrity, and continues "
            "the pipeline (enhance → validate → apply → render)."
        ),
    )
    args = parser.parse_args()

    # ── Argument validation ─────────────────────────────────────────
    if args.generate_requests and args.resume:
        parser.error("--generate-requests and --resume are mutually exclusive")

    if args.generate_requests and not args.input:
        parser.error("--input is required when using --generate-requests")

    if not args.resume and not args.generate_requests:
        if not args.input:
            parser.error("--input is required")
        if not args.output:
            parser.error("--output is required")

    # Resume mode may pull input/output from run.json so we defer
    # the required check until after the resume handler runs.

    # Backward compatibility: WX_DOC_LLM_ENHANCE env var → abc mode
    import os as _os
    if args.llm_enhance == "off" and _os.environ.get("WX_DOC_LLM_ENHANCE") == "1":
        args.llm_enhance = "abc"

    # Normalize new capability names to legacy mode names
    args.llm_enhance = normalize_mode(args.llm_enhance)

    # ── Handle --resume: load run.json and restore original CLI args ──
    if args.resume:
        run_info = read_run_info(args.resume)
        run_args = run_info.get("args", {})
        # Only fill in input/output/template from run.json if the user
        # didn't provide them on the CLI (allows override).
        if not args.input and "input" in run_args:
            args.input = Path(run_args["input"])
        if not args.output and "output" in run_args:
            args.output = Path(run_args["output"])
        if not args.template and run_args.get("template"):
            args.template = Path(run_args["template"])
        if not args.report and run_args.get("report"):
            args.report = Path(run_args["report"])
        if not args.report_md and run_args.get("report_md"):
            args.report_md = Path(run_args["report_md"])
        if not args.ast and run_args.get("ast"):
            args.ast = Path(run_args["ast"])
        if not args.source_ast and run_args.get("source_ast"):
            args.source_ast = Path(run_args["source_ast"])
        if "fail_on_risk" in run_args:
            args.fail_on_risk = bool(run_args["fail_on_risk"])
        if "llm_enhance" in run_args:
            args.llm_enhance = run_args["llm_enhance"]
        if "llm_hint" in run_args and run_args["llm_hint"]:
            args.llm_hint = run_args["llm_hint"]
        if "strict_normalize" in run_args:
            args.strict_normalize = run_args["strict_normalize"]
        if "table_row_height_cm" in run_args:
            args.table_row_height_cm = run_args["table_row_height_cm"]
        if "table_row_height_rule" in run_args:
            args.table_row_height_rule = run_args["table_row_height_rule"]

    # After --resume resolution, verify required paths exist.
    if args.resume and not args.input:
        raise SystemExit("--resume: run.json does not contain 'input', provide --input on CLI")
    if args.resume and not args.output:
        raise SystemExit("--resume: run.json does not contain 'output', provide --output on CLI")
    if args.resume:
        try:
            verify_source_snapshot(args.input, str(run_info.get("source_sha256") or ""))
        except ProtocolError as exc:
            raise SystemExit(f"--resume: {exc}") from exc

    global SKILL_VERSION
    SKILL_VERSION = _load_version()

    maybe_reexec_with_skill_venv()

    if DOCX_IMPORT_ERROR is not None:
        if args.input.suffix.lower() == ".docx":
            sys.path.insert(0, str(Path(__file__).resolve().parent))
            from format_docx_ooxml import convert_docx_ooxml
            convert_docx_ooxml(args.input, args.output, args.report, args.report_md)
            print(args.output)
            return
        raise SystemExit("python-docx or lxml not available.")

    report = new_report(SKILL_VERSION)
    template_profile = None

    if args.template is not None:
        template_profile = load_template_profile(args.template)
        out_doc = _open_template_renamed_media(args.template)
        clear_document_body(out_doc)
        numbering_ids = merge_template_numbering_ids(template_profile, {})
        report["template_profile"] = {
            "path": str(args.template),
            "resolved_styles": template_profile.get("resolved_styles", {}),
            "missing_roles": template_profile.get("missing_roles", []),
            "numbering_ids": template_profile.get("numbering_ids", {}),
            "table_style": template_profile.get("table_style", {}),
        }
    else:
        out_doc = Document()
        numbering_ids = ensure_fallback_style_setup(out_doc)

    report["non_text_objects"] = scan_non_text_objects(args.input)
    source_document_model = None
    normalized_document_model = None
    suffix = args.input.suffix.lower()

    # ── Phase 1: Generate LLM requests and stop ─────────────────────
    if args.generate_requests:
        _run_generate_requests(args, report)
        return

    # ── Phase 3: Resume — set up file protocol context ──────────────
    file_protocol_ctx = None
    if args.resume:
        run_info = read_run_info(args.resume)
        reqs = read_requests(Path(run_info["requests_path"]))
        resps = read_responses(Path(run_info["responses_path"]))
        file_protocol_ctx = {
            "mode": "resume",
            "requests": reqs,
            "responses": resps,
        }

    if suffix in {".md", ".markdown"}:
        models = convert_md(args.input, out_doc, report, args.table_row_height_cm, args.table_row_height_rule, numbering_ids, args, file_protocol_ctx=file_protocol_ctx)
        source_document_model = models["source"]
        normalized_document_model = models["normalized"]
    elif suffix == ".docx":
        models = convert_docx(args.input, out_doc, args.table_row_height_cm, args.table_row_height_rule, args.strict_normalize, report, numbering_ids, args, file_protocol_ctx=file_protocol_ctx)
        source_document_model = models["source"]
        normalized_document_model = models["normalized"]
    else:
        raise SystemExit(f"Unsupported input type: {args.input.suffix}")

    if file_protocol_ctx and file_protocol_ctx.get("new_requests"):
        run_info["request_stage"] = "ast"
        request_dir = Path(run_info["requests_path"]).parent
        write_requests_and_run(
            file_protocol_ctx.get("requests", []),
            run_info,
            request_dir,
        )
        print(
            f"Generated {len(file_protocol_ctx['new_requests'])} downstream "
            f"LLM request(s) in {request_dir.resolve()}"
        )
        print("Complete the new responses, then run --resume again.")
        return

    _finalize_and_audit_output(
        out_doc, report,
        args=args,
        template_profile=template_profile,
        normalized_model=normalized_document_model,
        source_model=source_document_model,
        models=models,
    )

    # Post-audit semantic review. Every mutation is applied to a cloned AST,
    # rendered into a fresh document, and accepted only after full re-audit.
    enhance_mode = normalize_mode(getattr(args, "llm_enhance", "off"))
    report["review_packet"] = build_review_packet(report, normalized_document_model or {})
    review_state = {
        "enabled": should_enhance(report, "C", enhance_mode),
        "triggered": False,
        "status": "not_requested",
        "rounds": 0,
        "repairable_findings": report["review_packet"].get("repairable_count", 0),
    }
    report["review_loop"] = review_state

    if review_state["enabled"] and not review_state["repairable_findings"]:
        review_state["status"] = "no_repairable_findings"

    if (
        review_state["enabled"]
        and review_state["repairable_findings"]
        and normalized_document_model is not None
    ):
        review_state["triggered"] = True
        review_state["status"] = "reviewing"
        baseline_report = deepcopy(report)
        baseline_model = deepcopy(normalized_document_model)
        candidate_report = deepcopy(report)
        candidate_model = deepcopy(normalized_document_model)
        before_applied = len(candidate_report.get("llm_enhancer", {}).get("applied", []))
        llm_call = _resolve_llm_call(args) if file_protocol_ctx is None else None
        candidate_model = _run_phase_enhancement(
            candidate_model, candidate_report, phase="C",
            llm_call=llm_call, hint=getattr(args, "llm_hint", None),
            file_protocol_ctx=file_protocol_ctx,
        )

        if file_protocol_ctx and file_protocol_ctx.get("new_requests"):
            run_info["request_stage"] = "review"
            request_dir = Path(run_info["requests_path"]).parent
            write_requests_and_run(
                file_protocol_ctx.get("requests", []), run_info, request_dir,
            )
            print(
                f"Generated {len(file_protocol_ctx['new_requests'])} post-audit "
                f"LLM review request(s) in {request_dir.resolve()}"
            )
            print("Complete the new responses, then run --resume again.")
            _cleanup_template_copy(out_doc)
            return

        applied = candidate_report.get("llm_enhancer", {}).get("applied", [])[before_applied:]
        review_state["decisions"] = applied
        if applied:
            review_state["rounds"] = 1
            candidate_doc = None
            try:
                candidate_model = normalize_document_model(candidate_model, candidate_report)
                candidate_report["table_semantics_audit"] = audit_model_table_semantics(candidate_model)
                candidate_report["caption_placement_model_audit"] = audit_model_caption_placement(candidate_model)
                candidate_doc = _render_review_candidate(
                    args, candidate_model, candidate_report,
                    template_profile=template_profile,
                    numbering_ids=numbering_ids,
                    models=models,
                )
                _finalize_and_audit_output(
                    candidate_doc, candidate_report,
                    args=args,
                    template_profile=template_profile,
                    normalized_model=candidate_model,
                    source_model=source_document_model,
                    models=models,
                )
                accepted, reason, comparison = accepted_candidate(
                    baseline_report, candidate_report, baseline_model, candidate_model,
                    decisions=applied,
                )
            except Exception as exc:
                accepted = False
                reason = "candidate_pipeline_error"
                comparison = {
                    "candidate_error": {
                        "type": type(exc).__name__,
                        "message": str(exc),
                    }
                }
            review_state.update({
                "status": "accepted" if accepted else "rolled_back",
                "accepted": accepted,
                "reason": reason,
                **comparison,
            })
            if accepted and candidate_doc is not None:
                candidate_report["review_packet"] = build_review_packet(
                    candidate_report, candidate_model,
                )
                _cleanup_template_copy(out_doc)
                out_doc = candidate_doc
                report = candidate_report
                normalized_document_model = candidate_model
                report["review_loop"] = review_state
            else:
                if candidate_doc is not None:
                    _cleanup_template_copy(candidate_doc)
                report["review_loop"] = review_state
        else:
            has_backend = llm_call is not None or file_protocol_ctx is not None
            review_state["status"] = "no_changes" if has_backend else "no_backend"

    remaining_findings = collect_audit_findings(report, normalized_document_model or {})
    if remaining_findings:
        report["unknown_pattern_packet"] = unknown_pattern_packet(
            report,
            normalized_document_model or {},
            input_hash=compute_source_sha256(args.input),
        )
    if report.get("review_loop", {}).get("status") == "rolled_back":
        report.setdefault("risk_warnings", []).append({
            "type": "document_review_rolled_back",
            "message": "Post-audit semantic repair was rejected and the baseline document was restored.",
            "reason": report["review_loop"].get("reason"),
        })

    args.output.parent.mkdir(parents=True, exist_ok=True)
    out_doc.save(args.output)
    _cleanup_template_copy(out_doc)

    if args.ast is not None and normalized_document_model is not None:
        args.ast.parent.mkdir(parents=True, exist_ok=True)
        args.ast.write_text(json.dumps(normalized_document_model, ensure_ascii=False, indent=2), encoding="utf-8")
    if args.source_ast is not None and source_document_model is not None:
        args.source_ast.parent.mkdir(parents=True, exist_ok=True)
        args.source_ast.write_text(json.dumps(source_document_model, ensure_ascii=False, indent=2), encoding="utf-8")
    if args.report is not None:
        args.report.parent.mkdir(parents=True, exist_ok=True)
        args.report.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    if args.report_md is not None:
        write_markdown_report(report, args.report_md)
    if args.fail_on_risk and report["risk_warnings"]:
        raise SystemExit("Conversion completed with risk warnings. Review the report before delivery.")
    print(args.output)


if __name__ == "__main__":
    main()
