from __future__ import annotations

import re
from collections import Counter
from typing import Any

from docx.oxml.ns import qn
from toc_contract import TOC_CUSTOM_STYLE_LEVELS


APPENDIX_TITLE_STYLE = "附录标题"
APPENDIX_HEADING_STYLES = {
    1: "附录一级标题",
    2: "附录二级标题",
    3: "附录三级标题",
}
APPENDIX_STYLE_ROLES = {
    APPENDIX_TITLE_STYLE: "appendix_title",
    **{
        style_name: f"appendix_heading_{level}"
        for level, style_name in APPENDIX_HEADING_STYLES.items()
    },
}

APPENDIX_TITLE_FONT = "黑体"
APPENDIX_TITLE_SIZE_HALF_POINTS = 28
APPENDIX_HEADING_SIZE_HALF_POINTS = 24
APPENDIX_LINE_SPACING_TWIPS = 300
APPENDIX_HALF_LINE_SPACING_TWIPS = 50

_APPENDIX_MARK_RE = re.compile(r"^附\s*录\s*([A-Z])(?:\s*|[：:].*)$", re.IGNORECASE)
_APPENDIX_CANDIDATE_RE = re.compile(
    r"^附\s*录(?:\s*([A-Z]))?(?:\s*[：:]\s*(.+)|\s+(.+))?$",
    re.IGNORECASE,
)
_APPENDIX_VISIBLE_HEADING_RE = re.compile(
    r"^\s*([A-Z])((?:[.．]\d+){1,3})(?:[\s、：:.)）-]+)?(.*)$",
    re.IGNORECASE,
)
_CLASSIFICATION_RE = re.compile(r"^[（(](规范性|资料性)[)）]$")
_CLASSIFICATION_TITLE_RE = re.compile(r"^[（(](规范性|资料性)[)）]\s*(.+)$")


def normalize_style_name(value: str | None) -> str:
    return (value or "").casefold().replace(" ", "")


def appendix_role_from_style(style_name: str | None) -> str | None:
    normalized = normalize_style_name(style_name)
    for name, role in APPENDIX_STYLE_ROLES.items():
        if normalized == normalize_style_name(name):
            return role
    return None


def appendix_heading_level(style_name: str | None) -> int | None:
    role = appendix_role_from_style(style_name)
    if not role or not role.startswith("appendix_heading_"):
        return None
    return int(role.rsplit("_", 1)[1])


def is_explicit_appendix_title(text: str) -> bool:
    lines = [line.strip() for line in (text or "").replace("\r", "").split("\n")]
    return any(_APPENDIX_MARK_RE.match(line) for line in lines if line)


def parse_appendix_marker(text: str) -> dict[str, str | None] | None:
    value = (text or "").replace("\r", "").strip()
    if "\n" in value:
        return None
    match = _APPENDIX_CANDIDATE_RE.fullmatch(value)
    if match is None:
        return None
    title = (match.group(2) or match.group(3) or "").strip()
    return {
        "appendix_id": match.group(1).upper() if match.group(1) else None,
        "title": title or None,
    }


def appendix_role_for_paragraph(style_name: str | None, text: str) -> str | None:
    role = appendix_role_from_style(style_name)
    if role:
        return role
    if is_explicit_appendix_title(text):
        return "appendix_title"
    return None


def appendix_index_to_id(index: int) -> str:
    value = max(1, int(index))
    result = ""
    while value:
        value, remainder = divmod(value - 1, 26)
        result = chr(ord("A") + remainder) + result
    return result


def parse_appendix_title(text: str, appendix_index: int) -> dict[str, Any]:
    raw_lines = (text or "").replace("\r", "").split("\n")
    lines = [line.strip() for line in raw_lines]
    appendix_id = None
    classification = None
    title_parts: list[str] = []

    for line in lines:
        if not line:
            continue
        marker = parse_appendix_marker(line)
        if marker:
            appendix_id = marker.get("appendix_id") or appendix_id
            if marker.get("title"):
                title_parts.append(str(marker["title"]))
            continue
        classification_only = _CLASSIFICATION_RE.match(line)
        if classification_only:
            classification = classification_only.group(1)
            continue
        combined = _CLASSIFICATION_TITLE_RE.match(line)
        if combined:
            classification = combined.group(1)
            title_parts.append(combined.group(2).strip())
            continue
        title_parts.append(line)

    appendix_id = appendix_id or appendix_index_to_id(appendix_index)
    title = "\n".join(part for part in title_parts if part).strip()
    title_lines = [""]
    if classification:
        title_lines.append(f"（{classification}）")
    if title:
        title_lines.append(title)
    return {
        "appendix_id": appendix_id,
        "classification": classification,
        "title": title,
        "title_lines": title_lines,
        "soft_break_count": max(0, len(title_lines) - 1),
    }


def _source_heading_level(block: dict) -> int | None:
    if block.get("block_type") == "heading":
        try:
            level = int(block.get("level") or 0)
        except (TypeError, ValueError):
            level = 0
        if level > 0:
            return level
    style_name = str(block.get("source", {}).get("style") or "")
    match = re.match(r"^heading\s*(\d+)$", style_name, re.IGNORECASE)
    return int(match.group(1)) if match else None


def _is_main_numbered_heading(block: dict) -> bool:
    if block.get("block_type") != "heading":
        return False
    numbering = block.get("source", {}).get("numbering", {})
    level = _source_heading_level(block) or 0
    return (
        numbering.get("num_fmt") == "decimal"
        and "%" in str(numbering.get("lvl_text") or "")
        and level >= 2
    )


def _has_page_boundary(blocks: list[dict], index: int) -> bool:
    layout = blocks[index].get("source", {}).get("layout", {})
    if layout.get("page_break_before"):
        return True
    if index > 0:
        previous_layout = blocks[index - 1].get("source", {}).get("layout", {})
        if previous_layout.get("contains_page_break"):
            return True
    return False


def _has_substantive_following_content(blocks: list[dict], index: int) -> bool:
    for block in blocks[index + 1:index + 6]:
        if block.get("block_type") in {"heading", "body", "table", "image", "caption", "list_item"}:
            text = str(block.get("text") or block.get("title") or "").strip()
            if text or block.get("block_type") in {"table", "image"}:
                return True
    return False


def _has_later_same_or_shallower_heading(blocks: list[dict], index: int, level: int) -> bool:
    for block in blocks[index + 1:]:
        if block.get("block_type") == "appendix":
            break
        later_level = _source_heading_level(block)
        if later_level is not None and later_level <= level:
            return True
    return False


def collect_appendix_candidates(model: dict) -> list[dict[str, Any]]:
    blocks = model.setdefault("document", {}).setdefault("blocks", [])
    total = max(1, len(blocks))
    candidates: list[dict[str, Any]] = []

    for index, block in enumerate(blocks):
        source = block.get("source", {})
        raw_text = str(source.get("raw_text") or block.get("text") or block.get("title") or "")
        marker = parse_appendix_marker(raw_text)
        style_role = appendix_role_from_style(source.get("style"))
        existing_appendix = block.get("block_type") == "appendix"
        if not existing_appendix and style_role != "appendix_title" and marker is None:
            continue
        if block.get("block_type") in {"caption", "table", "image", "list_item"}:
            continue

        evidence: list[str] = []
        negative_evidence: list[str] = []
        score = 0.0
        accepted = False

        if existing_appendix or style_role == "appendix_title":
            score = 1.0
            accepted = True
            evidence.append("dedicated_appendix_semantics")
        else:
            score = 0.60
            evidence.append("standalone_appendix_marker")
            if marker and marker.get("appendix_id"):
                score += 0.30
                evidence.append("explicit_appendix_id")
            position_ratio = index / max(1, total - 1)
            if position_ratio >= 0.70:
                score += 0.15
                evidence.append("late_document_region")
            if _has_page_boundary(blocks, index):
                score += 0.20
                evidence.append("page_boundary_before")
            if _has_substantive_following_content(blocks, index):
                score += 0.10
                evidence.append("substantive_following_content")

            heading_level = _source_heading_level(block)
            if heading_level == 1:
                score += 0.15
                evidence.append("root_heading_marker")
            elif heading_level == 2:
                score += 0.05
                evidence.append("second_level_heading_marker")

            if heading_level and not _has_later_same_or_shallower_heading(blocks, index, heading_level):
                score += 0.10
                evidence.append("no_later_peer_heading")

            if _is_main_numbered_heading(block) and (heading_level or 0) >= 3:
                score -= 0.40
                negative_evidence.append("nested_main_numbered_heading")
            accepted = score >= 0.70

        candidates.append({
            "block_index": index,
            "block_id": block.get("id"),
            "source_position": source.get("source_position"),
            "appendix_id": (
                block.get("appendix_id")
                or (marker or {}).get("appendix_id")
            ),
            "marker_title": (marker or {}).get("title"),
            "confidence": round(max(0.0, min(1.0, score)), 3),
            "evidence": evidence,
            "negative_evidence": negative_evidence,
            "accepted": accepted,
        })
    return candidates


def _retype_appendix_title(block: dict, candidate: dict, ordinal: int) -> None:
    raw_text = str(block.get("source", {}).get("raw_text") or block.get("text") or block.get("title") or "")
    title_data = parse_appendix_title(raw_text, ordinal)
    if candidate.get("appendix_id"):
        title_data["appendix_id"] = str(candidate["appendix_id"])
    block["block_type"] = "appendix"
    block["appendix_id"] = title_data["appendix_id"]
    block["title"] = title_data["title"]
    if title_data.get("classification"):
        block["classification"] = title_data["classification"]
    else:
        block.pop("classification", None)
    block["title_lines"] = title_data["title_lines"]
    block["soft_break_count"] = title_data["soft_break_count"]
    block["numbering"] = {"mode": "auto"}
    block["layout"] = {**dict(block.get("layout") or {}), "page_break_before": True}
    block["detection"] = {
        "confidence": candidate["confidence"],
        "evidence": list(candidate["evidence"]),
        "negative_evidence": list(candidate["negative_evidence"]),
    }
    for key in ("text", "role", "level", "list_type", "restart"):
        block.pop(key, None)


def _visible_appendix_heading(text: str, appendix_id: str) -> tuple[int, str] | None:
    match = _APPENDIX_VISIBLE_HEADING_RE.match(text or "")
    if match is None or match.group(1).upper() != appendix_id.upper():
        return None
    level = match.group(2).replace("．", ".").count(".")
    title = match.group(3).strip()
    return level, title


def _reclassify_appendix_headings(blocks: list[dict], start: int, end: int, appendix_id: str) -> None:
    generic_levels = [
        level
        for block in blocks[start + 1:end + 1]
        if block.get("role") != "appendix_heading"
        for level in [_source_heading_level(block)]
        if level is not None
    ]
    base_level = min(generic_levels) if generic_levels else None

    for block in blocks[start + 1:end + 1]:
        text = str(block.get("text") or "")
        visible = _visible_appendix_heading(text, appendix_id)
        relative_level = None
        heading_text = text
        level_source = None
        original_heading_level = _source_heading_level(block)
        if visible is not None:
            relative_level, heading_text = visible
            level_source = "visible_appendix_number"
        elif base_level is not None:
            source_level = original_heading_level
            if source_level is not None:
                relative_level = source_level - base_level + 1
                level_source = "relative_heading_hierarchy"
        if relative_level is None or relative_level not in {1, 2, 3}:
            continue
        block["block_type"] = "heading"
        block["role"] = "appendix_heading"
        block["level"] = relative_level
        block["text"] = heading_text
        block["numbering"] = {"mode": "auto"}
        block.pop("list_type", None)
        block.pop("restart", None)
        block["appendix_heading_detection"] = {
            "source": level_source,
            "source_heading_level": original_heading_level,
        }


def annotate_appendix_ranges(model: dict, parse_report: dict | None = None) -> list[dict]:
    document = model.setdefault("document", {})
    blocks = document.setdefault("blocks", [])
    candidates = collect_appendix_candidates(model)
    accepted_candidates = [candidate for candidate in candidates if candidate["accepted"]]
    for ordinal, candidate in enumerate(accepted_candidates, 1):
        _retype_appendix_title(blocks[candidate["block_index"]], candidate, ordinal)
    starts = [candidate["block_index"] for candidate in accepted_candidates]
    ranges: list[dict] = []

    for ordinal, start in enumerate(starts, 1):
        end = starts[ordinal] - 1 if ordinal < len(starts) else len(blocks) - 1
        title_block = blocks[start]
        appendix_id = str(title_block.get("appendix_id") or appendix_index_to_id(ordinal))
        title_block["appendix_id"] = appendix_id
        _reclassify_appendix_headings(blocks, start, end, appendix_id)
        for block in blocks[start:end + 1]:
            block["appendix_id"] = appendix_id
            context = dict(block.get("context") or {})
            context.update({"scope": "appendix", "appendix_id": appendix_id})
            block["context"] = context

        start_source = blocks[start].get("source", {}).get("source_position")
        end_source = blocks[end].get("source", {}).get("source_position") if end >= start else start_source
        ranges.append({
            "appendix_id": appendix_id,
            "classification": title_block.get("classification"),
            "title": title_block.get("title", ""),
            "start_block_id": blocks[start].get("id"),
            "end_block_id": blocks[end].get("id") if end >= start else blocks[start].get("id"),
            "start_block_index": start,
            "end_block_index": end,
            "start_source_position": start_source,
            "end_source_position": end_source,
            "block_count": max(0, end - start + 1),
            "detection": dict(title_block.get("detection") or {}),
            "page_break_before": bool(title_block.get("layout", {}).get("page_break_before")),
        })

    document["has_appendix"] = bool(ranges)
    document["appendices"] = ranges
    document["appendix_candidates"] = candidates
    if parse_report is not None:
        parse_report["appendix_detection"] = {
            "has_appendix": bool(ranges),
            "appendix_count": len(ranges),
            "candidate_count": len(candidates),
            "candidates": candidates,
            "ranges": ranges,
        }
    return ranges


def _style_by_name(doc, name: str):
    target = normalize_style_name(name)
    return next((style for style in doc.styles if normalize_style_name(style.name) == target), None)


def _style_num_values(style) -> tuple[str | None, str | None]:
    if style is None:
        return None, None
    p_pr = style._element.find(qn("w:pPr"))
    num_pr = p_pr.find(qn("w:numPr")) if p_pr is not None else None
    num_id = num_pr.find(qn("w:numId")) if num_pr is not None else None
    ilvl = num_pr.find(qn("w:ilvl")) if num_pr is not None else None
    return (
        num_id.get(qn("w:val")) if num_id is not None else None,
        ilvl.get(qn("w:val")) if ilvl is not None else None,
    )


def _on_off_value(element) -> bool | None:
    if element is None:
        return None
    value = (element.get(qn("w:val")) or "true").casefold()
    return value not in {"0", "false", "off", "no"}


def _effective_style_value(style, property_group: str, tag: str, attr: str = "val") -> str | None:
    current = style
    while current is not None:
        group = current._element.find(qn(f"w:{property_group}"))
        element = group.find(qn(f"w:{tag}")) if group is not None else None
        if element is not None and element.get(qn(f"w:{attr}")) is not None:
            return element.get(qn(f"w:{attr}"))
        current = current.base_style
    root = style._element.getroottree().getroot()
    default_path = (
        f"w:docDefaults/w:{property_group}Default/w:{property_group}/w:{tag}"
    )
    element = root.find(default_path, namespaces={"w": qn("w:styles").split("}")[0].lstrip("{")})
    if element is not None:
        return element.get(qn(f"w:{attr}"))
    return None


def _effective_style_font(style, attr: str) -> str | None:
    current = style
    while current is not None:
        r_pr = current._element.find(qn("w:rPr"))
        fonts = r_pr.find(qn("w:rFonts")) if r_pr is not None else None
        if fonts is not None and fonts.get(qn(f"w:{attr}")) is not None:
            return fonts.get(qn(f"w:{attr}"))
        current = current.base_style
    root = style._element.getroottree().getroot()
    fonts = root.find(
        "w:docDefaults/w:rPrDefault/w:rPr/w:rFonts",
        namespaces={"w": qn("w:styles").split("}")[0].lstrip("{")},
    )
    if fonts is not None:
        return fonts.get(qn(f"w:{attr}"))
    return None


def _effective_style_bold(style) -> bool | None:
    current = style
    while current is not None:
        r_pr = current._element.find(qn("w:rPr"))
        bold = r_pr.find(qn("w:b")) if r_pr is not None else None
        if bold is not None:
            return _on_off_value(bold)
        current = current.base_style
    return None


def _audit_style_format(style_name: str, style, *, title: bool, first_heading: bool) -> list[dict[str, Any]]:
    issues = []
    expected_size = APPENDIX_TITLE_SIZE_HALF_POINTS if title else APPENDIX_HEADING_SIZE_HALF_POINTS
    for font_attr in ("eastAsia", "ascii", "hAnsi"):
        actual = _effective_style_font(style, font_attr)
        if actual != APPENDIX_TITLE_FONT:
            issues.append({
                "type": "appendix_style_font",
                "style": style_name,
                "font_attribute": font_attr,
                "expected": APPENDIX_TITLE_FONT,
                "actual": actual,
            })
    actual_size = _effective_style_value(style, "rPr", "sz")
    if actual_size != str(expected_size):
        issues.append({
            "type": "appendix_style_size",
            "style": style_name,
            "expected": expected_size,
            "actual": actual_size,
        })
    actual_bold = _effective_style_bold(style)
    if actual_bold not in {None, False}:
        issues.append({
            "type": "appendix_style_bold",
            "style": style_name,
            "expected": False,
            "actual": actual_bold,
        })
    actual_line = _effective_style_value(style, "pPr", "spacing", "line")
    actual_line_rule = _effective_style_value(style, "pPr", "spacing", "lineRule")
    if actual_line != str(APPENDIX_LINE_SPACING_TWIPS) or actual_line_rule != "auto":
        issues.append({
            "type": "appendix_style_line_spacing",
            "style": style_name,
            "expected_line": APPENDIX_LINE_SPACING_TWIPS,
            "expected_rule": "auto",
            "actual_line": actual_line,
            "actual_rule": actual_line_rule,
        })
    if title or first_heading:
        before = _effective_style_value(style, "pPr", "spacing", "before")
        after = _effective_style_value(style, "pPr", "spacing", "after")
        if before != str(APPENDIX_HALF_LINE_SPACING_TWIPS) or after != str(APPENDIX_HALF_LINE_SPACING_TWIPS):
            issues.append({
                "type": "appendix_style_half_line_spacing",
                "style": style_name,
                "expected_before": APPENDIX_HALF_LINE_SPACING_TWIPS,
                "expected_after": APPENDIX_HALF_LINE_SPACING_TWIPS,
                "actual_before": before,
                "actual_after": after,
            })
    if title:
        alignment = _effective_style_value(style, "pPr", "jc")
        if alignment != "center":
            issues.append({
                "type": "appendix_title_alignment",
                "style": style_name,
                "expected": "center",
                "actual": alignment,
            })
    return issues


def _soft_break_types(paragraph) -> list[str]:
    result = []
    for br in paragraph._p.iter(qn("w:br")):
        result.append(br.get(qn("w:type")) or "textWrapping")
    return result


def _rendered_page_break_count_before(paragraphs, index: int) -> int:
    count = 0
    paragraph = paragraphs[index]
    p_pr = paragraph._p.find(qn("w:pPr"))
    page_break_before = p_pr.find(qn("w:pageBreakBefore")) if p_pr is not None else None
    if page_break_before is not None:
        value = (page_break_before.get(qn("w:val")) or "true").casefold()
        if value not in {"0", "false", "off", "no"}:
            count += 1
    if index <= 0:
        return count
    count += sum(
        element.get(qn("w:type")) == "page"
        for element in paragraphs[index - 1]._p.iter(qn("w:br"))
    )
    return count


def _numbering_levels(doc, num_id: str) -> dict[int, dict[str, str | None]]:
    root = doc.part.numbering_part.element
    num = next((item for item in root.findall(qn("w:num")) if item.get(qn("w:numId")) == num_id), None)
    if num is None:
        return {}
    abstract_ref = num.find(qn("w:abstractNumId"))
    abstract_id = abstract_ref.get(qn("w:val")) if abstract_ref is not None else None
    abstract = next(
        (item for item in root.findall(qn("w:abstractNum")) if item.get(qn("w:abstractNumId")) == abstract_id),
        None,
    )
    levels: dict[int, dict[str, str | None]] = {}
    for level in abstract.findall(qn("w:lvl")) if abstract is not None else []:
        ilvl = int(level.get(qn("w:ilvl"), "0"))
        num_fmt = level.find(qn("w:numFmt"))
        lvl_text = level.find(qn("w:lvlText"))
        p_style = level.find(qn("w:pStyle"))
        levels[ilvl] = {
            "num_fmt": num_fmt.get(qn("w:val")) if num_fmt is not None else None,
            "lvl_text": lvl_text.get(qn("w:val")) if lvl_text is not None else None,
            "p_style": p_style.get(qn("w:val")) if p_style is not None else None,
        }
    return levels


def audit_appendix_contract(doc, template_profile: dict | None = None) -> dict[str, Any]:
    paragraphs = list(doc.paragraphs)
    title_indices = [
        index for index, paragraph in enumerate(paragraphs, 1)
        if normalize_style_name(paragraph.style.name if paragraph.style else "")
        == normalize_style_name(APPENDIX_TITLE_STYLE)
    ]
    issues: list[dict[str, Any]] = []
    style_issues: list[dict[str, Any]] = []
    title_issues: list[dict[str, Any]] = []
    page_break_issues: list[dict[str, Any]] = []

    styles = {name: _style_by_name(doc, name) for name in APPENDIX_STYLE_ROLES}
    for name, style in styles.items():
        if style is None:
            style_issues.append({"type": "missing_style", "style": name})
        else:
            style_issues.extend(_audit_style_format(
                name,
                style,
                title=name == APPENDIX_TITLE_STYLE,
                first_heading=name == APPENDIX_HEADING_STYLES[1],
            ))

    expected_levels = {
        APPENDIX_TITLE_STYLE: 0,
        **{style_name: level for level, style_name in APPENDIX_HEADING_STYLES.items()},
    }
    bound_num_ids = set()
    for style_name, expected_level in expected_levels.items():
        num_id, ilvl = _style_num_values(styles.get(style_name))
        if num_id:
            bound_num_ids.add(num_id)
        if num_id is None or ilvl != str(expected_level):
            style_issues.append({
                "type": "style_numbering_binding",
                "style": style_name,
                "expected_ilvl": expected_level,
                "actual_num_id": num_id,
                "actual_ilvl": ilvl,
            })

    if len(bound_num_ids) > 1:
        style_issues.append({"type": "appendix_styles_use_different_num_ids", "num_ids": sorted(bound_num_ids)})
    if bound_num_ids:
        num_id = sorted(bound_num_ids)[0]
        levels = _numbering_levels(doc, num_id)
        expected_numbering = {
            0: ("upperLetter", "附  录  %1"),
            1: ("decimal", "%1.%2 "),
            2: ("decimal", "%1.%2.%3 "),
            3: ("decimal", "%1.%2.%3.%4 "),
        }
        for level, (num_fmt, lvl_text) in expected_numbering.items():
            actual = levels.get(level, {})
            if actual.get("num_fmt") != num_fmt or actual.get("lvl_text") != lvl_text:
                style_issues.append({
                    "type": "appendix_numbering_level",
                    "level": level,
                    "expected_num_fmt": num_fmt,
                    "expected_lvl_text": lvl_text,
                    "actual": actual,
                })

    for index in title_indices:
        paragraph = paragraphs[index - 1]
        break_types = _soft_break_types(paragraph)
        text_lines = paragraph.text.replace("\r", "").split("\n")
        expected_break_types = ["textWrapping"] * max(0, len(text_lines) - 1)
        if break_types != expected_break_types:
            title_issues.append({
                "type": "appendix_title_soft_breaks",
                "paragraph": index,
                "expected": expected_break_types,
                "actual": break_types,
            })
        if text_lines[0] != "":
            if not break_types:
                title_issues.append({
                    "type": "appendix_title_soft_breaks",
                    "paragraph": index,
                    "expected": ["textWrapping"],
                    "actual": break_types,
                })
            title_issues.append({
                "type": "appendix_title_line_structure",
                "paragraph": index,
                "actual_lines": text_lines,
            })
        page_break_count = _rendered_page_break_count_before(paragraphs, index - 1)
        if page_break_count == 0:
            page_break_issues.append({
                "type": "appendix_missing_page_break_before",
                "paragraph": index,
            })
        elif page_break_count > 1:
            page_break_issues.append({
                "type": "appendix_duplicate_page_break_before",
                "paragraph": index,
                "count": page_break_count,
            })
        for run_index, run in enumerate(paragraph.runs):
            if run.bold is True:
                title_issues.append({
                    "type": "appendix_title_direct_bold",
                    "paragraph": index,
                    "run": run_index,
                })
            if run.font.size is not None and int(round(run.font.size.pt * 2)) != APPENDIX_TITLE_SIZE_HALF_POINTS:
                title_issues.append({
                    "type": "appendix_title_direct_size",
                    "paragraph": index,
                    "run": run_index,
                    "actual_half_points": int(round(run.font.size.pt * 2)),
                })

    expected_heading_styles = set(APPENDIX_HEADING_STYLES.values())
    appendix_heading_paragraphs = [
        {
            "paragraph": index,
            "style": paragraph.style.name if paragraph.style else "",
            "text": paragraph.text[:120],
        }
        for index, paragraph in enumerate(paragraphs, 1)
        if (paragraph.style.name if paragraph.style else "") in expected_heading_styles
    ]
    if title_indices:
        toc_instructions = [
            element.text or ""
            for element in doc.element.body.iter(qn("w:instrText"))
            if "TOC" in (element.text or "").upper()
        ]
        expected_tokens = [
            f"{style_name},{level}"
            for style_name, level in TOC_CUSTOM_STYLE_LEVELS.items()
        ]
        if not any(all(token in instruction for token in expected_tokens) for instruction in toc_instructions):
            issues.append({
                "type": "appendix_toc_style_mapping",
                "expected": TOC_CUSTOM_STYLE_LEVELS,
                "instructions": toc_instructions,
            })
    issues.extend(style_issues)
    issues.extend(title_issues)
    issues.extend(page_break_issues)
    return {
        "has_appendix": bool(title_indices),
        "appendix_count": len(title_indices),
        "title_paragraphs": title_indices,
        "appendix_heading_paragraphs": appendix_heading_paragraphs,
        "style_issues": style_issues,
        "title_issues": title_issues,
        "page_break_issues": page_break_issues,
        "issues": issues,
        "passed": not issues,
    }


def audit_appendix_preservation(
    doc,
    model: dict | None,
    template_profile: dict | None = None,
) -> dict[str, Any]:
    contract = audit_appendix_contract(doc, template_profile)
    blocks = (model or {}).get("document", {}).get("blocks", [])
    expected_titles = [block for block in blocks if block.get("block_type") == "appendix"]
    expected_headings = {
        level: [
            str(block.get("text") or "")
            for block in blocks
            if block.get("block_type") == "heading"
            and block.get("role") == "appendix_heading"
            and int(block.get("level") or 0) == level
        ]
        for level in APPENDIX_HEADING_STYLES
    }
    expected_captions = Counter(
        str(block.get("source", {}).get("raw_text") or "").strip()
        for block in blocks
        if block.get("block_type") == "caption"
        and block.get("appendix_id")
        and str(block.get("source", {}).get("raw_text") or "").strip()
    )
    rendered_titles = [
        paragraph for paragraph in doc.paragraphs
        if appendix_role_from_style(paragraph.style.name if paragraph.style else "") == "appendix_title"
    ]
    rendered_headings = {
        level: [
            paragraph.text
            for paragraph in doc.paragraphs
            if normalize_style_name(paragraph.style.name if paragraph.style else "")
            == normalize_style_name(style_name)
        ]
        for level, style_name in APPENDIX_HEADING_STYLES.items()
    }
    rendered_captions = Counter(
        paragraph.text.strip()
        for paragraph in doc.paragraphs
        if normalize_style_name(paragraph.style.name if paragraph.style else "")
        == normalize_style_name("Caption")
    )
    issues = list(contract["issues"])

    if len(rendered_titles) != len(expected_titles):
        issues.append({
            "type": "appendix_title_count_mismatch",
            "expected": len(expected_titles),
            "actual": len(rendered_titles),
        })
    for ordinal, expected in enumerate(expected_titles, 1):
        if ordinal > len(rendered_titles):
            break
        actual = parse_appendix_title(rendered_titles[ordinal - 1].text, ordinal)
        for field in ("appendix_id", "classification", "title", "title_lines"):
            if actual.get(field) != expected.get(field):
                issues.append({
                    "type": "appendix_title_semantics_mismatch",
                    "appendix": expected.get("appendix_id"),
                    "field": field,
                    "expected": expected.get(field),
                    "actual": actual.get(field),
                })

    for level in APPENDIX_HEADING_STYLES:
        if rendered_headings[level] != expected_headings[level]:
            issues.append({
                "type": "appendix_heading_preservation",
                "level": level,
                "expected": expected_headings[level],
                "actual": rendered_headings[level],
            })
    for caption_text, expected_count in expected_captions.items():
        actual_count = rendered_captions.get(caption_text, 0)
        if actual_count < expected_count:
            issues.append({
                "type": "appendix_caption_preservation",
                "text": caption_text,
                "expected_count": expected_count,
                "actual_count": actual_count,
            })

    return {
        "source_has_appendix": bool(expected_titles),
        "source_appendix_count": len(expected_titles),
        "rendered_appendix_count": len(rendered_titles),
        "expected_heading_counts": {
            str(level): len(values) for level, values in expected_headings.items()
        },
        "rendered_heading_counts": {
            str(level): len(values) for level, values in rendered_headings.items()
        },
        "expected_caption_counts": dict(expected_captions),
        "rendered_appendix_caption_counts": {
            text: rendered_captions.get(text, 0) for text in expected_captions
        },
        "contract": contract,
        "issues": issues,
        "passed": not issues,
    }
