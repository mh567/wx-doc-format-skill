from __future__ import annotations

import re


HEADING_PATTERNS = [
    (re.compile(r"^第([一二三四五六七八九十百千零\d]+)章\s*"), 1),
    (re.compile(r"^第([一二三四五六七八九十百千零\d]+)节\s*"), 2),
    (re.compile(r"^第([一二三四五六七八九十百千零\d]+)条\s*"), 3),
    (re.compile(r"^([一二三四五六七八九十百千零\d]+)[、，,\.\s]\s*"), 1),
    (re.compile(r"^\(([一二三四五六七八九十百千零\d]+)\)\s*"), 2),
    (re.compile(r"^（([一二三四五六七八九十百千零\d]+)）\s*"), 2),
]

LIST_PATTERNS = [
    (re.compile(r"^[a-zA-Z]\)\s*"), "letter"),
    (re.compile(r"^\d+\.\.?\s*"), "decimal"),
    (re.compile(r"^\d+\)\s*"), "decimal"),
    (re.compile(r"^（\d+）\s*"), "decimal"),
    (re.compile(r"^\(\d+\)\s*"), "decimal"),
    (re.compile(r"^[•·]\s*"), "bullet2"),
    (re.compile(r"^——\s*"), "dash"),
    (re.compile(r"^—\s*"), "dash"),
    (re.compile(r"^[●○]\s*"), "bullet2"),
]

DATE_LIKE = re.compile(r"^\d{4}\s*年\s*\d{1,2}\s*月")
TOC_TITLES = {"目次", "目录", "目  次"}
CAPTION_PATTERN = re.compile(r"^([图表])[：:\s]*(?:(?:[A-Z]\.)?\d+(?:[-\.](?:[A-Z]\.)?\d+)*\s*)?(.*)")
FORMULA_PATTERN = re.compile(r"^(S|S\.)?\(?\d+(?:\.\d+)*\)?[：:]\s*")
NOTE_FORMULA_LINE = re.compile(r"^（\d+(?:\.\d+)*）\s*")


def strip_heading_marker(text: str) -> str:
    text = text.strip()
    text = re.sub(r"^\d+(?:\.\d+)+\s*", "", text)
    # Trailing page numbers or tabs: "概述	1", "依据文件 1"
    text = re.sub(r"[\t ]+\d+$", "", text)
    text = re.sub(r"^\d+\s+", "", text)
    for pattern, _ in HEADING_PATTERNS:
        text = pattern.sub("", text, count=1)
    return text.strip()


def heading_level_from_text(text: str) -> int | None:
    text = text.strip()
    if len(text) > 30:
        return None
    m = re.match(r"^\d+(?:\.\d+)+\s+", text)
    if m:
        parts = m.group().strip().split(".")
        return min(len(parts), 6)
    m = re.match(r"^\d+\s+", text)
    if m:
        return 1
    for pattern, level in HEADING_PATTERNS:
        if pattern.match(text):
            return level
    return None


def heading_style_for_level(level: int) -> str:
    return f"Heading {level}"


def heading_level_from_style(style_name: str) -> int | None:
    if style_name and style_name.startswith("Heading"):
        try:
            return int(style_name.split()[-1])
        except (ValueError, IndexError):
            return None
    return None


def normalize_heading_style_level(style_name: str | None, level_shift: int = 0) -> str | None:
    if not style_name:
        return None
    level = heading_level_from_style(style_name)
    if level is None:
        return style_name
    new_level = max(1, level - level_shift)
    return heading_style_for_level(new_level)


def resolved_heading_level(style: str | None, num_level: int | None, text: str) -> int:
    heading = heading_level_from_style(style or "")
    if heading is not None:
        return heading
    if num_level is not None:
        return num_level + 1
    inferred = heading_level_from_text(text)
    if inferred is not None:
        return inferred
    return 1


def heading_number_source(num_id: int | None, num_level: int | None, text: str) -> str:
    if num_id is not None:
        return "docx-numbering"
    if heading_level_from_text(text) is not None:
        return "docx-text"
    return "assumed"


def existing_heading_number(text: str) -> bool:
    if not text:
        return False
    if re.match(r"^\d+(?:\.\d+)+\s+", text):
        return True
    if re.match(r"^\d+\s+", text):
        return True
    for pattern, _ in HEADING_PATTERNS:
        if pattern.match(text):
            return True
    return False


def strip_manual_number(text: str) -> str:
    return strip_heading_marker(text)


def compact_heading_text(text: str) -> str:
    if " " in text or "　" in text:
        return text.split()[0]
    return text


def is_compact_function_heading_text(text: str) -> bool:
    return len(text) <= 16 and not text.endswith(("。", "；", ";", "，", ","))


def is_compact_numbered_function_heading(text: str) -> bool:
    text = text.strip()
    m = re.match(r"^\d+(?:\.\d+)+\s*", text)
    if not m:
        return False
    rest = text[m.end():].strip()
    if not rest:
        return False
    return is_compact_function_heading_text(rest)


def looks_like_list_item(text: str) -> bool:
    for pattern, _ in LIST_PATTERNS:
        if pattern.match(text):
            return True
    return False


def list_style_for_text(text: str) -> str:
    kind = list_kind_for_text(text)
    if kind in {"letter", "decimal"}:
        return "1.1一级列项-编号"
    if kind == "dash":
        return "1.2一级列项-无编号"
    if kind == "bullet2":
        return "2.2二级列项-无编号"
    return "1.1一级列项-编号"


def list_kind_for_text(text: str) -> str:
    for pattern, kind in LIST_PATTERNS:
        if pattern.match(text):
            return kind
    return "letter"


def strip_list_marker(text: str) -> str:
    for pattern, _ in LIST_PATTERNS:
        text = pattern.sub("", text, count=1)
    # Trailing page numbers or tabs: "术语和定义\t2"
    text = re.sub(r"[\t ]+\d+$", "", text)
    return text.strip()


def list_level_from_text(text: str, fallback: int | None = None) -> int:
    kind = list_kind_for_text(text)
    if kind in {"decimal", "bullet2"}:
        return 1
    return fallback or 0


def compact_heading_text(text: str) -> str:
    text = strip_heading_marker(text).lstrip("▲").strip().rstrip("。．.")
    return text


def is_compact_function_heading_text(text: str) -> bool:
    clean_text = compact_heading_text(text)
    if not clean_text or len(clean_text) > 24:
        return False
    if re.search(r"[，,；;：:]", clean_text):
        return False
    if re.match(r"^(支持|要求|依据|基础属性|权限属性|联系属性|组织属性|岗位属性)", clean_text):
        return False
    return True


def is_compact_numbered_function_heading(text: str) -> bool:
    return bool(re.match(r"^\d+[.．]\S", text)) and is_compact_function_heading_text(text)


def is_formula_text(text: str) -> bool:
    if FORMULA_PATTERN.match(text):
        return True
    if NOTE_FORMULA_LINE.match(text):
        return True
    return False


def is_appendix_title(text: str) -> bool:
    return bool(re.match(r"^附\s*录\s", text)) or bool(re.match(r"^附录([A-Z]|[一二三四五六七八九十百千零])", text))


def is_caption_text(text: str) -> bool:
    return bool(CAPTION_PATTERN.match(text))


def is_date_like_text(text: str) -> bool:
    return bool(DATE_LIKE.match(text))


def is_toc_title(text: str) -> bool:
    return text.strip() in TOC_TITLES


_TOC_ENTRY_RE = re.compile(r'[\t.·…]{1,}\d{1,4}$')
def is_toc_entry(text: str) -> bool:
    """Detect TOC entry paragraphs (title<tab/dots>page_number)."""
    t = text.strip()
    return len(t) < 200 and bool(_TOC_ENTRY_RE.search(t))


def is_front_matter_text(paragraph, text: str) -> bool:
    style = paragraph.style.name if paragraph.style is not None else ""
    if style.startswith("Heading"):
        return True
    if is_date_like_text(text) or is_toc_title(text):
        return True
    # Cover-page metadata prefixes — "版本：", "文档编号：", "密级：", etc.
    if re.match(r'^(版本|Version|文档编号|文件编号|密级|状态|作者|日期|修订|审核|批准|修订记录)\s*[：:]', text):
        return True
    # Centered short/medium text on the cover page is front matter.
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER and len(text) <= 80:
        return True
    return False


def clean_note_prefix(text: str) -> str:
    for prefix in ("**备注：**", "**编写提示：**", "备注：", "编写提示：", "【备注提示】", "【编写样例】"):
        if text.startswith(prefix):
            return text[len(prefix):].strip()
    return text


def paragraph_direct_format_score(paragraph) -> tuple[bool, int | None]:
    """Return (mostly_bold, max_size_pt).

    mostly_bold  — True when over half of runs are bold.
    max_size_pt — largest font size in points, or None.
    """
    try:
        bold_count = 0
        run_count = 0
        max_size = 0.0
        for run in paragraph.runs:
            run_count += 1
            if run.bold:
                bold_count += 1
            if run.font.size:
                sz = run.font.size.pt if hasattr(run.font.size, "pt") else float(run.font.size)
                if sz > max_size:
                    max_size = sz
        if run_count == 0:
            return False, None
        mostly_bold = bold_count > run_count / 2
        return mostly_bold, int(max_size) if max_size > 0 else None
    except Exception:
        return False, None


def looks_like_visual_heading(paragraph) -> bool:
    try:
        text = paragraph.text.strip()
        if not text:
            return False
        if len(text) > 30:
            return False
        if text.endswith(("。", "；", ";")):
            return False
        if is_date_like_text(text):
            return False
        if is_toc_title(text):
            return False
        # Cover-page metadata should never be treated as a heading.
        if re.match(r'^(版本|Version|文档编号|文件编号|密级|状态|作者|日期|修订|审核|批准)\s*[：:]', text):
            return False
        # Parenthetical annotations with specific keywords are not headings.
        if re.search(r'[（(][^）)]*(仅供参考|示例|待补充|待定)[）)]', text):
            return False
        # Attribute-like descriptions (e.g. "身份状态：在职、离职") are not headings.
        if re.search(r'[：:].+[，,、]', text):
            return False
        # API-doc section labels are not headings.
        _api_labels = {
            "请求参数", "请求示例", "返回参数", "返回示例",
            "接口说明", "请求体", "响应体", "成功响应", "失败响应",
        }
        if text.rstrip("：:") in _api_labels:
            return False
        if not text[0].isalnum() and not (text[0] >= "一" and text[0] <= "鿿"):
            return False
        mostly_bold, max_size = paragraph_direct_format_score(paragraph)
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        centered = paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
        return mostly_bold or centered or (max_size is not None and max_size >= 14)
    except Exception:
        return False


def caption_parts(text: str) -> tuple[str, str | None, str | None, str]:
    """Extract caption type and text from a caption string.
    Always discards the number (handled by auto-numbering in the template).
    Returns (caption_type, label, None, caption_text).
    """
    match = CAPTION_PATTERN.match(text.strip())
    if not match:
        return "unknown", None, None, text.strip()
    label, caption_text = match.groups()
    caption_text = (caption_text or "").strip()
    # "图表1" → label=图, caption_text=表1: ambiguous, treat as one token
    # and return empty caption_text (it has no descriptive title, just a number)
    if caption_text.startswith(("表", "图")):
        return ("figure" if label == "图" else "table"), label, None, ""
    return ("figure" if label == "图" else "table"), label, None, caption_text


def model_list_type_from_style(style_name: str) -> tuple[int, str, bool]:
    if "1.1一级列项-编号" in style_name:
        return 0, "lower_letter_paren", False
    if "2.1二级列项-有编号" in style_name:
        return 1, "decimal_paren", False
    if "1.2一级列项-无编号" in style_name:
        return 0, "dash", False
    if "2.2二级列项-无编号" in style_name:
        return 1, "bullet_dot", False
    return 0, "lower_letter_paren", False


def table_rows_for_model(table, table_type: str, header_rows: int) -> list[list[dict]]:
    rows = []
    for ri, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            if table_type == "code_sample":
                cell_role = "code"
            elif ri < header_rows:
                cell_role = "header"
            else:
                cell_role = "body"
            cells.append({"text": text, "cell_role": cell_role})
        rows.append(cells)
    return rows


def scan_non_text_objects(src: Path) -> dict:
    counts = {"media_files": 0, "relationships": 0, "drawings": 0}
    try:
        import zipfile
        from pathlib import Path
        with zipfile.ZipFile(src) as zf:
            names = zf.namelist()
            media = [n for n in names if n.startswith("word/media/") and not n.endswith("/")]
            counts["media_files"] = len(media)
            if "word/document.xml" in names:
                from lxml import etree
                doc_xml = zf.read("word/document.xml")
                root = etree.fromstring(doc_xml)
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                counts["drawings"] = len(root.findall(".//w:drawing", ns))
        if "word/_rels/document.xml.rels" in names:
            from lxml import etree
            rels_xml = zf.read("word/_rels/document.xml.rels")
            rels_root = etree.fromstring(rels_xml)
            ns = {"": "http://schemas.openxmlformats.org/package/2006/relationships"}
            counts["relationships"] = len(rels_root.findall(":Relationship", ns))
    except Exception:
        pass
    return counts


def build_document_model_from_output(doc, source_path: Path, report: dict) -> dict:
    from document_model import (
        new_document_model, append_block, heading_block, body_block,
        list_item_block, table_block, caption_block, image_block, appendix_block
    )
    from pathlib import Path
    model = new_document_model(source_path, "docx", report.get("skill_version", "unknown"))
    block_index = 1
    images_found = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else ""
        block_id = f"b{block_index:04d}"
        if style_name.startswith("Heading"):
            level = heading_level_from_style(style_name) or 1
            append_block(model, heading_block(block_id, text, level))
        elif "列项" in style_name:
            level, list_type, restart = model_list_type_from_style(style_name)
            append_block(model, list_item_block(block_id, text, level, list_type, restart=restart))
        elif "文档标题" in style_name:
            append_block(model, heading_block(block_id, text, 0, role="title"))
        elif "附录" in style_name:
            append_block(model, appendix_block(block_id, text))
        elif is_caption_text(text):
            cap_type, label, raw_number, cap_text = caption_parts(text)
            append_block(model, caption_block(block_id, cap_text, cap_type, label=label, raw_number=raw_number))
        elif style_name in ("Caption", "题注"):
            append_block(model, caption_block(block_id, text, "unknown"))
        elif paragraph_has_graphics(paragraph):
            append_block(model, image_block(block_id, alt_text=text or ""))
            images_found += 1
        else:
            role = "body"
            if "注-" in style_name:
                role = "note"
            append_block(model, body_block(block_id, text, source={"role": role}))
        block_index += 1
    for table_idx, table in enumerate(doc.tables, 1):
        table_type = "data"
        header_rows = 1
        if looks_like_code_sample_table(table):
            table_type = "code_sample"
            header_rows = 0
        rows = table_rows_for_model(table, table_type, header_rows)
        block_id = f"b{block_index:04d}"
        append_block(model, table_block(block_id, table_type, rows, header_rows=header_rows))
        block_index += 1
    if block_index <= 2:
        text = "".join(p.text for p in doc.paragraphs)
        if not text.strip() and not doc.tables:
            return model
    report["non_text_objects"]["images_in_output"] = images_found
    return model


# — XML namespaces and attribute names used when scanning for image rIds —
_R_EMBED_QN  = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
_R_ID_QN     = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
_R_PICT_QN   = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}pict"

_IMAGE_RELTYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"


def paragraph_has_graphics(paragraph) -> bool:
    """True when the paragraph contains at least one image rId reference.

    Instead of hard‑coding every XML element name (a:blip, v:imagedata,
    w:drawing, wp:inline, o:OLEObject, w:pict …) we simply walk the full
    XML tree and look for the standard image‑relationship attributes.
    This handles every embedding mechanism that Office documents use.
    """
    try:
        el = paragraph._element
        part = paragraph.part
    except Exception:
        return False
    return _element_has_image_rid(el, part)


def _element_has_image_rid(el, part) -> bool:
    """Return True if *el* contains an r:embed / r:id / r:pict that points
    to an image relationship in *part*."""
    for _candidate in _iter_image_rids(el, part):
        return True
    return False


def _iter_image_rids(el, part):
    """Yield (attribute_name, rId, rel) for every image rId found in *el*."""
    import re
    try:
        xml_str = el.xpath("string()")  # triggers etree serialization
    except Exception:
        pass
    # Scan all descendant elements for the three known attribute names.
    for descendant in el.iter():
        for attr_name in (_R_EMBED_QN, _R_ID_QN, _R_PICT_QN):
            rid = descendant.get(attr_name)
            if not rid:
                continue
            try:
                rel = part.rels[rid]
            except Exception:
                # Try by iterating
                for r in part.rels.values():
                    if r.rId == rid:
                        rel = r
                        break
                else:
                    continue
            if rel.reltype == _IMAGE_RELTYPE:
                yield (attr_name, rid, rel)


def iter_blocks(doc):
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def paragraph_num_info(paragraph) -> tuple[int | None, int | None]:
    try:
        p_pr = paragraph._p.pPr
        if p_pr is None:
            return None, None
        num_pr = p_pr.numPr
        if num_pr is None:
            return None, None
        ilvl = None
        num_id = None
        try:
            ilvl = num_pr.ilvl.val if num_pr.ilvl is not None else None
        except AttributeError:
            pass
        try:
            num_id = num_pr.numId.val if num_pr.numId is not None else None
        except AttributeError:
            pass
        return ilvl, num_id
    except Exception:
        return None, None


def paragraph_direct_num_info(paragraph) -> tuple[int | None, int | None]:
    return paragraph_num_info(paragraph)


def paragraph_numbering_descriptor(paragraph) -> tuple[str | None, str | None]:
    ilvl, num_id = paragraph_num_info(paragraph)
    return (str(num_id) if num_id is not None else None,
            str(ilvl) if ilvl is not None else None)


def source_numbering_heading_level(paragraph) -> int | None:
    ilvl, num_id = paragraph_num_info(paragraph)
    if num_id is None:
        return None
    try:
        numbering = paragraph._element.getroottree().getroot().find(".//" + "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numbering")
        if numbering is None:
            return None
        nums = numbering.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num")
        for num in nums:
            if num.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId") == str(num_id):
                abstract_id = num.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId")
                if abstract_id is not None:
                    abstracts = numbering.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNum")
                    for ab in abstracts:
                        if ab.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId") == abstract_id.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"):
                            lvl = ab.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lvl")
                            if lvl is not None:
                                p_style = lvl.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle")
                                if p_style is not None:
                                    from docx_pipeline import heading_level_from_pstyle
                                    from docx.oxml.ns import qn
                                    return heading_level_from_pstyle(p_style.get(qn("w:val")))
    except Exception:
        pass
    return None


def source_heading_level_shift(doc) -> int:
    first_heading_level = None
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        level = heading_level_from_style(style_name)
        if level is not None and level < first_heading_level if first_heading_level else True:
            first_heading_level = level
            break
    if first_heading_level is not None and first_heading_level > 1:
        return first_heading_level - 1
    return 0




def looks_like_code_sample_table(table) -> bool:
    try:
        rows = table.rows
        if len(rows) < 2:
            return False
        cell_text = " ".join(cell.text for cell in rows[0].cells).strip().lower()
        return any(kw in cell_text for kw in ["请求", "响应", "http", "json", "参数", "字段", "说明", "名称"])
    except Exception:
        return False


def looks_like_api_example_table(table) -> bool:
    """Return True if table is a single-cell API request/response example."""
    try:
        rows = table.rows
        if len(rows) != 1 or len(rows[0].cells) != 1:
            return False
        text = rows[0].cells[0].text.strip()
        # JSON block, XML, HTTP request/response, or code snippet
        if text.startswith(("{", "[", "<")):
            return True
        upper60 = text[:60].upper()
        for kw in ("POST ", "GET ", "PUT ", "DELETE ", "PATCH ",
                   "HTTP/1.", "<<HEADER>>", "CONTENT-TYPE:", "CONTENT-MD5",
                   "JSON", "XML", "<?XML", "PLAIN TEXT"):
            if kw in upper60:
                return True
        # JSON or XML body anywhere in first 200 chars
        body200 = text[:200]
        if '{' in body200 and ('"' in body200 or ':' in body200):
            return True
        return False
    except Exception:
        return False


def set_table_autofit_to_window(table) -> None:
    try:
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else None
        if tbl_pr is None:
            return
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            from docx.oxml import OxmlElement
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:type"), "pct")
        tbl_w.set(qn("w:w"), "5000")
    except Exception:
        pass


def set_template_table_properties(table, row_height_cm: float, row_height_rule: str,
                                     table_body_style: str | None = None) -> None:
    """Compatibility wrapper around the canonical table normalizer."""
    from table_formatting import normalize_table

    profile = None
    if table_body_style:
        profile = {"resolved_styles": {"table_body": table_body_style}}
    normalize_table(table, profile, row_height_cm, row_height_rule)
