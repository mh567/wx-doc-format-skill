from __future__ import annotations

from dataclasses import dataclass

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm


TABLE_DEFAULT_CELL_MARGINS = {"top": 0, "bottom": 0, "left": 108, "right": 108}
TABLE_BORDER_NAMES = ("top", "left", "bottom", "right", "insideH", "insideV")
PARAGRAPH_ALLOWED_TAGS = {qn("w:pStyle"), qn("w:jc")}
RUN_ALLOWED_TAGS = {qn("w:rStyle"), qn("w:lang"), qn("w:bCs"), qn("w:iCs")}


@dataclass(frozen=True)
class TableFormatPolicy:
    role: str
    paragraph_style: str
    alignment: object
    row_height_cm: float
    row_height_rule: str
    table_style_id: str | None
    table_look: dict[str, str]


def table_role(table_type: str | None) -> str:
    if table_type in {"code_sample", "callout", "layout"}:
        return str(table_type)
    return "data"


def build_table_policy(
    profile: dict | None,
    row_height_cm: float,
    row_height_rule: str,
    *,
    role: str = "data",
) -> TableFormatPolicy:
    profile = profile or {}
    resolved = profile.get("resolved_styles", {})
    table_style = profile.get("table_style", {})
    return TableFormatPolicy(
        role=table_role(role),
        paragraph_style=resolved.get("table_body", "表正文"),
        alignment=(
            WD_ALIGN_PARAGRAPH.LEFT
            if table_role(role) in {"code_sample", "callout"}
            else WD_ALIGN_PARAGRAPH.CENTER
        ),
        row_height_cm=row_height_cm,
        row_height_rule=row_height_rule,
        table_style_id=table_style.get("style_id"),
        table_look=dict(table_style.get("tbl_look") or {}),
    )


def _get_or_add(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _set_attr(element, name: str, value: str) -> bool:
    key = qn(f"w:{name}")
    if element.get(key) == value:
        return False
    element.set(key, value)
    return True


def _normalize_table_properties(table, policy: TableFormatPolicy, path: str) -> list[dict]:
    corrections: list[dict] = []
    tbl_pr = table._tbl.tblPr

    if policy.table_style_id:
        tbl_style = _get_or_add(tbl_pr, "w:tblStyle")
        if _set_attr(tbl_style, "val", str(policy.table_style_id)):
            corrections.append({"type": "table_style_normalized", "table": path})

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is not None:
        tbl_pr.remove(tbl_ind)
        corrections.append({"type": "table_indent_removed", "table": path})

    tbl_w = _get_or_add(tbl_pr, "w:tblW")
    width_changed = _set_attr(tbl_w, "type", "auto")
    width_changed = _set_attr(tbl_w, "w", "0") or width_changed
    if width_changed:
        corrections.append({"type": "table_width_normalized", "table": path})

    tbl_layout = _get_or_add(tbl_pr, "w:tblLayout")
    if _set_attr(tbl_layout, "type", "autofit"):
        corrections.append({"type": "table_layout_normalized", "table": path})

    tbl_cell_mar = _get_or_add(tbl_pr, "w:tblCellMar")
    margin_changed = False
    expected_margin_tags = {qn(f"w:{name}") for name in TABLE_DEFAULT_CELL_MARGINS}
    for child in list(tbl_cell_mar):
        if child.tag not in expected_margin_tags:
            tbl_cell_mar.remove(child)
            margin_changed = True
    for name, width in TABLE_DEFAULT_CELL_MARGINS.items():
        margin = _get_or_add(tbl_cell_mar, f"w:{name}")
        margin_changed = _set_attr(margin, "w", str(width)) or margin_changed
        margin_changed = _set_attr(margin, "type", "dxa") or margin_changed
    if margin_changed:
        corrections.append({"type": "table_default_cell_margins_normalized", "table": path})

    if policy.table_look:
        tbl_look = _get_or_add(tbl_pr, "w:tblLook")
        changed = False
        for name, value in policy.table_look.items():
            changed = _set_attr(tbl_look, name, value) or changed
        if changed:
            corrections.append({"type": "table_look_normalized", "table": path})
    else:
        tbl_look = tbl_pr.find(qn("w:tblLook"))
        if tbl_look is not None:
            tbl_pr.remove(tbl_look)
            corrections.append({"type": "table_look_removed", "table": path})

    borders = _get_or_add(tbl_pr, "w:tblBorders")
    changed = False
    for name in TABLE_BORDER_NAMES:
        edge = _get_or_add(borders, f"w:{name}")
        for attr, value in (("val", "single"), ("sz", "4"), ("space", "0"), ("color", "000000")):
            changed = _set_attr(edge, attr, value) or changed
    if changed:
        corrections.append({"type": "table_borders_normalized", "table": path})
    return corrections


def _normalize_row(row, policy: TableFormatPolicy, path: str) -> list[dict]:
    corrections: list[dict] = []
    target_rule = (
        WD_ROW_HEIGHT_RULE.AT_LEAST
        if policy.row_height_rule == "at-least"
        else WD_ROW_HEIGHT_RULE.EXACTLY
    )
    if row.height_rule != target_rule:
        row.height_rule = target_rule
        corrections.append({"type": "table_row_height_rule_normalized", "row": path})
    if row.height is None or abs(row.height.cm - policy.row_height_cm) > 0.02:
        row.height = Cm(policy.row_height_cm)
        corrections.append({"type": "table_row_height_normalized", "row": path})
    return corrections


def _normalize_cell(cell, policy: TableFormatPolicy, path: str) -> list[dict]:
    corrections: list[dict] = []
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is not None:
        tc_pr.remove(tc_mar)
        corrections.append({"type": "table_cell_margin_override_removed", "cell": path})

    if cell.vertical_alignment != WD_CELL_VERTICAL_ALIGNMENT.CENTER:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        corrections.append({"type": "table_cell_vertical_alignment_normalized", "cell": path})

    for paragraph_index, paragraph in enumerate(cell.paragraphs, 1):
        paragraph_path = f"{path}.p{paragraph_index}"
        corrections.extend(_normalize_paragraph(paragraph, policy, paragraph_path))
    return corrections


def _normalize_paragraph(paragraph, policy: TableFormatPolicy, path: str) -> list[dict]:
    corrections: list[dict] = []
    current_style = paragraph.style.name if paragraph.style is not None else ""
    if current_style != policy.paragraph_style:
        paragraph.style = policy.paragraph_style
        corrections.append({"type": "table_paragraph_style_normalized", "paragraph": path})

    p_pr = paragraph._p.get_or_add_pPr()
    removed = []
    for child in list(p_pr):
        if child.tag not in PARAGRAPH_ALLOWED_TAGS:
            removed.append(child.tag.rsplit("}", 1)[-1])
            p_pr.remove(child)
    if removed:
        corrections.append({
            "type": "table_paragraph_direct_format_removed",
            "paragraph": path,
            "properties": removed,
        })

    if paragraph.alignment != policy.alignment:
        paragraph.alignment = policy.alignment
        corrections.append({"type": "table_paragraph_alignment_normalized", "paragraph": path})

    run_changes = 0
    for run_element in paragraph._p.iter(qn("w:r")):
        r_pr = run_element.find(qn("w:rPr"))
        if r_pr is None:
            continue
        for child in list(r_pr):
            if child.tag not in RUN_ALLOWED_TAGS:
                r_pr.remove(child)
                run_changes += 1
    if run_changes:
        corrections.append({
            "type": "table_run_direct_format_removed",
            "paragraph": path,
            "properties_removed": run_changes,
        })
    return corrections


def normalize_table(
    table,
    profile: dict | None,
    row_height_cm: float,
    row_height_rule: str,
    *,
    role: str = "data",
    path: str = "1",
) -> list[dict]:
    if table_role(role) == "layout":
        return []
    policy = build_table_policy(
        profile,
        row_height_cm,
        row_height_rule,
        role=role,
    )
    corrections = _normalize_table_properties(table, policy, path)
    seen_cells: set[object] = set()
    for row_index, row in enumerate(table.rows, 1):
        corrections.extend(_normalize_row(row, policy, f"{path}.r{row_index}"))
        for cell_index, cell in enumerate(row.cells, 1):
            cell_key = cell._tc
            if cell_key in seen_cells:
                continue
            seen_cells.add(cell_key)
            cell_path = f"{path}.r{row_index}.c{cell_index}"
            corrections.extend(_normalize_cell(cell, policy, cell_path))
            for nested_index, nested in enumerate(cell.tables, 1):
                corrections.extend(normalize_table(
                    nested,
                    profile,
                    row_height_cm,
                    row_height_rule,
                    role=role,
                    path=f"{cell_path}.t{nested_index}",
                ))
    return corrections


def normalize_document_tables(
    doc,
    profile: dict | None,
    row_height_cm: float,
    row_height_rule: str,
    *,
    table_roles: list[str] | None = None,
) -> list[dict]:
    corrections: list[dict] = []
    roles = table_roles or []
    for table_index, table in enumerate(doc.tables, 1):
        role = roles[table_index - 1] if table_index <= len(roles) else "data"
        corrections.extend(normalize_table(
            table,
            profile,
            row_height_cm,
            row_height_rule,
            role=role,
            path=str(table_index),
        ))
    return corrections


def _audit_table(table, policy: TableFormatPolicy, path: str, audit: dict) -> None:
    audit["table_count"] += 1
    if policy.role == "layout":
        audit["layout_table_count"] += 1
        return
    tbl_pr = table._tbl.tblPr
    tbl_style = tbl_pr.find(qn("w:tblStyle"))
    actual_style = tbl_style.get(qn("w:val")) if tbl_style is not None else None
    if policy.table_style_id and actual_style != str(policy.table_style_id):
        audit["invalid_table_styles"].append({
            "table": path,
            "expected": str(policy.table_style_id),
            "actual": actual_style,
        })

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None or tbl_layout.get(qn("w:type")) != "autofit":
        audit["table_layout_issues"].append({"table": path})

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if (
        tbl_w is None
        or tbl_w.get(qn("w:type")) != "auto"
        or tbl_w.get(qn("w:w")) != "0"
    ):
        audit["table_width_issues"].append({"table": path})

    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    default_margin_issue = False
    for name, width in TABLE_DEFAULT_CELL_MARGINS.items():
        margin = tbl_cell_mar.find(qn(f"w:{name}")) if tbl_cell_mar is not None else None
        if (
            margin is None
            or margin.get(qn("w:w")) != str(width)
            or margin.get(qn("w:type")) != "dxa"
        ):
            default_margin_issue = True
    if default_margin_issue:
        audit["table_default_cell_margin_issues"].append({"table": path})

    seen_cells: set[object] = set()
    expected_rule = (
        WD_ROW_HEIGHT_RULE.AT_LEAST
        if policy.row_height_rule == "at-least"
        else WD_ROW_HEIGHT_RULE.EXACTLY
    )
    for row_index, row in enumerate(table.rows, 1):
        row_path = f"{path}.r{row_index}"
        audit["row_count"] += 1
        if row.height_rule != expected_rule:
            audit["row_height_rule_issues"].append({"row": row_path})
        if row.height is None or abs(row.height.cm - policy.row_height_cm) > 0.02:
            audit["row_height_issues"].append({"row": row_path})

        for cell_index, cell in enumerate(row.cells, 1):
            cell_key = cell._tc
            if cell_key in seen_cells:
                continue
            seen_cells.add(cell_key)
            cell_path = f"{row_path}.c{cell_index}"
            audit["cell_count"] += 1
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is not None:
                audit["cell_margin_issues"].append({"cell": cell_path})
            if cell.vertical_alignment != WD_CELL_VERTICAL_ALIGNMENT.CENTER:
                audit["cell_vertical_alignment_issues"].append({"cell": cell_path})

            for paragraph_index, paragraph in enumerate(cell.paragraphs, 1):
                paragraph_path = f"{cell_path}.p{paragraph_index}"
                audit["paragraph_count"] += 1
                style_name = paragraph.style.name if paragraph.style is not None else ""
                if style_name != policy.paragraph_style:
                    audit["paragraph_style_issues"].append({
                        "paragraph": paragraph_path,
                        "expected": policy.paragraph_style,
                        "actual": style_name,
                    })
                if paragraph.alignment != policy.alignment:
                    audit["paragraph_alignment_issues"].append({"paragraph": paragraph_path})
                p_pr = paragraph._p.get_or_add_pPr()
                forbidden = [
                    child.tag.rsplit("}", 1)[-1]
                    for child in p_pr
                    if child.tag not in PARAGRAPH_ALLOWED_TAGS
                ]
                if forbidden:
                    audit["paragraph_direct_format_issues"].append({
                        "paragraph": paragraph_path,
                        "properties": forbidden,
                    })
                run_forbidden = []
                for run_element in paragraph._p.iter(qn("w:r")):
                    r_pr = run_element.find(qn("w:rPr"))
                    if r_pr is None:
                        continue
                    run_forbidden.extend(
                        child.tag.rsplit("}", 1)[-1]
                        for child in r_pr
                        if child.tag not in RUN_ALLOWED_TAGS
                    )
                if run_forbidden:
                    audit["run_direct_format_issues"].append({
                        "paragraph": paragraph_path,
                        "properties": run_forbidden,
                    })
            for nested_index, nested in enumerate(cell.tables, 1):
                _audit_table(nested, policy, f"{cell_path}.t{nested_index}", audit)


def audit_document_tables(
    doc,
    profile: dict | None,
    row_height_cm: float,
    row_height_rule: str,
    *,
    table_roles: list[str] | None = None,
) -> dict:
    audit = {
        "table_count": 0,
        "layout_table_count": 0,
        "row_count": 0,
        "cell_count": 0,
        "paragraph_count": 0,
        "invalid_table_styles": [],
        "table_layout_issues": [],
        "table_width_issues": [],
        "table_default_cell_margin_issues": [],
        "row_height_rule_issues": [],
        "row_height_issues": [],
        "cell_margin_issues": [],
        "cell_vertical_alignment_issues": [],
        "paragraph_style_issues": [],
        "paragraph_alignment_issues": [],
        "paragraph_direct_format_issues": [],
        "run_direct_format_issues": [],
        "table_body_style_spacing_issues": [],
    }
    roles = table_roles or []
    if any((roles[index] if index < len(roles) else "data") != "layout" for index in range(len(doc.tables))):
        resolved = (profile or {}).get("resolved_styles", {})
        style_name = resolved.get("table_body", "表正文")
        try:
            style = doc.styles[style_name]
            spacing = style._element.find(qn("w:pPr") + "/" + qn("w:spacing"))
        except (KeyError, TypeError):
            spacing = None
        if (
            spacing is None
            or spacing.get(qn("w:line")) != "0"
            or spacing.get(qn("w:lineRule")) != "atLeast"
        ):
            audit["table_body_style_spacing_issues"].append({
                "style": style_name,
                "expected_line": "0",
                "expected_line_rule": "atLeast",
            })
    for table_index, table in enumerate(doc.tables, 1):
        role = roles[table_index - 1] if table_index <= len(roles) else "data"
        policy = build_table_policy(
            profile,
            row_height_cm,
            row_height_rule,
            role=role,
        )
        _audit_table(table, policy, str(table_index), audit)
    issue_keys = [key for key in audit if key.endswith("issues") or key == "invalid_table_styles"]
    audit["passed"] = not any(audit[key] for key in issue_keys)
    return audit
