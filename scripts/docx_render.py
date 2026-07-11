from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Callable, Iterable

from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.table import Table

from text_utils import (
    heading_level_from_style,
    heading_level_from_text,
    heading_style_for_level,
    is_front_matter_text,
    looks_like_visual_heading,
    looks_like_list_item,
    list_kind_for_text,
    list_style_for_text,
    normalize_heading_style_level,
    paragraph_has_graphics,
    paragraph_num_info,
    resolved_heading_level,
    source_numbering_heading_level,
    looks_like_code_sample_table, looks_like_api_example_table,
    source_heading_level_shift,
    strip_list_marker,
)
from word_model_renderer import (
    style_from_profile,
    list_style_for_model,
    _new_list_num,
    _set_list_numbering,
)
from text_utils import set_template_table_properties as _set_template_table_properties
from docx_pipeline import infer_docx_role


def _body_insert_before_section(doc, new_el, *, qn_fn: Callable[[str], str]) -> None:
    body = doc.element.body
    sect_pr = body[-1] if len(body) and body[-1].tag == qn_fn("w:sectPr") else None
    if sect_pr is not None:
        body.insert(len(body) - 1, new_el)
    else:
        body.append(new_el)


def _append_table_clone_new(doc, table, *, qn_fn=qn, table_class=Table):
    new_tbl = deepcopy(table._tbl)
    _body_insert_before_section(doc, new_tbl, qn_fn=qn_fn)
    return table_class(new_tbl, doc)


def _clone_related_media_rels_new(src_paragraph, dst_doc, new_el, *, qn_fn=qn, image_reltype, media_map=None) -> int:
    """Re‑wire cloned XML so every image rId points at the matching part in *dst_doc*.

    Scans *all* descendant elements for r:embed / r:id / r:pict attributes
    (a:blip, v:imagedata, o:OLEObject, w:pict …) and updates them to point
    at the corresponding ImagePart that was created by `_precopy_source_media`.

    The generic scan avoids hard‑coding element names and therefore works
    with every embedding mechanism that Office documents use.
    """
    from text_utils import _iter_image_rids

    count = 0
    src_part = src_paragraph.part
    dst_part = dst_doc.part

    for attr_name, rid, rel in _iter_image_rids(new_el, src_part):
        src_pname = str(rel.target_part.partname)
        new_pname = media_map.get(src_pname, src_pname) if media_map else src_pname

        # Find the rId that already points to new_pname in dst_part …
        existing_rid = None
        for dst_rel in dst_part.rels.values():
            if str(dst_rel.target_part.partname) == new_pname:
                existing_rid = dst_rel.rId
                break

        if existing_rid is None:
            # … or create one.
            for part in dst_doc.part.package.iter_parts():
                if str(part.partname) == new_pname:
                    existing_rid = dst_part.rels.add_relationship(
                        rel.reltype, part, part.partname,
                    )
                    break

        if existing_rid is not None:
            # Patch the XML element's attribute in‑place (XPath returns
            # live references).
            for descendant in new_el.iter():
                if descendant.get(attr_name) == rid:
                    descendant.set(attr_name, existing_rid)
                    count += 1
                    break

    return count

def _make_paragraph_clone(doc, paragraph: Paragraph) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    _body_insert_before_section(doc, new_p, qn_fn=qn)
    return Paragraph(new_p, doc)


def _append_paragraph_clone_new(doc, paragraph, *, qn_fn=qn, paragraph_class=Paragraph,
                                  paragraph_has_graphics_fn=None, normalize_graphics_paragraph_fn=None, image_reltype=None, media_map=None):
    new_p = _make_paragraph_clone(doc, paragraph)
    media_count = 0
    if image_reltype:
        media_count = _clone_related_media_rels_new(paragraph, doc, new_p._element, qn_fn=qn_fn, image_reltype=image_reltype, media_map=media_map)
    if normalize_graphics_paragraph_fn and (paragraph_has_graphics_fn or paragraph_has_graphics)(paragraph):
        try:
            normalize_graphics_paragraph_fn(new_p)
        except Exception:
            pass
    return new_p, media_count


def _normalize_graphics_paragraph_new(paragraph, *, center_alignment=None, pt=None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn as _qn2
    _center = center_alignment if center_alignment is not None else WD_ALIGN_PARAGRAPH.CENTER
    paragraph.alignment = _center

    # Image paragraphs use "No Spacing" style to avoid inheriting
    # Normal's firstLine=640 indent.  The style is natively indent-free.
    # Also strip any direct-formatting indent/spacing that may have been
    # deep‑copied from the source paragraph.
    from docx.oxml import OxmlElement as _Oxml
    ppr = paragraph._element.find(_qn2('w:pPr'))
    if ppr is None:
        ppr = _Oxml('w:pPr')
        paragraph._element.insert(0, ppr)
    # Remove indent, spacing, justification from clone — pStyle + jc handle it.
    for tag in ('w:ind', 'w:spacing', 'w:jc', 'w:numPr', 'w:pBdr', 'w:shd',
                'w:textDirection', 'w:outlineLvl', 'w:keepNext', 'w:keepLines',
                'w:pageBreakBefore', 'w:widowControl', 'w:autoSpaceDE',
                'w:autoSpaceDN', 'w:contextualSpacing', 'w:snapToGrid'):
        for el in ppr.findall(_qn2(tag)):
            ppr.remove(el)
    try:
        paragraph.style = paragraph._parent.styles["No Spacing"]
    except Exception:
        pass
    # Re‑assert centering (python-docx .alignment property does this reliably)
    paragraph.alignment = _center

    # Remove inline image margins (distL/distR on wp:inline or wp:anchor)
    # which cause visual left/right offsets independent of paragraph indents.
    for drawing in paragraph._element.findall('.//' + _qn2('w:drawing')):
        for wp_el in drawing:
            tag = wp_el.tag.split('}')[-1]
            if tag in ('inline', 'anchor'):
                for attr in ('distL', 'distR', 'distT', 'distB'):
                    if wp_el.get(attr) is not None:
                        wp_el.set(attr, '0')


def _element_has_graphics_new(element, *, qn_fn=qn):
    try:
        drawings = element.findall(".//" + qn_fn("w:drawing"))
        if drawings:
            return True
        inline = element.findall(".//" + qn_fn("w:inline"))
        if inline:
            return True
        blip = element.findall(".//" + "{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
        if blip:
            return True
    except Exception:
        pass
    return False


def _remove_non_graphic_children(new_el, *, qn_fn=qn, graphic_tags: Iterable[str]) -> None:
    for child in list(new_el):
        if child.tag == qn_fn("w:r"):
            drawing = child.find(qn_fn("w:drawing")) or child.find(qn_fn("w:inline"))
            if drawing is None:
                new_el.remove(child)
            continue
        if child.tag in graphic_tags:
            continue
        _remove_non_graphic_children(child, qn_fn=qn_fn, graphic_tags=graphic_tags)


def _append_graphics_only_paragraph_clone_new(doc, paragraph, *, qn_fn=qn, paragraph_class=Paragraph,
                                                normalize_graphics_paragraph_fn=None, image_reltype=None):
    new_p = deepcopy(paragraph._p)
    _remove_non_graphic_children(new_p, qn_fn=qn_fn, graphic_tags={qn_fn("w:pPr"), qn_fn("w:rPr")})
    _body_insert_before_section(doc, new_p, qn_fn=qn_fn)
    new_para = Paragraph(new_p, doc)
    media_count = 0
    if image_reltype:
        media_count = _clone_related_media_rels_new(paragraph, doc, new_p, qn_fn=qn_fn, image_reltype=image_reltype)
    if normalize_graphics_paragraph_fn:
        try:
            normalize_graphics_paragraph_fn(new_para)
        except Exception:
            pass
    return new_para, media_count




def _insert_seq_caption(doc, caption_type: str, caption_text: str, *, template_profile=None):
    """Insert a paragraph with a Word SEQ field for auto-numbering.
    
    caption_type: 'table' or 'figure'
    caption_text: the caption text without number prefix
    Uses the template's caption style.
    """
    from docx.oxml.ns import qn as _qn
    from lxml import etree as _etree
    
    style_name = style_from_profile(template_profile, "caption", "Caption")
    
    p = doc.add_paragraph()
    if style_name:
        try:
            p.style = doc.styles[style_name]
        except Exception:
            pass
    
    seq_name = "Table" if caption_type == "table" else "Figure"
    prefix = "表 " if caption_type == "table" else "图 "
    
    # Build SEQ field: <w:fldSimple w:instr="SEQ Table \* ARABIC">
    fld = _etree.SubElement(p._element, _qn('w:fldSimple'))
    fld.set(_qn('w:instr'), f'SEQ {seq_name} \\* ARABIC')
    
    r = _etree.SubElement(fld, _qn('w:r'))
    t = _etree.SubElement(r, _qn('w:t'))
    t.text = '1'  # placeholder, Word fills this in on update
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    
    # Add separator and caption text
    r2 = _etree.SubElement(p._element, _qn('w:r'))
    t2 = _etree.SubElement(r2, _qn('w:t'))
    t2.text = ' ' + caption_text
    t2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    
    return p


def render_docx_direct(
    src_doc,
    dst_doc,
    report: dict,
    row_height_cm: float,
    row_height_rule: str,
    numbering_ids: dict,
    *,
    template_profile: dict | None = None,
    strict_normalize: bool = True,
    role_overrides: dict | None = None,
    heading_level_overrides: dict[int, int] | None = None,
    table_type_overrides: dict[int, str] | None = None,
) -> None:
    """
    Render a DOCX source directly by iterating blocks, using template styles only.
    Images and table cloning are preserved; paragraphs use template styles.
    """
    heading_shift = source_heading_level_shift(src_doc)
    
    # Pre-copy all media from source to destination with unique names
    _media_map = _precopy_source_media(src_doc, dst_doc)
    if _media_map:
        report.setdefault("media_relationships_preserved", 
            report.get("media_relationships_preserved", 0) + len(_media_map))


    if heading_shift:
        report.setdefault("content_warnings", []).append({
            "type": "heading_level_shift",
            "message": "源文档标题样式从非一级开始，已整体上移标题层级。",
            "shift": heading_shift,
        })

    active_list_nums: dict[int, int] = {}
    seen_content = False
    structural_started = False

    def _iter_blocks(d):
        for child in d.element.body.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, d)
            elif child.tag == qn("w:tbl"):
                yield Table(child, d)

    def _add_styled(text, style, role):
        s = style
        if s is None:
            if role == "note":
                s = style_from_profile(template_profile, "note", "3.1注-无编号注")
            elif role == "numbered_note":
                s = style_from_profile(template_profile, "numbered_note", "3.2注-有编号注")
            elif role == "formula":
                s = style_from_profile(template_profile, "formula", "Normal")
            elif role == "title":
                s = style_from_profile(template_profile, "title", "文档标题")
            elif role == "caption":
                s = style_from_profile(template_profile, "caption", "Caption")
            elif role == "body":
                s = style_from_profile(template_profile, "body", "Normal")
            else:
                s = style_from_profile(template_profile, "body", "Normal")
        try:
            dst_doc.add_paragraph(text, style=s)
        except Exception:
            dst_doc.add_paragraph(text)

    last_was_caption = False
    _para_idx = 0
    for block in _iter_blocks(src_doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()

            if paragraph_has_graphics(block):
                if text:
                    inferred_text, style, role = infer_docx_role(block, strict_normalize, report)
                    if role_overrides is not None and _para_idx in role_overrides:
                        role = role_overrides[_para_idx]
                    source_style = block.style.name if block.style is not None else ""
                    if role == "heading" and heading_level_from_style(source_style) is not None:
                        style = normalize_heading_style_level(style, heading_shift)
                    # Apply heading level override from Phase B enhancement
                    if (role == "heading"
                            and heading_level_overrides is not None
                            and _para_idx in heading_level_overrides):
                        override_level = heading_level_overrides[_para_idx]
                        style = heading_style_for_level(override_level)
                    if role == "caption":
                        from text_utils import caption_parts
                        ct, label, _, cap_text = caption_parts(inferred_text)
                        if ct == "unknown":
                            if "图" in inferred_text[:4]:
                                ct, cap_text = "figure", inferred_text
                            else:
                                ct, cap_text = "table", inferred_text
                        _insert_seq_caption(dst_doc, ct, cap_text, template_profile=template_profile)
                    else:
                        _handle_inferred(dst_doc, inferred_text, style, role, report, numbering_ids, active_list_nums, template_profile=template_profile)
                    _, media_count = _append_graphics_only_paragraph_clone_new(dst_doc, block, qn_fn=qn, image_reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image", media_map=_media_map, normalize_graphics_paragraph_fn=_normalize_graphics_paragraph_new)
                    split_record = {"text": inferred_text, "role": role}
                    if role == "heading":
                        split_record["level"] = resolved_heading_level(style, None, inferred_text)
                    report.setdefault("semantic_object_splits", []).append(split_record)
                    report.setdefault("mixed_text_graphic_paragraphs_split", []).append(split_record)
                    report["graphic_paragraphs_preserved"] = report.get("graphic_paragraphs_preserved", 0) + 1
                    report["media_relationships_preserved"] = report.get("media_relationships_preserved", 0) + media_count
                    active_list_nums = {}
                    last_was_caption = False
                    seen_content = True
                    _para_idx += 1
                    continue

                _, media_count = _append_paragraph_clone_new(dst_doc, block, qn_fn=qn, image_reltype="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image", media_map=_media_map, paragraph_has_graphics_fn=paragraph_has_graphics, normalize_graphics_paragraph_fn=_normalize_graphics_paragraph_new)
                report["graphic_paragraphs_preserved"] = report.get("graphic_paragraphs_preserved", 0) + 1
                report["media_relationships_preserved"] = report.get("media_relationships_preserved", 0) + media_count
                active_list_nums = {}
                last_was_caption = False
                seen_content = True
                _para_idx += 1
                continue

            if not text:
                continue

            if strict_normalize and not structural_started and seen_content and is_front_matter_text(block, text):
                # Source-style headings (Heading 1-9) should never be treated as
                # front matter — they define the document structure.
                if block.style and block.style.name and block.style.name.startswith("Heading"):
                    pass  # fall through to normal heading handling
                else:
                    _add_styled(text, None, "body")
                    active_list_nums = {}
                    last_was_caption = False
                    continue

            if strict_normalize and not seen_content and looks_like_visual_heading(block):
                report.setdefault("suspect_visual_headings", []).append({"text": text, "assigned_level": "title"})
                _add_styled(text, style_from_profile(template_profile, "title", "文档标题"), "title")
                last_was_caption = False
                seen_content = True
                continue

            inferred_text, style, role = infer_docx_role(block, strict_normalize, report)
            if role_overrides is not None and _para_idx in role_overrides:
                role = role_overrides[_para_idx]
            source_style = block.style.name if block.style is not None else ""
            if role == "heading" and heading_level_from_style(source_style) is not None:
                style = normalize_heading_style_level(style, heading_shift)

            # Apply heading level override from Phase B enhancement
            if (role == "heading"
                    and heading_level_overrides is not None
                    and _para_idx in heading_level_overrides):
                override_level = heading_level_overrides[_para_idx]
                style = heading_style_for_level(override_level)

            if role == "caption":
                from text_utils import caption_parts
                ct, label, _, cap_text = caption_parts(inferred_text)
                if ct == "unknown":
                    # Text-based caption — detect from content
                    if "图" in inferred_text[:4]:
                        ct, cap_text = "figure", inferred_text
                    else:
                        ct, cap_text = "table", inferred_text
                _insert_seq_caption(dst_doc, ct, cap_text, template_profile=template_profile)
            else:
                _handle_inferred(dst_doc, inferred_text, style, role, report, numbering_ids, active_list_nums, template_profile=template_profile)

            last_was_caption = (role == "caption")
            if role == "heading":
                structural_started = True
                active_list_nums = {}
            seen_content = True
            _para_idx += 1

        else:
            is_api_example = looks_like_api_example_table(block)
            # Apply table type override from Phase C enhancement
            if table_type_overrides is not None and _para_idx in table_type_overrides:
                override_type = table_type_overrides[_para_idx]
                is_api_example = (override_type == "code_sample")
            # Auto-insert table caption if the preceding paragraph wasn't one
            if not last_was_caption:
                _insert_seq_caption(dst_doc, 'table', ' ', template_profile=template_profile)
                report.setdefault('captions_auto_generated', 0)
                report['captions_auto_generated'] = report['captions_auto_generated'] + 1
            
            new_table = append_table_clone(dst_doc, block)
            _table_body_style = style_from_profile(template_profile, "table_body", "表正文")
            _set_template_table_properties(new_table, row_height_cm, row_height_rule, table_body_style=_table_body_style)
            report["tables_processed"] = report.get("tables_processed", 0) + 1
            last_was_caption = False
            active_list_nums = {}
            seen_content = True
            _para_idx += 1


def _handle_inferred(doc, text: str, style: str | None, role: str, report: dict, numbering_ids: dict, active_list_nums: dict, *, template_profile: dict | None = None) -> dict:
    """Apply inferred paragraph to doc using template styles."""
    if role == "heading":
        heading_level = resolved_heading_level(style, None, text)
        if heading_level <= 0 or role == "title":
            from text_utils import strip_heading_marker as _shm
            try:
                doc.add_paragraph(_shm(text), style=style_from_profile(None, "title", "文档标题"))
            except Exception:
                doc.add_paragraph(text)
            return {}

        resolved_style = style_from_profile(template_profile, f"heading_{heading_level}", style)
        from text_utils import strip_heading_marker as _shm
        clean_text = _shm(text)
        try:
            doc.add_paragraph(clean_text, style=resolved_style)
        except Exception:
            doc.add_paragraph(clean_text)
        report.setdefault("automatic_numbers", []).append(
            {"type": "heading", "text": text, "level": heading_level, "source": "docx-render"}
        )
        return {}

    if role in ("list", "list_item"):
        kind = list_kind_for_text(text)
        from text_utils import strip_list_marker as _slm
        clean_text = _slm(text)
        style_name = list_style_for_text(text)
        try:
            doc.add_paragraph(clean_text, style=style_name)
        except Exception:
            doc.add_paragraph(clean_text)

        if kind in {"letter", "decimal"}:
            list_key = (0, "letter")  # every section starts from primary (letter) list
            if list_key not in active_list_nums:
                aid = numbering_ids.get("list_letter_abstract")
                if aid is not None:
                    active_list_nums[list_key] = _new_list_num(doc, aid)
            nid = active_list_nums.get(list_key)
            if nid is not None:
                # Write numPr directly via XML (python-docx may clear indirect set)
                pPr = doc.paragraphs[-1]._element.find(qn("w:pPr"))
                if pPr is None:
                    pPr = OxmlElement("w:pPr")
                    doc.paragraphs[-1]._element.insert(0, pPr)
                num_pr = pPr.find(qn("w:numPr"))
                if num_pr is None:
                    num_pr = OxmlElement("w:numPr")
                    pPr.append(num_pr)
                else:
                    for c in list(num_pr):
                        num_pr.remove(c)
                nid_el = OxmlElement("w:numId")
                nid_el.set(qn("w:val"), str(nid))
                num_pr.append(nid_el)
                ilvl_el = OxmlElement("w:ilvl")
                ilvl_el.set(qn("w:val"), "0")
                num_pr.append(ilvl_el)
            report.setdefault("automatic_numbers", []).append(
                {"type": "list", "text": clean_text, "source": "docx-render"}
            )
        return active_list_nums

    # Fallback: use role-based template style
    if role == "note":
        s = style_from_profile(None, "note", "3.1注-无编号注")
    elif role == "numbered_note":
        s = style_from_profile(None, "numbered_note", "3.2注-有编号注")
    elif role == "caption":
        s = style_from_profile(None, "caption", "Caption")
    else:
        s = style_from_profile(None, "body", "Normal")
    try:
        doc.add_paragraph(text, style=s)
    except Exception:
        doc.add_paragraph(text)
    return {}




# ---- Public API (old parameter names for format_document.py compatibility) ----

def append_table_clone(doc, table, *, qn=None, table_class=None):
    from docx.oxml.ns import qn as _qn
    from docx.table import Table as _Table
    return _append_table_clone_new(doc, table, qn_fn=qn or _qn, table_class=table_class or _Table)


def append_table_clone_with_dependencies(doc, table, *, qn=None, table_class=None):
    return append_table_clone(doc, table, qn=qn, table_class=table_class)


def clone_related_media_rels(src_paragraph, dst_doc, new_el, *, qn=None, image_reltype):
    from docx.oxml.ns import qn as _qn
    return _clone_related_media_rels_new(src_paragraph, dst_doc, new_el, qn_fn=qn or _qn, image_reltype=image_reltype)


def clone_related_media_rels_with_dependencies(src_paragraph, dst_doc, new_el, *, qn=None, image_reltype=None):
    return clone_related_media_rels(src_paragraph, dst_doc, new_el, qn=qn, image_reltype=image_reltype)


def append_paragraph_clone(doc, paragraph, *, qn=None, paragraph_class=None,
                            paragraph_has_graphics=None, normalize_graphics_paragraph=None, image_reltype=None):
    from docx.oxml.ns import qn as _qn
    from docx.text.paragraph import Paragraph as _Paragraph
    return _append_paragraph_clone_new(doc, paragraph, qn_fn=qn or _qn, paragraph_class=paragraph_class or _Paragraph,
                                        paragraph_has_graphics_fn=paragraph_has_graphics,
                                        normalize_graphics_paragraph_fn=normalize_graphics_paragraph,
                                        image_reltype=image_reltype)


def append_paragraph_clone_with_dependencies(doc, paragraph, *, qn=None, paragraph_class=None,
                                              paragraph_has_graphics=None, normalize_graphics_paragraph=None, image_reltype=None):
    return append_paragraph_clone(doc, paragraph, qn=qn, paragraph_class=paragraph_class,
                                   paragraph_has_graphics=paragraph_has_graphics,
                                   normalize_graphics_paragraph=normalize_graphics_paragraph,
                                   image_reltype=image_reltype)


def normalize_graphics_paragraph(paragraph, *, center_alignment=None, pt=None):
    from docx.shared import Pt as _Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _center = center_alignment if center_alignment is not None else WD_ALIGN_PARAGRAPH.CENTER
    _pt = pt if pt is not None else _Pt
    _normalize_graphics_paragraph_new(paragraph, center_alignment=_center, pt=_pt)


def normalize_graphics_paragraph_with_dependencies(paragraph, *, center_alignment=None, pt=None):
    return normalize_graphics_paragraph(paragraph, center_alignment=center_alignment, pt=pt)


def element_has_graphics(element, *, qn=None):
    from docx.oxml.ns import qn as _qn
    return _element_has_graphics_new(element, qn_fn=qn or _qn)


def element_has_graphics_with_dependencies(element, *, qn=None):
    return element_has_graphics(element, qn=qn)


def append_graphics_only_paragraph_clone(doc, paragraph, *, qn=None, paragraph_class=None,
                                          normalize_graphics_paragraph=None, image_reltype=None):
    from docx.oxml.ns import qn as _qn
    from docx.text.paragraph import Paragraph as _Paragraph
    return _append_graphics_only_paragraph_clone_new(doc, paragraph, qn_fn=qn or _qn,
                                                      paragraph_class=paragraph_class or _Paragraph,
                                                      normalize_graphics_paragraph_fn=normalize_graphics_paragraph,
                                                      image_reltype=image_reltype)


def append_graphics_only_paragraph_clone_with_dependencies(doc, paragraph, *, qn=None, paragraph_class=None,
                                                            normalize_graphics_paragraph=None, image_reltype=None):
    return append_graphics_only_paragraph_clone(doc, paragraph, qn=qn, paragraph_class=paragraph_class,
                                                 normalize_graphics_paragraph=normalize_graphics_paragraph,
                                                 image_reltype=image_reltype)


def _precopy_source_media(src_doc, dst_doc) -> dict:
    """Copy all media from source to destination, converting to PNG when
    the source format is not PNG or JPEG.  EMF, WMF, BMP, GIF, TIFF, SVG
    and any other raster/vector format are all normalised to PNG so that
    downstream code only needs to deal with one well‑known format.

    Creates ImagePart objects directly in the destination package.
    Returns a dict mapping old partname -> new partname.
    """
    import hashlib
    from docx.opc.packuri import PackURI
    from docx.parts.image import ImagePart

    # PNG/JPEG are passed through as-is.
    # EMF/WMF are also kept in original format because Pillow cannot
    # reliably re-encode them on all platforms (the WMF loader may
    # fail at render time even when Image.open succeeds).
    _DIRECT_FORMATS = frozenset({'png', 'jpg', 'jpeg', 'emf', 'wmf'})

    media_map: dict[str, str] = {}
    seen: set[str] = set()

    for rel in src_doc.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            source_part = rel.target_part
        except Exception:
            continue
        pname = str(source_part.partname)
        if not pname.startswith('/word/media/'):
            continue
        if pname in seen:
            continue
        seen.add(pname)
        if not hasattr(source_part, 'blob'):
            continue
        try:
            source_blob = source_part.blob
        except Exception:
            continue

        ext = pname.rsplit('.', 1)[-1].lower() if '.' in pname else 'png'
        sig = hashlib.md5(source_blob).hexdigest()[:8]

        if ext in _DIRECT_FORMATS:
            if ext == 'png':
                content_type = 'image/png'
            elif ext in ('jpg', 'jpeg'):
                content_type = 'image/jpeg'
            elif ext == 'emf':
                content_type = 'image/x-emf'
            elif ext == 'wmf':
                content_type = 'image/x-wmf'
            else:
                content_type = 'image/png'
            new_partname_str = f'/word/media/img_{sig}.{ext}'
            blob = source_blob
        else:
            content_type = 'image/png'
            new_partname_str = f'/word/media/img_{sig}.png'
            blob = _convert_image_to_png(source_blob, pname)

        if blob is None:
            continue

        image_part = ImagePart(PackURI(new_partname_str), content_type, blob)
        dst_doc.part.relate_to(image_part, rel.reltype)
        media_map[pname] = new_partname_str

    return media_map



def _convert_image_to_png(blob: bytes, partname: str = '') -> bytes | None:
    """Try to decode *blob* with Pillow and return a PNG byte‑string.
    Returns *None* when the blob cannot be decoded.
    """
    import io
    try:
        from PIL import Image
    except ImportError:
        return None
    try:
        img = Image.open(io.BytesIO(blob))
    except Exception:
        return None
    out = io.BytesIO()
    try:
        if img.mode in ('RGBA', 'LA', 'P', 'PA'):
            img = img.convert('RGBA')
        else:
            img = img.convert('RGB')
        img.save(out, format='PNG')
    except Exception:
        return None
    return out.getvalue()
