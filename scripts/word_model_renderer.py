from __future__ import annotations

from typing import Callable

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from table_formatting import normalize_table

from list_style_mapping import (
    normalize_wx_list_type,
    wx_list_style_name,
    wx_numbering_abstract_key,
)

def style_from_profile(template_profile: dict | None, role: str, fallback: str) -> str:
    if template_profile:
        return template_profile.get("resolved_styles", {}).get(role, fallback)
    return fallback


def list_style_for_model(list_type: str, level: int, template_profile: dict | None = None) -> str:
    return wx_list_style_name(list_type, level, template_profile)


def _new_list_num(doc, abstract_num_id: int) -> int:
    """Clone an abstract numbering definition into a new num instance with start=1.

    Creates a proper <w:num> element with <w:lvlOverride> wrapping <w:startOverride>
    so that the list restarts from the first value.
    """
    numbering = doc.part.numbering_part.element
    max_id = 0
    for num in numbering.findall(qn("w:num")):
        try:
            nid = int(num.get(qn("w:numId")))
            if nid > max_id:
                max_id = nid
        except (ValueError, TypeError):
            pass
    new_id = max_id + 1

    num_el = OxmlElement("w:num")
    num_el.set(qn("w:numId"), str(new_id))

    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_num_id))
    num_el.append(ref)

    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num_el.append(lvl_override)

    numbering.append(num_el)
    return new_id


def _set_list_numbering(paragraph, num_id: int, ilvl: int = 0) -> None:
    """Attach existing numId to a paragraph (the style already defines the rest)."""
    try:
        p_pr = paragraph._element.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        else:
            for child in list(num_pr):
                num_pr.remove(child)
        nid_el = OxmlElement("w:numId")
        nid_el.set(qn("w:val"), str(num_id))
        num_pr.append(nid_el)
        ilvl_el = OxmlElement("w:ilvl")
        ilvl_el.set(qn("w:val"), str(ilvl))
        num_pr.append(ilvl_el)
    except Exception:
        pass


def render_document_model(
    model: dict,
    doc,
    report: dict,
    row_height_cm: float,
    row_height_rule: str,
    numbering_ids: dict,
    *,
    template_profile: dict | None = None,
) -> None:
    """
    Render a normalized AST model into a template-created document.
    This function ONLY sets paragraph styles from the template.
    No manual numbering XML, no run-level fonts, no direct formatting.
    
    The template's style definitions already carry:
      - Heading styles bound to numbering (numId=1)
      - List styles bound to numbering (numId=3, 8, etc.)
      - Note, caption, appendix styles with their own numbering
    """
    active_list_nums: dict[int, int] = {}
    heading_num_id = numbering_ids.get("heading")

    for block in model.get("document", {}).get("blocks", []):
        block_type = block.get("block_type")

        # --- Heading ---
        if block_type == "heading":
            text = block.get("text", "")
            level = int(block.get("level") or 0)
            role = block.get("role", "heading")

            if role == "title" or level <= 0:
                style = style_from_profile(template_profile, "title", "文档标题")
                try:
                    doc.add_paragraph(text, style=style)
                except Exception:
                    doc.add_paragraph(text)
                active_list_nums = {}
                continue

            style = style_from_profile(template_profile, f"heading_{level}", f"Heading {level}")
            try:
                doc.add_paragraph(text, style=style)
            except Exception:
                doc.add_paragraph(text)

            report.setdefault("automatic_numbers", []).append(
                {"type": "heading", "text": text, "level": level, "source": "model"}
            )
            active_list_nums = {}

        # --- List Item ---
        elif block_type == "list_item":
            text = block.get("text", "")
            level = int(block.get("level") or 0)
            list_type = normalize_wx_list_type(
                block.get("list_type", "lower_letter_paren"), level,
            )
            restart = block.get("restart", False)
            style_name = list_style_for_model(list_type, level, template_profile)

            try:
                doc.add_paragraph(text, style=style_name)
            except Exception:
                doc.add_paragraph(text)

            # Only numbered lists need manual numId management for restart.
            # Key by (level, list_type) so that letter‑style and decimal‑style
            # lists maintain independent numbering within the same section.
            abstract_key = wx_numbering_abstract_key(list_type, level)
            if abstract_key is not None:
                list_key = (level, list_type)
                if restart or list_key not in active_list_nums:
                    aid = numbering_ids.get(abstract_key)
                    if aid is not None:
                        active_list_nums[list_key] = _new_list_num(doc, aid)
                nid = active_list_nums.get(list_key)
                if nid is not None:
                    _set_list_numbering(doc.paragraphs[-1], nid, 0)

                report.setdefault("automatic_numbers", []).append(
                    {"type": "list", "text": text, "source": "model"}
                )

        # --- Table ---
        elif block_type == "table":
            from text_utils import looks_like_code_sample_table

            rows_data = block.get("rows", [])
            if rows_data:
                col_count = max(len(r) for r in rows_data)
                table = doc.add_table(rows=len(rows_data), cols=col_count)
                for ri, row_data in enumerate(rows_data):
                    for ci in range(col_count):
                        text = row_data[ci].get("text", "") if ci < len(row_data) else ""
                        table.rows[ri].cells[ci].text = text
                normalize_table(
                    table,
                    template_profile,
                    row_height_cm,
                    row_height_rule,
                    role=block.get("table_type", "data"),
                    path=str(len(doc.tables)),
                )
            active_list_nums = {}

        # --- Appendix ---
        elif block_type == "appendix":
            text = block.get("title", "")
            style = style_from_profile(template_profile, "appendix_title", "附录标题")
            try:
                doc.add_paragraph(text, style=style)
            except Exception:
                doc.add_paragraph(text)
            active_list_nums = {}

        # --- Caption ---
        elif block_type == "caption":
            caption_type = block.get("caption_type", "table")
            label = "图" if caption_type == "figure" else "表"
            caption_text = block.get("text", "")
            text = f"{label}  {caption_text}" if caption_text else f"{label}"
            style = style_from_profile(template_profile, "caption", "Caption")
            try:
                doc.add_paragraph(text, style=style)
            except Exception:
                doc.add_paragraph(text)
            active_list_nums = {}

        # --- Body / Note / Formula ---
        else:
            source_role = block.get("role") or block.get("source", {}).get("role")
            if source_role == "note":
                style = style_from_profile(template_profile, "note", "3.1注-无编号注")
            elif source_role == "numbered_note":
                style = style_from_profile(template_profile, "numbered_note", "3.2注-有编号注")
            elif source_role == "formula":
                style = style_from_profile(template_profile, "formula", "Normal")
            else:
                style = style_from_profile(template_profile, "body", "Normal")
            try:
                doc.add_paragraph(block.get("text", ""), style=style)
            except Exception:
                doc.add_paragraph(block.get("text", ""))
            active_list_nums = {}
