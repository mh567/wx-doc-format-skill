from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt


SCRIPTS = Path(__file__).resolve().parents[1] / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from front_matter import (
    analyze_front_matter,
    audit_output_structure,
    front_matter_source_positions,
    inject_document_title,
)
from template_finalizer import insert_table_of_contents
from toc_detector import detect_toc_regions, selected_source_positions


def _title_paragraph(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(24)
    return paragraph


def _cover_and_toc_document() -> Document:
    doc = Document()
    _title_paragraph(doc, "统一身份认证系统集成方案")
    subtitle = doc.add_paragraph("零信任与生物特征认证")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "文档版本"
    table.cell(0, 1).text = "V1.0"
    table.cell(1, 0).text = "编制日期"
    table.cell(1, 1).text = "2026年7月"
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("目  录")
    doc.add_paragraph("1 项目概述")
    doc.add_paragraph("2 总体设计")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("1 项目概述", style="Heading 1")
    doc.add_paragraph("正文内容。")
    doc.add_paragraph("2 总体设计", style="Heading 1")
    return doc


def test_cover_before_source_toc_is_removed_and_title_extracted() -> None:
    doc = _cover_and_toc_document()
    report: dict = {}
    toc_context = detect_toc_regions(doc, report)

    context = analyze_front_matter(doc, toc_context, Path("备用文件名.docx"), report)

    assert context["status"] == "cover_detected"
    assert context["title"]["text"] == "统一身份认证系统集成方案"
    assert context["title"]["source"] == "source"
    assert front_matter_source_positions(context) == {0, 1, 2, 3}
    assert selected_source_positions(toc_context).isdisjoint(front_matter_source_positions(context))


def test_title_style_without_cover_is_normalized_once() -> None:
    doc = Document()
    style = doc.styles.add_style("文档标题", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("访问控制技术方案", style=style)
    doc.add_paragraph("项目概述", style="Heading 1")
    context = analyze_front_matter(doc, None, Path("降级标题.docx"), {})

    assert context["status"] == "title_normalized"
    assert context["cover_end_source_position"] is None
    assert context["excluded_source_positions"] == [0]
    assert context["title"]["text"] == "访问控制技术方案"


def test_document_without_cover_uses_filename_title() -> None:
    doc = Document()
    doc.add_paragraph("项目概述", style="Heading 1")
    doc.add_paragraph("正文内容。")

    context = analyze_front_matter(doc, None, Path("零信任集成方案-重构稿.docx"), {})

    assert context["status"] == "no_cover"
    assert context["excluded_source_positions"] == []
    assert context["title"]["text"] == "零信任集成方案"
    assert context["title"]["source"] == "filename"


def test_long_preamble_is_not_treated_as_cover() -> None:
    doc = Document()
    doc.add_paragraph("本文档用于说明系统的建设背景、业务边界、实施原则和验收要求，该段属于正文前言，需要在转换后完整保留。" * 2)
    doc.add_paragraph("项目概述", style="Heading 1")

    context = analyze_front_matter(doc, None, Path("系统说明.docx"), {})

    assert context["cover_end_source_position"] is None
    assert context["excluded_source_positions"] == []


def test_title_injection_keeps_exactly_one_ast_title() -> None:
    model = {
        "document": {
            "blocks": [
                {"id": "old", "block_type": "heading", "role": "title", "level": 0, "text": "旧标题"},
                {"id": "h1", "block_type": "heading", "role": "heading", "level": 1, "text": "项目概述"},
            ],
        },
    }
    context = {"title": {"text": "新标题", "source": "source", "source_position": 2}}

    inject_document_title(model, context)

    titles = [block for block in model["document"]["blocks"] if block.get("role") == "title"]
    assert len(titles) == 1
    assert model["document"]["blocks"][0]["text"] == "新标题"


def test_output_structure_requires_toc_then_title() -> None:
    doc = Document()
    title_style = doc.styles.add_style("文档标题", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("零信任集成方案", style=title_style)
    doc.add_paragraph("项目概述", style="Heading 1")
    profile = {"resolved_styles": {"title": "文档标题", "heading_1": "Heading 1"}}

    insert_table_of_contents(doc, profile)
    result = audit_output_structure(doc, profile)

    assert result["passed"] is True
    assert result["document_title_text"] == "零信任集成方案"
    assert result["document_title_count"] == 1


def test_output_structure_rejects_table_before_title() -> None:
    doc = Document()
    title_style = doc.styles.add_style("文档标题", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_table(rows=1, cols=1).cell(0, 0).text = "封面残留"
    doc.add_paragraph("零信任集成方案", style=title_style)
    profile = {"resolved_styles": {"title": "文档标题", "heading_1": "Heading 1"}}
    insert_table_of_contents(doc, profile)

    result = audit_output_structure(doc, profile)

    assert result["passed"] is False
    assert any(issue["type"] == "title_missing_after_toc" for issue in result["issues"])
