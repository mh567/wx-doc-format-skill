from __future__ import annotations

import base64
import sys
from pathlib import Path

from docx import Document


SCRIPTS = Path(__file__).resolve().parents[1] / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from caption_placement import (
    audit_model_caption_placement,
    audit_rendered_caption_placement,
)
from document_model import caption_block, image_block, new_document_model, table_block
from docx_render import render_docx_direct
from docx_pipeline import parse_docx_to_model_simple
from docx.text.paragraph import Paragraph
from text_utils import iter_blocks, looks_like_code_sample_table
from model_normalization import normalize_document_model_simple
from main import clear_document_body


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "assets" / "wx_template.docx"


def _table(block_id: str, source_position: int, table_type: str = "data") -> dict:
    return table_block(
        block_id,
        table_type,
        [[{"text": "序号", "cell_role": "header"}], [{"text": "1", "cell_role": "body"}]],
        header_rows=1,
        source={"source_position": source_position},
    )


def test_authored_table_caption_below_moves_above() -> None:
    model = new_document_model("fixture.docx", "docx", "test")
    model["document"]["blocks"] = [
        _table("t1", 0, "code_sample"),
        caption_block("c1", "接口清单", "table", source={"source_position": 1}),
    ]

    normalized = normalize_document_model_simple(model, {})

    assert [block["id"] for block in normalized["document"]["blocks"]] == ["c1", "t1"]
    assert normalized["document"]["blocks"][0]["association"]["moved"] is True
    assert audit_model_caption_placement(normalized)["passed"] is True


def test_authored_figure_caption_above_moves_below() -> None:
    model = new_document_model("fixture.docx", "docx", "test")
    model["document"]["blocks"] = [
        caption_block("c1", "架构图", "figure", source={"source_position": 0}),
        image_block("i1", source={"source_position": 1}),
    ]

    normalized = normalize_document_model_simple(model, {})

    assert [block["id"] for block in normalized["document"]["blocks"]] == ["i1", "c1"]
    assert audit_model_caption_placement(normalized)["passed"] is True


def test_ambiguous_caption_is_preserved_without_auto_generation() -> None:
    model = new_document_model("fixture.docx", "docx", "test")
    model["document"]["blocks"] = [
        _table("t1", 0),
        caption_block("c1", "待确认题注", "table", source={"source_position": 1}),
        _table("t2", 2),
    ]

    normalized = normalize_document_model_simple(model, {})

    assert [block["id"] for block in normalized["document"]["blocks"]] == ["t1", "c1", "t2"]
    assert not any(block.get("_auto_generated") for block in normalized["document"]["blocks"])
    assert normalized["caption_placement"]["issues"][0]["type"] == "ambiguous_caption_association"


def test_caption_normalization_is_idempotent() -> None:
    model = new_document_model("fixture.docx", "docx", "test")
    model["document"]["blocks"] = [
        _table("t1", 0),
        caption_block("c1", "数据表", "table", source={"source_position": 1}),
    ]

    first = normalize_document_model_simple(model, {})
    second = normalize_document_model_simple(first, {})

    assert [block["id"] for block in first["document"]["blocks"]] == ["c1", "t1"]
    assert [block["id"] for block in second["document"]["blocks"]] == ["c1", "t1"]
    assert second["caption_placement"]["moved"] == 0
    assert second["caption_placement"]["inserted"] == 0


def test_auto_generated_caption_is_reused_on_second_normalization() -> None:
    model = new_document_model("fixture.docx", "docx", "test")
    model["document"]["blocks"] = [_table("t1", 0)]

    first = normalize_document_model_simple(model, {})
    first_ids = [block["id"] for block in first["document"]["blocks"]]
    first_caption = first["document"]["blocks"][0]
    second = normalize_document_model_simple(first, {})

    assert first_ids == ["caption-auto-1", "t1"]
    assert [block["id"] for block in second["document"]["blocks"]] == first_ids
    assert first_caption["_auto_generated"] is True
    assert second["document"]["blocks"][0]["association"]["object_id"] == "t1"
    assert second["caption_placement"]["moved"] == 0
    assert second["caption_placement"]["inserted"] == 0


def test_direct_renderer_uses_model_table_caption_position() -> None:
    source = Document()
    table = source.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "序号"
    table.cell(1, 0).text = "1"
    source.add_paragraph("表1 接口清单", style="Caption")
    model = new_document_model("fixture.docx", "docx", "test")
    model["document"]["blocks"] = [
        _table("t1", 0, "code_sample"),
        caption_block("c1", "接口清单", "table", source={"source_position": 1}),
    ]
    normalized = normalize_document_model_simple(model, {})
    output = Document(TEMPLATE)
    clear_document_body(output)

    render_docx_direct(source, output, {}, 0.69, "at-least", {}, model=normalized)

    audit = audit_rendered_caption_placement(output)
    assert audit["passed"] is True
    assert audit["caption_count"] == 1


def test_direct_renderer_moves_figure_caption_after_graphic(tmp_path: Path) -> None:
    png = tmp_path / "pixel.png"
    png.write_bytes(base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wl2nK0AAAAASUVORK5CYII="
    ))
    source = Document()
    source.add_paragraph("图1 架构图", style="Caption")
    source.add_paragraph().add_run().add_picture(str(png))
    model = new_document_model("fixture.docx", "docx", "test")
    model["document"]["blocks"] = [
        caption_block("c1", "架构图", "figure", source={"source_position": 0}),
        image_block("i1", source={"source_position": 1}),
    ]
    normalized = normalize_document_model_simple(model, {})
    output = Document()

    render_docx_direct(source, output, {}, 0.69, "at-least", {}, model=normalized)

    audit = audit_rendered_caption_placement(output)
    assert audit["passed"] is True
    assert audit["caption_count"] == 1


def test_graphics_only_paragraph_is_present_in_source_ast(tmp_path: Path) -> None:
    png = tmp_path / "pixel.png"
    png.write_bytes(base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Wl2nK0AAAAASUVORK5CYII="
    ))
    source = Document()
    source.add_paragraph().add_run().add_picture(str(png))

    model = parse_docx_to_model_simple(
        tmp_path / "fixture.docx",
        source,
        True,
        0,
        skill_version=lambda: "test",
        new_report=lambda: {},
        iter_blocks=iter_blocks,
        paragraph_class=Paragraph,
        infer_docx_role=lambda paragraph, strict, report, numbering=None: ("", None, "body"),
        looks_like_code_sample_table=looks_like_code_sample_table,
        caption_pattern=None,
    )

    assert [block["block_type"] for block in model["document"]["blocks"]] == ["image"]
    assert model["document"]["blocks"][0]["source"]["source_position"] == 0
