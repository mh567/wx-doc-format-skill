from __future__ import annotations

from pathlib import Path
import json
import sys
import re
import tempfile

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from main import convert_md, convert_docx, clear_document_body, merge_template_numbering_ids, skill_version, build_document_model_from_output_wrapper, audit_document_wrapper
from reporting import new_report
from template_profile import load_template_profile
from template_finalizer import apply_template_finalizer
from text_utils import looks_like_code_sample_table, set_table_autofit_to_window
from md_pipeline import parse_md_to_model
from model_normalization import normalize_document_model
from word_model_renderer import style_from_profile


# Template path: relative to repo root, falls back to assets/

def _template_path() -> Path:
    """Find the template docx.  Looks in assets/ relative to repo root."""
    repo = Path(__file__).resolve().parents[1]
    p = repo / "assets" / "wx_template.docx"
    if p.exists():
        return p
    p = repo / "wx_template.docx"
    if p.exists():
        return p
    raise FileNotFoundError("Template not found in assets/ or repo root")



def _md_source(text: str) -> Path:
    p = Path(tempfile.mkdtemp()) / "source.md"
    p.write_text(text, encoding="utf-8")
    return p


# ============================================================
# Step 1: 解析测试
# ============================================================

def test_md_parse_creates_source_ast():
    rpt = new_report("test")
    model = parse_md_to_model(
        _md_source("# 测试文档\n\n第一章 测试标题\n\n正文内容。\n\na) 第一项\nb) 第二项\n"),
        rpt, skill_version=lambda: "test"
    )
    blocks = model.get("document", {}).get("blocks", [])
    types = [b["block_type"] for b in blocks]
    assert "heading" in types
    assert "body" in types
    assert "list_item" in types
    assert len(blocks) >= 4


def test_md_parse_heading_levels():
    rpt = new_report("test")
    model = parse_md_to_model(
        _md_source("# 文档标题\n\n## 一级标题\n\n### 二级标题\n"),
        rpt, skill_version=lambda: "test"
    )
    blocks = model.get("document", {}).get("blocks", [])
    assert blocks[0]["block_type"] == "heading"
    assert blocks[0].get("level") == 0
    assert blocks[1]["block_type"] == "heading"
    assert blocks[1].get("level") == 1
    assert blocks[2]["block_type"] == "heading"
    assert blocks[2].get("level") == 2


def test_md_parse_inferred_heading():
    rpt = new_report("test")
    model = parse_md_to_model(
        _md_source("第一章 推断标题\n\n1.1 二级标题\n"),
        rpt, skill_version=lambda: "test"
    )
    blocks = model.get("document", {}).get("blocks", [])
    bs = [(b["block_type"], b.get("level")) for b in blocks]
    assert ("heading", 1) in bs
    assert ("heading", 2) in bs


def test_md_parse_list_restart():
    rpt = new_report("test")
    model = parse_md_to_model(
        _md_source("a) 第一项\nb) 第二项\n"),
        rpt, skill_version=lambda: "test"
    )
    blocks = model.get("document", {}).get("blocks", [])
    assert blocks[0]["block_type"] == "list_item"
    assert blocks[0].get("restart") == True
    assert blocks[1].get("restart") == False


def test_md_parse_note():
    rpt = new_report("test")
    model = parse_md_to_model(
        _md_source("备注：这是一个注释。\n"),
        rpt, skill_version=lambda: "test"
    )
    blocks = model.get("document", {}).get("blocks", [])
    assert blocks[0]["block_type"] == "body"
    assert blocks[0].get("source", {}).get("role") == "note"


# ============================================================
# Step 2: 规范化测试
# ============================================================

def test_normalization_removes_manual_heading_number():
    rpt = new_report("test")
    source = parse_md_to_model(
        _md_source("第一章 测试标题\n"),
        rpt, skill_version=lambda: "test"
    )
    model = normalize_document_model(source, rpt)
    blocks = model.get("document", {}).get("blocks", [])
    assert blocks[0]["block_type"] == "heading"
    assert "测试标题" in blocks[0].get("text", "")
    assert "第一章" not in blocks[0].get("text", "")


def test_normalization_sets_heading_numbering_auto():
    rpt = new_report("test")
    source = parse_md_to_model(
        _md_source("1.1 二级标题\n"),
        rpt, skill_version=lambda: "test"
    )
    normalized = normalize_document_model(source, rpt)
    blocks = normalized.get("document", {}).get("blocks", [])
    assert blocks[0].get("numbering", {}).get("mode") == "auto"


def test_normalization_marks_list_restart():
    rpt = new_report("test")
    source = parse_md_to_model(
        _md_source("第一节 小节\n\na) 第一项\nb) 第二项\n"),
        rpt, skill_version=lambda: "test"
    )
    normalized = normalize_document_model(source, rpt)
    list_items = [b for b in normalized.get("document", {}).get("blocks", []) if b["block_type"] == "list_item"]
    assert len(list_items) >= 2
    assert list_items[0].get("restart") == True
    assert list_items[1].get("restart") == False


# ============================================================
# Step 3: 模板渲染测试
# ============================================================

def test_render_heading_uses_template_style():
    profile = load_template_profile(str(_template_path()))
    doc = Document(str(_template_path()))
    clear_document_body(doc)
    nids = merge_template_numbering_ids(profile, {})
    rpt = new_report("test")
    rpt['template_profile'] = {"resolved_styles": profile.get("resolved_styles", {})}

    convert_md(
        _md_source("# 文档标题\n\n第一章 一级标题\n\n## 二级标题\n"),
        doc, rpt, 0.69, "at-least", nids
    )

    styles = [p.style.name for p in doc.paragraphs if p.text.strip()]
    assert "文档标题" in styles
    assert any("heading 1" in s or "Heading 1" in s for s in styles)
    assert any("heading 2" in s or "Heading 2" in s or "Heading 1" in s for s in styles)


def test_render_body_uses_template_style():
    profile = load_template_profile(str(_template_path()))
    doc = Document(str(_template_path()))
    clear_document_body(doc)
    nids = merge_template_numbering_ids(profile, {})
    rpt = new_report("test")
    rpt['template_profile'] = {"resolved_styles": profile.get("resolved_styles", {})}

    convert_md(_md_source("正文内容。\n"), doc, rpt, 0.69, "at-least", nids)

    body_paras = [p for p in doc.paragraphs if p.text.strip()]
    assert len(body_paras) >= 1
    body_style = style_from_profile(profile, "body", "Normal")
    assert body_paras[0].style.name == body_style


def test_render_list_uses_template_style():
    profile = load_template_profile(str(_template_path()))
    doc = Document(str(_template_path()))
    clear_document_body(doc)
    nids = merge_template_numbering_ids(profile, {})
    rpt = new_report("test")
    rpt['template_profile'] = {"resolved_styles": profile.get("resolved_styles", {})}

    convert_md(_md_source("a) 第一项\nb) 第二项\n"), doc, rpt, 0.69, "at-least", nids)

    list_paras = [p for p in doc.paragraphs if p.text.strip()]
    assert len(list_paras) >= 2
    list_style = style_from_profile(profile, "list_letter", "1.1一级列项-编号")
    assert list_paras[0].style.name == list_style


def test_render_no_unexpected_styles():
    profile = load_template_profile(str(_template_path()))
    doc = Document(str(_template_path()))
    clear_document_body(doc)
    nids = merge_template_numbering_ids(profile, {})
    rpt = new_report(skill_version())
    rpt['template_profile'] = {"resolved_styles": profile.get("resolved_styles", {}), "missing_roles": [], "numbering_ids": nids}

    convert_md(_md_source("# 标题\n\n正文。\n\na) 列项\n"), doc, rpt, 0.69, "at-least", nids)

    result = apply_template_finalizer(
        doc, profile, 0.69, "at-least",
        row_height_rule_enum=WD_ROW_HEIGHT_RULE, cm=Cm,
        left_alignment=WD_ALIGN_PARAGRAPH.LEFT,
        center_alignment=WD_ALIGN_PARAGRAPH.CENTER,
        set_table_autofit_to_window=set_table_autofit_to_window,
        looks_like_code_sample_table=looks_like_code_sample_table,
    )

    unexpected = result.get("style_audit", {}).get("unexpected_styles", [])
    assert len(unexpected) == 0, f"Unexpected styles: {unexpected}"


# ============================================================
# 工具函数测试
# ============================================================

def test_heading_level_from_text():
    from text_utils import heading_level_from_text
    assert heading_level_from_text("第一章 测试") == 1
    assert heading_level_from_text("1.1 测试") == 2
    assert heading_level_from_text("1.1.1 测试") == 3
    assert heading_level_from_text("一、测试") == 1
    assert heading_level_from_text("这是一段超过三十个字的正文内容，不应该被认为是标题") is None
    assert heading_level_from_text("") is None


def test_is_caption_text():
    from text_utils import is_caption_text
    assert is_caption_text("表 69 测试表")
    assert is_caption_text("图 1 架构图")
    assert is_caption_text("表 107-1 接口说明")
    assert not is_caption_text("这是一个普通段落")
    assert not is_caption_text("")


def test_looks_like_list_item():
    from text_utils import looks_like_list_item
    assert looks_like_list_item("a) 第一项")
    assert looks_like_list_item("1) 第一项")
    assert looks_like_list_item("—— 无编号项")
    assert looks_like_list_item("• 中点项")
    assert not looks_like_list_item("普通正文内容")


def test_strip_heading_marker():
    from text_utils import strip_heading_marker
    assert strip_heading_marker("第一章 测试") == "测试"
    assert strip_heading_marker("1.1 二级标题") == "二级标题"
    assert strip_heading_marker("一、核心定位") == "核心定位"
    assert strip_heading_marker("无编号的标题") == "无编号的标题"


def test_strip_list_marker():
    from text_utils import strip_list_marker
    assert strip_list_marker("a) 内容") == "内容"
    assert strip_list_marker("1) 内容") == "内容"
    assert strip_list_marker("—— 内容") == "内容"
    assert strip_list_marker("正文") == "正文"


def test_style_from_profile():
    profile = load_template_profile(str(_template_path()))
    assert style_from_profile(profile, "body", "Normal") == "Normal"
    assert style_from_profile(profile, "title", "文档标题") == "文档标题"
    assert style_from_profile(profile, "heading_1", "Heading 1") in ("heading 1", "Heading 1")
    assert style_from_profile(profile, "list_letter", "1.1一级列项-编号") == "1.1一级列项-编号"
    assert style_from_profile(None, "body", "Normal") == "Normal"


# ============================================================
# 完整管线测试
# ============================================================

def test_full_pipeline_md_no_risk(tmp_path):
    profile = load_template_profile(str(_template_path()))
    doc = Document(str(_template_path()))
    clear_document_body(doc)
    nids = merge_template_numbering_ids(profile, {})
    rpt = new_report(skill_version())
    rpt['template_profile'] = {"path": str(_template_path()), "resolved_styles": profile.get("resolved_styles", {}), "missing_roles": [], "numbering_ids": nids}

    models = convert_md(
        _md_source("# 测试\n\n第一章 标题\n\n正文。\n\na) 列项"),
        doc, rpt, 0.69, "at-least", nids
    )

    result = apply_template_finalizer(
        doc, profile, 0.69, "at-least",
        row_height_rule_enum=WD_ROW_HEIGHT_RULE, cm=Cm,
        left_alignment=WD_ALIGN_PARAGRAPH.LEFT,
        center_alignment=WD_ALIGN_PARAGRAPH.CENTER,
        set_table_autofit_to_window=set_table_autofit_to_window,
        looks_like_code_sample_table=looks_like_code_sample_table,
    )

    output_path = tmp_path / "output.docx"
    doc.save(str(output_path))

    assert models["source"] is not None
    assert models["normalized"] is not None
    assert len(result.get("style_audit", {}).get("unexpected_styles", [])) == 0
    assert output_path.exists()
    assert output_path.stat().st_size > 1000


def test_full_pipeline_markdown_table(tmp_path):
    profile = load_template_profile(str(_template_path()))
    doc = Document(str(_template_path()))
    clear_document_body(doc)
    nids = merge_template_numbering_ids(profile, {})
    rpt = new_report(skill_version())
    rpt['template_profile'] = {"path": str(_template_path()), "resolved_styles": profile.get("resolved_styles", {}), "missing_roles": [], "numbering_ids": nids}

    models = convert_md(
        _md_source("# 测试\n\n| 名称 | 说明 |\n|------|------|\n| 字段1 | 说明1 |\n| 字段2 | 说明2 |\n"),
        doc, rpt, 0.69, "at-least", nids
    )

    result = apply_template_finalizer(
        doc, profile, 0.69, "at-least",
        row_height_rule_enum=WD_ROW_HEIGHT_RULE, cm=Cm,
        left_alignment=WD_ALIGN_PARAGRAPH.LEFT,
        center_alignment=WD_ALIGN_PARAGRAPH.CENTER,
        set_table_autofit_to_window=set_table_autofit_to_window,
        looks_like_code_sample_table=looks_like_code_sample_table,
    )

    output_path = tmp_path / "output.docx"
    doc.save(str(output_path))

    assert rpt.get("tables_processed", 0) >= 1
    assert len(result.get("style_audit", {}).get("unexpected_styles", [])) == 0
