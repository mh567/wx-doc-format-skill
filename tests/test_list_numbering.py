from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SCRIPTS = Path(__file__).resolve().parents[1] / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from list_detector import analyze_docx_lists, audit_list_preservation
from list_style_mapping import normalize_wx_list_type, wx_list_style_name
from llm_enhancer import enhance_document_model, validate_patch
from template_profile import load_template_profile, template_numbering_ids
from docx_render import _handle_inferred
from word_model_renderer import render_document_model


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "assets" / "wx_template.docx"


def _style_num_id(doc: Document, style_name: str) -> int:
    num_pr = doc.styles[style_name].element.pPr.numPr
    assert num_pr is not None and num_pr.numId is not None
    return int(num_pr.numId.val)


def _set_direct_numbering(paragraph, num_id: int, ilvl: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def test_list_number_style_is_detected() -> None:
    doc = Document()
    doc.add_paragraph("功能一", style="List Number")
    doc.add_paragraph("功能二", style="List Number")
    report: dict = {}

    context = analyze_docx_lists(doc, report)

    assert report["source_lists"]["detected"] == 2
    assert all(item["status"] == "detected" for item in context.values())
    assert all(item["source_list_type"] == "decimal_paren" for item in context.values())
    assert all(item["list_type"] == "lower_letter_paren" for item in context.values())
    assert sum(item["restart"] for item in context.values()) == 1


def test_list_bullet_style_is_detected() -> None:
    doc = Document()
    doc.add_paragraph("功能一", style="List Bullet")
    doc.add_paragraph("功能二", style="List Bullet")

    context = analyze_docx_lists(doc, {})

    assert all(item["status"] == "detected" for item in context.values())
    assert all(item["source_list_type"] == "bullet_dot" for item in context.values())
    assert all(item["list_type"] == "dash" for item in context.values())


def test_list_paragraph_and_custom_style_use_numbering_evidence() -> None:
    doc = Document()
    num_id = _style_num_id(doc, "List Number")
    custom = doc.styles.add_style("项目功能项", WD_STYLE_TYPE.PARAGRAPH)
    first = doc.add_paragraph("功能一", style="List Paragraph")
    second = doc.add_paragraph("功能二", style=custom)
    _set_direct_numbering(first, num_id)
    _set_direct_numbering(second, num_id)

    context = analyze_docx_lists(doc, {})

    assert [item["status"] for item in context.values()] == ["detected", "detected"]


def test_multilevel_numbering_preserves_ilvl() -> None:
    doc = Document()
    num_id = _style_num_id(doc, "List Number")
    first = doc.add_paragraph("一级功能", style="List Number")
    second = doc.add_paragraph("二级功能", style="List Number")
    _set_direct_numbering(first, num_id, 0)
    _set_direct_numbering(second, num_id, 1)

    context = analyze_docx_lists(doc, {})

    assert [item["ilvl"] for item in context.values()] == [0, 1]
    assert [item["list_type"] for item in context.values()] == [
        "lower_letter_paren", "decimal_paren",
    ]


def test_wx_list_mapping_follows_level_and_ordered_semantics() -> None:
    assert normalize_wx_list_type("decimal_paren", 0) == "lower_letter_paren"
    assert normalize_wx_list_type("lower_letter_paren", 1) == "decimal_paren"
    assert normalize_wx_list_type("bullet_dot", 0) == "dash"
    assert normalize_wx_list_type("dash", 1) == "bullet_dot"
    assert wx_list_style_name("decimal_paren", 0) == "1.1一级列项-编号"
    assert wx_list_style_name("lower_letter_paren", 1) == "2.1二级列项-有编号"


def test_direct_and_model_renderers_share_wx_list_styles() -> None:
    profile = load_template_profile(TEMPLATE)
    numbering_ids = template_numbering_ids(profile)
    cases = [
        (0, "decimal_paren", "1.1一级列项-编号"),
        (1, "lower_letter_paren", "2.1二级列项-有编号"),
        (0, "bullet_dot", "1.2一级列项-无编号"),
        (1, "dash", "2.2二级列项-无编号"),
    ]
    for level, list_type, expected_style in cases:
        direct_doc = Document(TEMPLATE)
        _handle_inferred(
            direct_doc,
            "功能项",
            None,
            "list_item",
            {},
            numbering_ids,
            {},
            template_profile=profile,
            list_meta={"level": level, "list_type": list_type, "restart": True},
        )
        assert direct_doc.paragraphs[-1].style.name == expected_style

        model_doc = Document(TEMPLATE)
        render_document_model(
            {
                "document": {
                    "blocks": [{
                        "id": "b0001",
                        "block_type": "list_item",
                        "text": "功能项",
                        "level": level,
                        "list_type": list_type,
                        "restart": True,
                    }],
                },
            },
            model_doc,
            {},
            0.8,
            "at-least",
            numbering_ids,
            template_profile=profile,
        )
        assert model_doc.paragraphs[-1].style.name == expected_style


def test_numbered_heading_is_protected() -> None:
    doc = Document()
    num_id = _style_num_id(doc, "List Number")
    heading = doc.add_paragraph("项目概述", style="Heading 1")
    _set_direct_numbering(heading, num_id)

    context = analyze_docx_lists(doc, {})

    item = next(iter(context.values()))
    assert item["status"] == "ignored"
    assert "protected_role" in item["evidence"]


def test_numbered_toc_style_is_protected() -> None:
    doc = Document()
    num_id = _style_num_id(doc, "List Number")
    toc_style = doc.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    paragraph = doc.add_paragraph("项目概述", style=toc_style)
    _set_direct_numbering(paragraph, num_id)

    context = analyze_docx_lists(doc, {})

    assert next(iter(context.values()))["status"] == "ignored"


def test_isolated_normal_numbering_is_ambiguous() -> None:
    doc = Document()
    num_id = _style_num_id(doc, "List Number")
    paragraph = doc.add_paragraph("可能是残留编号的正文。")
    _set_direct_numbering(paragraph, num_id)
    report: dict = {}

    context = analyze_docx_lists(doc, report)

    item = next(iter(context.values()))
    assert item["status"] == "ambiguous"
    assert report["source_lists"]["ambiguous"] == 1


def test_invalid_num_id_is_ignored() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("无效编号定义。")
    _set_direct_numbering(paragraph, 9999)

    context = analyze_docx_lists(doc, {})

    item = next(iter(context.values()))
    assert item["status"] == "ignored"
    assert "invalid_numbering_definition" in item["evidence"]


def test_ambiguous_llm_candidate_must_preserve_numbering_shape() -> None:
    model = {
        "schema_version": "1.0",
        "document": {
            "blocks": [{
                "id": "b0001",
                "block_type": "body",
                "text": "候选段落",
                "source": {
                    "style": "Normal",
                    "numbering": {
                        "status": "ambiguous",
                        "ilvl": 0,
                        "source_list_type": "decimal_paren",
                        "list_type": "lower_letter_paren",
                    },
                },
            }],
        },
    }
    patch = {
        "schema_version": "1.0",
        "phase": "A",
        "decisions": [{
            "block_id": "b0001",
            "operation": "retype",
            "from": {"block_type": "body"},
            "to": {
                "block_type": "list_item",
                "level": 1,
                "list_type": "lower_letter_paren",
            },
            "confidence": 0.90,
        }],
    }

    errors = validate_patch(patch, model, frozenset({"retype"}))

    assert any("preserve level and list_type" in error.get("message", "") for error in errors)


def test_ambiguous_llm_candidate_accepts_exact_numbering_shape() -> None:
    model = {
        "schema_version": "1.0",
        "document": {
            "blocks": [{
                "id": "b0001",
                "block_type": "body",
                "text": "候选段落",
                "source": {
                    "style": "Normal",
                    "numbering": {
                        "status": "ambiguous",
                        "ilvl": 0,
                        "source_list_type": "decimal_paren",
                        "list_type": "lower_letter_paren",
                    },
                },
            }],
        },
    }
    patch = {
        "schema_version": "1.0",
        "phase": "A",
        "decisions": [{
            "block_id": "b0001",
            "operation": "retype",
            "from": {"block_type": "body"},
            "to": {
                "block_type": "list_item",
                "level": 0,
                "list_type": "lower_letter_paren",
            },
            "confidence": 0.90,
        }],
    }

    errors = validate_patch(patch, model, frozenset({"retype"}))

    assert errors == []


def test_enhancer_applies_exact_ambiguous_list_decision() -> None:
    model = {
        "schema_version": "1.0",
        "document": {
            "blocks": [{
                "id": "b0001",
                "block_type": "body",
                "text": "可能的编号列表项",
                "source": {
                    "style": "Normal",
                    "numbering": {
                        "status": "ambiguous",
                        "ilvl": 0,
                        "source_list_type": "decimal_paren",
                        "list_type": "lower_letter_paren",
                    },
                },
            }],
        },
    }
    response = """{
      "schema_version": "1.0",
      "phase": "A",
      "decisions": [{
        "block_id": "b0001",
        "operation": "retype",
        "from": {"block_type": "body"},
        "to": {"block_type": "list_item", "level": 0, "list_type": "lower_letter_paren"},
        "confidence": 0.90
      }]
    }"""
    report = {"parse_report": {"ambiguous_numbered_paragraphs": 1}}

    enhanced = enhance_document_model(
        model,
        report,
        phase="list_detect",
        llm_call=lambda _: response,
    )

    assert enhanced["document"]["blocks"][0]["block_type"] == "list_item"
    assert len(report["llm_enhancer"]["applied"]) == 1


def test_enhancer_keeps_low_confidence_ambiguous_candidate_as_body() -> None:
    model = {
        "schema_version": "1.0",
        "document": {
            "blocks": [{
                "id": "b0001",
                "block_type": "body",
                "text": "可能的编号列表项",
                "source": {
                    "style": "Normal",
                    "numbering": {
                        "status": "ambiguous",
                        "ilvl": 0,
                        "source_list_type": "decimal_paren",
                        "list_type": "lower_letter_paren",
                    },
                },
            }],
        },
    }
    response = """{
      "schema_version": "1.0",
      "phase": "A",
      "decisions": [{
        "block_id": "b0001",
        "operation": "retype",
        "from": {"block_type": "body"},
        "to": {"block_type": "list_item", "level": 0, "list_type": "lower_letter_paren"},
        "confidence": 0.60
      }]
    }"""
    report = {"parse_report": {"ambiguous_numbered_paragraphs": 1}}

    enhanced = enhance_document_model(
        model,
        report,
        phase="list_detect",
        llm_call=lambda _: response,
    )

    assert enhanced["document"]["blocks"][0]["block_type"] == "body"
    assert len(report["llm_enhancer"]["skipped"]) == 1


def test_rendered_list_audit_excludes_numbered_title_style() -> None:
    doc = Document()
    title_style = doc.styles.add_style("文档标题", WD_STYLE_TYPE.PARAGRAPH)
    list_style = doc.styles.add_style("1.1一级列项-编号", WD_STYLE_TYPE.PARAGRAPH)
    title = doc.add_paragraph("文档标题", style=title_style)
    list_num_id = _style_num_id(doc, "List Number")
    _set_direct_numbering(title, list_num_id)
    doc.add_paragraph("功能一", style=list_style)
    model = {
        "document": {
            "blocks": [{
                "id": "b0001",
                "block_type": "list_item",
                "text": "功能一",
                "source": {"numbering": {"status": "detected"}},
            }],
        },
    }

    result = audit_list_preservation(
        doc, model, {"detected": 1}, load_template_profile(TEMPLATE),
    )

    assert result["rendered_list_items"] == 1
    assert result["passed"] is True


def test_rendered_list_audit_rejects_secondary_style_for_level_zero() -> None:
    doc = Document()
    secondary = doc.styles.add_style("2.1二级列项-有编号", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("功能一", style=secondary)
    model = {
        "document": {
            "blocks": [{
                "id": "b0001",
                "block_type": "list_item",
                "text": "功能一",
                "level": 0,
                "list_type": "lower_letter_paren",
                "source": {"numbering": {"status": "detected"}},
            }],
        },
    }

    result = audit_list_preservation(doc, model, {"detected": 1})

    assert result["style_level_mismatches"] == [{
        "block_id": "b0001",
        "level": 0,
        "list_type": "lower_letter_paren",
        "expected_style": "1.1一级列项-编号",
        "rendered_style": "2.1二级列项-有编号",
    }]
    assert result["passed"] is False
