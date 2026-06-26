from pathlib import Path
import base64
import sys
import zipfile


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

from format_document import (
    DEFAULT_TABLE_ROW_HEIGHT_CM,
    DEFAULT_TABLE_ROW_HEIGHT_RULE,
    audit_document,
    convert_docx,
    ensure_auto_numbering,
    ensure_fallback_styles,
    existing_heading_number,
    heading_level_from_text,
    infer_docx_role,
    is_caption_text,
    is_date_like_text,
    is_toc_title,
    normalize_graphics_paragraph,
    new_report,
    new_num_for_abstract,
    normalize_table,
    paragraph_numbering_descriptor,
    paragraph_has_graphics,
    resolved_heading_level,
    scan_non_text_objects,
    source_numbering_heading_level,
    strip_heading_marker,
)
from update_installed_skill import read_version


def test_report_includes_skill_version():
    assert new_report()["skill_version"] == (ROOT / "VERSION").read_text(encoding="utf-8").strip()


def test_chinese_number_heading_is_level_one():
    assert heading_level_from_text("一、核心定位") == 1
    assert heading_level_from_text("二、总体设计思路") == 1


def test_numbered_decimal_heading_levels_still_work():
    assert heading_level_from_text("1 总则") == 1
    assert heading_level_from_text("1.1 总体架构") == 2
    assert heading_level_from_text("11.1.1 全局视角") == 3


def test_date_and_caption_are_not_heading_markers():
    assert is_date_like_text("2026 年 6 月")
    assert is_toc_title("目  录")
    assert is_caption_text("图 2-1　AI 安全能力中台总体逻辑架构")
    assert heading_level_from_text("2026 年 6 月") is None
    assert heading_level_from_text("图 2-1　AI 安全能力中台总体逻辑架构") is None


def test_chinese_number_marker_is_stripped():
    assert strip_heading_marker("一、核心定位") == "核心定位"
    assert strip_heading_marker("十五、最后收尾备忘") == "最后收尾备忘"
    assert strip_heading_marker("（一）智能体平台技术路线") == "智能体平台技术路线"
    assert strip_heading_marker("(1) 数字括号标题") == "数字括号标题"
    assert existing_heading_number("一、核心定位")
    assert existing_heading_number("（一）智能体平台技术路线")


def test_heading_style_level_survives_without_num_id():
    assert resolved_heading_level("Heading 2", 1, "分层架构设计") == 2
    assert resolved_heading_level("Heading 2", None, "统一认证中心") == 2


def add_test_numbering(doc, paragraph, num_id: int, abstract_id: int, num_fmt: str, lvl_text: str) -> None:
    numbering = doc.part.numbering_part.element
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_id))
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), lvl_text)
    level.append(fmt)
    level.append(text)
    abstract_num.append(level)
    numbering.append(abstract_num)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def test_chinese_source_numbering_short_heading_is_level_one():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("AI能力中台定位").bold = True
    add_test_numbering(doc, paragraph, 101, 201, "chineseCounting", "%1、")

    assert paragraph_numbering_descriptor(paragraph) == ("chineseCounting", "%1、")
    assert source_numbering_heading_level(paragraph) == 1
    assert infer_docx_role(paragraph, True, new_report()) == ("AI能力中台定位", "Heading 1", "heading")


def test_decimal_source_numbering_short_heading_is_level_two():
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("目标架构图").bold = True
    add_test_numbering(doc, paragraph, 102, 202, "decimal", "%1.")

    assert paragraph_numbering_descriptor(paragraph) == ("decimal", "%1.")
    assert source_numbering_heading_level(paragraph) == 2
    assert infer_docx_role(paragraph, True, new_report()) == ("目标架构图", "Heading 2", "heading")


def test_new_list_numbering_instances_restart_at_one():
    doc = Document()
    numbering_ids = ensure_auto_numbering(doc)
    num_id = new_num_for_abstract(doc, numbering_ids["list_letter_abstract"])
    numbering = doc.part.numbering_part.element
    num = next(element for element in numbering.findall(qn("w:num")) if element.get(qn("w:numId")) == str(num_id))
    start = num.find(".//" + qn("w:startOverride"))
    assert start is not None
    assert start.get(qn("w:val")) == "1"


def test_audit_tracks_heading_sequence_and_list_restart_groups():
    doc = Document()
    ensure_fallback_styles(doc)
    numbering_ids = ensure_auto_numbering(doc)
    heading = doc.add_paragraph("测试章节", style="Heading 1")
    from format_document import apply_numbering

    apply_numbering(heading, numbering_ids["heading"], 0)
    num_id = new_num_for_abstract(doc, numbering_ids["list_letter_abstract"])
    item = doc.add_paragraph("测试列项", style="1.1一级列项-编号")
    apply_numbering(item, num_id, 0)

    audit = audit_document(doc, 0.69, "at-least")
    assert audit["heading_sequence"][0]["text"] == "测试章节"
    assert audit["list_restart_groups"][0]["restart_at_one"] is True
    assert audit["ordered_list_nums_without_restart"] == []


def test_default_table_row_height_rule_is_at_least():
    doc = Document()
    ensure_fallback_styles(doc)
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "这是一段较长的表格文字，用于验证默认最小行高允许单元格内容自然换行展示。"

    normalize_table(table, DEFAULT_TABLE_ROW_HEIGHT_CM, DEFAULT_TABLE_ROW_HEIGHT_RULE)
    audit = audit_document(doc, DEFAULT_TABLE_ROW_HEIGHT_CM, DEFAULT_TABLE_ROW_HEIGHT_RULE)

    assert DEFAULT_TABLE_ROW_HEIGHT_RULE == "at-least"
    assert table.rows[0].height_rule == WD_ROW_HEIGHT_RULE.AT_LEAST
    assert audit["table_rows_bad_height"] == []
    assert audit["table_cells_may_clip"] == []


def test_table_autofits_to_window():
    doc = Document()
    ensure_fallback_styles(doc)
    table = doc.add_table(rows=1, cols=1)

    normalize_table(table, DEFAULT_TABLE_ROW_HEIGHT_CM, DEFAULT_TABLE_ROW_HEIGHT_RULE)

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    tbl_grid = OxmlElement("w:tblGrid")
    grid_col = OxmlElement("w:gridCol")
    grid_col.set(qn("w:w"), "1200")
    tbl_grid.append(grid_col)
    table._tbl.append(tbl_grid)
    tc_pr = table.cell(0, 0)._tc.get_or_add_tcPr()
    tc_w = tc_pr.get_or_add_tcW()
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), "1200")
    no_wrap = OxmlElement("w:noWrap")
    tc_pr.append(no_wrap)

    normalize_table(table, DEFAULT_TABLE_ROW_HEIGHT_CM, DEFAULT_TABLE_ROW_HEIGHT_RULE)

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    tc_w = table.cell(0, 0)._tc.tcPr.find(qn("w:tcW"))
    assert table.autofit is True
    assert table._tbl.find(qn("w:tblGrid")) is not None
    assert tbl_w is not None
    assert tbl_w.get(qn("w:type")) == "pct"
    assert tbl_w.get(qn("w:w")) == "5000"
    assert tbl_layout is not None
    assert tbl_layout.get(qn("w:type")) == "autofit"
    assert tc_w is not None
    assert tc_w.get(qn("w:type")) == "auto"
    assert tc_w.get(qn("w:w")) == "0"
    assert no_wrap.getparent() is None


def test_graphics_paragraph_layout_is_cleared_and_centered():
    doc = Document()
    paragraph = doc.add_paragraph("placeholder")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.left_indent = Cm(2)
    paragraph.paragraph_format.first_line_indent = Cm(1)

    normalize_graphics_paragraph(paragraph)

    assert paragraph.style.name == "Normal"
    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    p_pr = paragraph._p.pPr
    assert p_pr.find(qn("w:numPr")) is None
    ind = p_pr.find(qn("w:ind"))
    assert ind is not None
    assert ind.get(qn("w:left")) == "0"
    assert ind.get(qn("w:firstLine")) == "0"


def test_caption_role_precedes_visual_heading():
    doc = Document()
    ensure_fallback_styles(doc)
    paragraph = doc.add_paragraph("图 2-1　AI 安全能力中台总体逻辑架构")
    paragraph.alignment = 1

    text, style, role = infer_docx_role(paragraph, True, new_report())

    assert text == "图 2-1　AI 安全能力中台总体逻辑架构"
    assert style == "Caption"
    assert role == "caption"


def test_updater_reads_version(tmp_path):
    skill_dir = tmp_path / "skill"
    skill_dir.mkdir()
    (skill_dir / "VERSION").write_text("9.9.9\n", encoding="utf-8")
    assert read_version(skill_dir) == "9.9.9"
    assert read_version(tmp_path / "missing") == "unknown"


def test_docx_image_paragraphs_are_preserved(tmp_path):
    image_path = tmp_path / "tiny.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
        )
    )
    src_path = tmp_path / "source.docx"
    output_path = tmp_path / "output.docx"

    src_doc = Document()
    src_doc.add_paragraph("1 总体定位", style="Heading 1")
    src_doc.add_picture(str(image_path))
    src_doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
    src_doc.paragraphs[-1].paragraph_format.left_indent = Cm(2)
    src_doc.save(src_path)

    out_doc = Document()
    ensure_fallback_styles(out_doc)
    numbering_ids = ensure_auto_numbering(out_doc)
    report = new_report()
    report["non_text_objects"] = scan_non_text_objects(src_path)
    convert_docx(src_path, out_doc, 0.69, "at-least", True, report, numbering_ids)
    out_doc.save(output_path)

    assert report["graphic_paragraphs_preserved"] == 1
    assert report["media_relationships_preserved"] == 1
    converted = Document(output_path)
    assert len(converted.inline_shapes) == 1
    assert converted.paragraphs[1].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraph_has_graphics(Document(src_path).paragraphs[1])
    with zipfile.ZipFile(output_path) as zf:
        assert any(name.startswith("word/media/") for name in zf.namelist())


def test_heading_text_and_image_in_one_paragraph_are_split(tmp_path):
    image_path = tmp_path / "tiny.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
        )
    )
    src_path = tmp_path / "source.docx"
    output_path = tmp_path / "output.docx"

    src_doc = Document()
    paragraph = src_doc.add_paragraph()
    paragraph.add_run("11.1.1 全局视角")
    paragraph.add_run().add_picture(str(image_path))
    src_doc.save(src_path)

    out_doc = Document()
    ensure_fallback_styles(out_doc)
    numbering_ids = ensure_auto_numbering(out_doc)
    report = new_report()
    report["non_text_objects"] = scan_non_text_objects(src_path)
    convert_docx(src_path, out_doc, 0.69, "at-least", True, report, numbering_ids)
    out_doc.save(output_path)

    converted = Document(output_path)
    assert converted.paragraphs[0].style.name == "Heading 3"
    assert converted.paragraphs[0].text == "全局视角"
    assert not paragraph_has_graphics(converted.paragraphs[0])
    assert converted.paragraphs[1].text == ""
    assert paragraph_has_graphics(converted.paragraphs[1])
    assert report["mixed_text_graphic_paragraphs_split"] == [
        {"text": "11.1.1 全局视角", "role": "heading", "level": 3}
    ]
    assert report["semantic_object_splits"] == [
        {"text": "11.1.1 全局视角", "role": "heading", "level": 3}
    ]
    assert report["media_relationships_preserved"] == 1


def test_body_text_and_image_in_one_paragraph_are_split(tmp_path):
    image_path = tmp_path / "tiny.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
        )
    )
    src_path = tmp_path / "source.docx"
    output_path = tmp_path / "output.docx"

    src_doc = Document()
    paragraph = src_doc.add_paragraph()
    paragraph.add_run("这是一段正文说明，用来描述后续图片。")
    paragraph.add_run().add_picture(str(image_path))
    src_doc.save(src_path)

    out_doc = Document()
    ensure_fallback_styles(out_doc)
    numbering_ids = ensure_auto_numbering(out_doc)
    report = new_report()
    report["non_text_objects"] = scan_non_text_objects(src_path)
    convert_docx(src_path, out_doc, 0.69, "at-least", True, report, numbering_ids)
    out_doc.save(output_path)

    converted = Document(output_path)
    assert converted.paragraphs[0].style.name == "Normal"
    assert converted.paragraphs[0].text == "这是一段正文说明，用来描述后续图片。"
    assert not paragraph_has_graphics(converted.paragraphs[0])
    assert converted.paragraphs[1].text == ""
    assert paragraph_has_graphics(converted.paragraphs[1])
    assert report["semantic_object_splits"] == [
        {"text": "这是一段正文说明，用来描述后续图片。", "role": "body"}
    ]
    assert report["media_relationships_preserved"] == 1


def test_list_text_and_image_in_one_paragraph_are_split(tmp_path):
    image_path = tmp_path / "tiny.png"
    image_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
        )
    )
    src_path = tmp_path / "source.docx"
    output_path = tmp_path / "output.docx"

    src_doc = Document()
    paragraph = src_doc.add_paragraph()
    paragraph.add_run("a) 第一项说明")
    paragraph.add_run().add_picture(str(image_path))
    src_doc.save(src_path)

    out_doc = Document()
    ensure_fallback_styles(out_doc)
    numbering_ids = ensure_auto_numbering(out_doc)
    report = new_report()
    report["non_text_objects"] = scan_non_text_objects(src_path)
    convert_docx(src_path, out_doc, 0.69, "at-least", True, report, numbering_ids)
    out_doc.save(output_path)

    converted = Document(output_path)
    assert converted.paragraphs[0].style.name == "1.1一级列项-编号"
    assert converted.paragraphs[0].text == "第一项说明"
    assert not paragraph_has_graphics(converted.paragraphs[0])
    assert converted.paragraphs[1].text == ""
    assert paragraph_has_graphics(converted.paragraphs[1])
    assert report["semantic_object_splits"] == [
        {"text": "a) 第一项说明", "role": "list"}
    ]
    assert report["media_relationships_preserved"] == 1


def test_media_scan_ignores_zip_directory_entries(tmp_path):
    docx_path = tmp_path / "source.docx"
    with zipfile.ZipFile(docx_path, "w") as zf:
        zf.writestr("word/media/", "")
        zf.writestr("word/media/image1.png", b"png")
        zf.writestr("word/document.xml", "<w:document><w:drawing/></w:document>")

    counts = scan_non_text_objects(docx_path)

    assert counts["media_files"] == 1
    assert counts["drawings"] == 1
