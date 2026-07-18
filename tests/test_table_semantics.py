from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


SCRIPTS = Path(__file__).resolve().parents[1] / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from document_model import caption_block, new_document_model, table_block
from llm_enhancer import _collect_caption_targets
from model_normalization import normalize_document_model_simple
from table_formatting import normalize_table
from table_semantics import (
    audit_model_table_semantics,
    classify_docx_table,
    table_caption_eligible,
)
from template_profile import load_template_profile


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "assets" / "wx_template.docx"


def _model_with_table(table_type: str, *, header_rows: int = 0) -> dict:
    model = new_document_model("fixture.docx", "docx", "test")
    model["document"]["blocks"].append(
        table_block(
            "b1",
            table_type,
            [[{"text": "说明内容", "cell_role": "header"}]],
            header_rows=header_rows,
        )
    )
    return model


def test_single_cell_natural_language_is_callout() -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "边界说明：该内容用于突出一段完整说明。"

    result = classify_docx_table(table)

    assert result.table_type == "callout"
    assert result.header_rows == 0
    assert result.visual_cell_count == 1
    assert result.caption_eligible is False
    assert "natural_language_content" in result.evidence


def test_single_cell_code_payload_is_code_sample() -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = 'POST /api/auth\nContent-Type: application/json\n{"user":"demo"}'

    result = classify_docx_table(table)

    assert result.table_type == "code_sample"
    assert result.header_rows == 0
    assert result.caption_eligible is False


def test_structured_interface_catalog_is_data_table() -> None:
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    values = [
        ("序号", "接口名称", "接口能力"),
        ("1", "SDK_Init", "SDK初始化"),
        ("2", "Capture", "采集指纹"),
    ]
    for row, row_values in zip(table.rows, values):
        for cell, value in zip(row.cells, row_values):
            cell.text = value

    result = classify_docx_table(table, multi_cell_code_sample=True)

    assert result.table_type == "data"
    assert result.caption_eligible is True
    assert "legacy_header_signal_rejected_without_payload" in result.evidence


def test_multi_cell_payload_example_remains_code_sample() -> None:
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "请求示例"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = 'POST /api/auth\nContent-Type: application/json\n{"user":"demo","fingerprint":"abcdef"}'
    table.cell(1, 1).text = "登录请求"

    result = classify_docx_table(table, multi_cell_code_sample=True)

    assert result.table_type == "code_sample"
    assert result.caption_eligible is False


def test_single_merged_visual_cell_uses_unique_tc_count() -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "合并后的说明框"

    result = classify_docx_table(table)

    assert result.visual_cell_count == 1
    assert result.table_type == "callout"


def test_single_cell_nested_table_is_layout() -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).add_table(rows=1, cols=1)

    result = classify_docx_table(table)

    assert result.table_type == "layout"
    assert result.caption_eligible is False


def test_only_data_tables_receive_automatic_captions() -> None:
    callout = _model_with_table("callout", header_rows=1)
    normalized_callout = normalize_document_model_simple(callout, {})
    callout_blocks = normalized_callout["document"]["blocks"]

    assert [block["block_type"] for block in callout_blocks] == ["table"]
    assert callout_blocks[0]["header_rows"] == 0
    assert callout_blocks[0]["rows"][0][0]["cell_role"] == "body"

    data = _model_with_table("data", header_rows=1)
    normalized_data = normalize_document_model_simple(data, {})
    data_blocks = normalized_data["document"]["blocks"]

    assert [block["block_type"] for block in data_blocks] == ["caption", "table"]
    assert data_blocks[0]["_auto_generated"] is True


def test_authored_caption_for_callout_is_preserved_and_valid() -> None:
    model = _model_with_table("callout")
    model["document"]["blocks"].insert(
        0,
        caption_block("c1", "作者题注", "table"),
    )

    normalized = normalize_document_model_simple(model, {})
    audit = audit_model_table_semantics(normalized)

    assert [block["block_type"] for block in normalized["document"]["blocks"]] == ["caption", "table"]
    assert normalized["document"]["blocks"][0].get("_auto_generated") is None
    assert audit["passed"] is True


def test_enhanced_caption_collector_rejects_callout_targets() -> None:
    model = _model_with_table("callout")
    generated = caption_block("c1", "", "table")
    generated["_auto_generated"] = True
    model["document"]["blocks"].insert(0, generated)

    targets = _collect_caption_targets(model)
    audit = audit_model_table_semantics(model)

    assert targets == []
    assert audit["passed"] is False
    assert audit["issues"][0]["type"] == "auto_caption_on_ineligible_table"


def test_callout_format_policy_is_left_aligned() -> None:
    profile = load_template_profile(TEMPLATE)
    doc = Document(TEMPLATE)
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "说明内容"

    normalize_table(table, profile, 0.69, "at-least", role="callout")

    assert table.cell(0, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert table_caption_eligible("callout") is False
