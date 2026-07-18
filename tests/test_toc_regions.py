from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SCRIPTS = Path(__file__).resolve().parents[1] / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from llm_enhancer import enhance_document_model, normalize_mode, validate_patch
from toc_detector import detect_toc_regions, selected_source_positions


def _manual_toc_document() -> Document:
    doc = Document()
    doc.add_paragraph("系统建设方案")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("目  录")
    doc.add_paragraph("1 项目概述")
    doc.add_paragraph("2 总体设计")
    doc.add_paragraph("3 实施计划")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("1 项目概述", style="Heading 1")
    doc.add_paragraph("正文内容。")
    doc.add_paragraph("2 总体设计", style="Heading 1")
    doc.add_paragraph("正文内容。")
    doc.add_paragraph("3 实施计划", style="Heading 1")
    return doc


def _ambiguous_context() -> tuple[dict, dict]:
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("1 项目概述")
    doc.add_paragraph("2 总体设计")
    doc.add_paragraph("说明文字，没有明确分页边界。")
    report: dict = {}
    context = detect_toc_regions(doc, report)
    assert context["source_context"]["toc_status"] == "ambiguous"
    return context, report


def test_manual_toc_is_selected_by_rules() -> None:
    report: dict = {}
    context = detect_toc_regions(_manual_toc_document(), report)

    assert report["source_toc"]["status"] == "detected"
    assert report["source_toc"]["method"] == "rules"
    assert selected_source_positions(context)
    candidate = context["document"]["blocks"][0]
    assert candidate["entry_count"] == 3
    assert "entries_repeat_as_headings" in candidate["evidence"]


def test_no_toc_keeps_source_untouched() -> None:
    doc = Document()
    doc.add_paragraph("项目概述", style="Heading 1")
    doc.add_paragraph("正文内容。")
    report: dict = {}

    context = detect_toc_regions(doc, report)

    assert report["source_toc"]["status"] == "no_candidate"
    assert selected_source_positions(context) == set()


def test_body_directory_heading_is_not_removed_by_rules() -> None:
    doc = Document()
    doc.add_paragraph("项目概述", style="Heading 1")
    doc.add_paragraph("目录")
    doc.add_paragraph("1 文件管理")
    doc.add_paragraph("2 权限管理")
    report: dict = {}

    context = detect_toc_regions(doc, report)

    assert report["source_toc"]["status"] != "detected"
    assert selected_source_positions(context) == set()


def test_word_toc_field_is_strong_evidence() -> None:
    doc = Document()
    paragraph = doc.add_paragraph()
    run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.text = 'TOC \\o "1-3" \\h \\z'
    run.append(instruction)
    paragraph._p.append(run)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("项目概述", style="Heading 1")
    report: dict = {}

    context = detect_toc_regions(doc, report)

    assert report["source_toc"]["status"] == "detected"
    assert "toc_field" in context["document"]["blocks"][0]["evidence"]


def test_toc_style_is_detected() -> None:
    doc = Document()
    toc_style = doc.styles.add_style("TOC Custom", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("1 项目概述", style=toc_style)
    doc.add_paragraph("2 总体设计", style=toc_style)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("1 项目概述", style="Heading 1")
    doc.add_paragraph("2 总体设计", style="Heading 1")
    report: dict = {}

    context = detect_toc_regions(doc, report)

    assert report["source_toc"]["status"] == "detected"
    assert "toc_style" in context["document"]["blocks"][0]["evidence"]


def test_table_based_manual_toc_is_detected() -> None:
    doc = Document()
    doc.add_paragraph("目录")
    table = doc.add_table(rows=3, cols=1)
    table.cell(0, 0).text = "1 项目概述"
    table.cell(1, 0).text = "2 总体设计"
    table.cell(2, 0).text = "3 实施计划"
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("1 项目概述", style="Heading 1")
    doc.add_paragraph("2 总体设计", style="Heading 1")
    doc.add_paragraph("3 实施计划", style="Heading 1")
    report: dict = {}

    context = detect_toc_regions(doc, report)

    assert report["source_toc"]["status"] == "detected"
    assert context["document"]["blocks"][0]["entry_count"] == 3


def test_llm_capability_selects_only_existing_candidate() -> None:
    context, report = _ambiguous_context()
    candidate = context["document"]["blocks"][0]
    response = json.dumps({
        "schema_version": "1.0",
        "phase": "S",
        "decisions": [{
            "block_id": candidate["id"],
            "operation": "exclude_toc_region",
            "to": {
                "candidate_id": candidate["candidate_id"],
                "start_source_position": candidate["start_source_position"],
                "end_source_position": candidate["end_source_position"],
            },
            "confidence": 0.91,
            "reason": "目录候选得到确认",
        }],
    }, ensure_ascii=False)

    enhance_document_model(
        context,
        report,
        phase="S",
        llm_call=lambda prompt: response,
    )

    assert report["source_toc"]["status"] == "detected"
    assert report["source_toc"]["method"] == "llm"
    assert selected_source_positions(context)


def test_llm_cannot_expand_candidate_boundaries() -> None:
    context, _ = _ambiguous_context()
    candidate = context["document"]["blocks"][0]
    patch = {
        "schema_version": "1.0",
        "phase": "S",
        "decisions": [{
            "block_id": candidate["id"],
            "operation": "exclude_toc_region",
            "to": {
                "candidate_id": candidate["candidate_id"],
                "start_source_position": candidate["start_source_position"],
                "end_source_position": candidate["end_source_position"] + 1,
            },
            "confidence": 0.95,
        }],
    }

    errors = validate_patch(patch, context, frozenset({"exclude_toc_region"}))

    assert errors
    assert any("cannot change" in error.get("message", "") for error in errors)


def test_low_confidence_llm_selection_is_skipped() -> None:
    context, report = _ambiguous_context()
    candidate = context["document"]["blocks"][0]
    response = json.dumps({
        "schema_version": "1.0",
        "phase": "S",
        "decisions": [{
            "block_id": candidate["id"],
            "operation": "exclude_toc_region",
            "to": {
                "candidate_id": candidate["candidate_id"],
                "start_source_position": candidate["start_source_position"],
                "end_source_position": candidate["end_source_position"],
            },
            "confidence": 0.50,
            "reason": "证据不足",
        }],
    }, ensure_ascii=False)

    enhance_document_model(
        context,
        report,
        phase="S",
        llm_call=lambda prompt: response,
    )

    assert report["source_toc"]["status"] == "ambiguous"
    assert selected_source_positions(context) == set()
    skipped = report["llm_enhancer"]["skipped"]
    assert any(item.get("skip_reason") == "low_confidence" for item in skipped)


def test_all_mode_includes_source_capability() -> None:
    assert normalize_mode("all") == "sab"
