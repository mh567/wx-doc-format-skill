"""Unit tests for llm_enhancer.py.

Covers JSON extraction, patch validation, patch application, suspicion
scoring, should-enhance decisions, and the core enhance_document_model
entry point — all without actual LLM calls (fake_llm is used).
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

# ── Path setup ──────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

# ruff: noqa: E402 — imports after sys.path manipulation

from document_model import new_document_model
from llm_enhancer import (
    _append_phase_summary,
    _apply_token_budget,
    _build_phase_a_prompt,
    _build_phase_b_prompt,
    _collect_phase_b_sections,
    _collect_phase_c_table_groups,
    _collect_suspicious_sections,
    _iter_batches,
    _prevalidate_patch_schema,
    _record_phase_metric,
    _resolve_llm_call,
    ALLOWED_OPS_BY_PHASE,
    LOW_CONFIDENCE_THRESHOLD,
    PATCH_SCHEMA_VERSION,
    PHASE_B_SECTION_BATCH_SIZE,
    PHASE_C_TABLE_BATCH_SIZE,
    apply_patch_to_model,
    build_role_overrides_from_docx,
    compute_suspicion_score,
    enhance_document_model,
    extract_json_object,
    should_enhance,
    validate_patch,
)


# ═════════════════════════════════════════════════════════════════════
#  Helpers
# ═════════════════════════════════════════════════════════════════════

def fake_llm(patch_json: str):
    """Return an llm_call callable that always returns *patch_json*."""
    def _call(prompt: str) -> str:
        return patch_json
    return _call


def make_minimal_model(blocks=None):
    """Return a minimal document model, optionally with *blocks*."""
    model = new_document_model("test.md", "md", "test")
    if blocks is not None:
        model["document"]["blocks"] = blocks
    return model


def _block(bid: str, block_type: str = "body", **kwargs):
    """Build a minimal block dict."""
    b = {"id": bid, "block_type": block_type, "text": kwargs.pop("text", "test content")}
    b.update(kwargs)
    return b


# ═════════════════════════════════════════════════════════════════════
#  extract_json_object
# ═════════════════════════════════════════════════════════════════════

class TestExtractJsonObject:
    def test_extract_json_object_basic(self):
        result = extract_json_object('{"key": "value"}')
        assert result == {"key": "value"}

    def test_extract_json_object_with_markdown_wrapper(self):
        raw = '这是一段说明文字。\n\n```json\n{"name": "测试", "count": 3}\n```'
        result = extract_json_object(raw)
        assert result == {"name": "测试", "count": 3}

    def test_extract_json_object_empty_string_returns_none(self):
        assert extract_json_object("") is None

    def test_extract_json_object_invalid_json_returns_none(self):
        assert extract_json_object("{broken json}") is None

    def test_extract_json_object_nested_braces(self):
        raw = '前文 {"a": {"b": [1, 2]}} 后文'
        result = extract_json_object(raw)
        assert result == {"a": {"b": [1, 2]}}


# ═════════════════════════════════════════════════════════════════════
#  validate_patch
# ═════════════════════════════════════════════════════════════════════

class TestValidatePatch:
    def test_validate_patch_valid_phase_a(self):
        blocks = [_block("b0001"), _block("b0002")]
        model = make_minimal_model(blocks)
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "A",
            "decisions": [{
                "block_id": "b0001",
                "operation": "retype",
                "from": {"block_type": "body"},
                "to": {"block_type": "list_item", "level": 0,
                        "list_type": "lower_letter_paren"},
                "confidence": 0.85,
                "reason": "consecutive_functional_points",
            }],
        }
        errors = validate_patch(patch, model, ALLOWED_OPS_BY_PHASE["A"])
        assert errors == []

    def test_validate_patch_invalid_schema_version(self):
        model = make_minimal_model()
        patch = {"schema_version": "0.9", "phase": "A", "decisions": []}
        errors = validate_patch(patch, model, {"retype"})
        assert any("schema_version" in str(e) for e in errors)

    def test_validate_patch_unknown_block_id(self):
        model = make_minimal_model([_block("b0001")])
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "A",
            "decisions": [{
                "block_id": "b9999",
                "operation": "retype",
                "to": {"block_type": "body"},
                "confidence": 0.95,
            }],
        }
        errors = validate_patch(patch, model, {"retype"})
        assert len(errors) == 1
        assert "b9999" in errors[0].get("message", "")

    def test_validate_patch_disallowed_operation(self):
        model = make_minimal_model([_block("b0001")])
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "A",
            "decisions": [{
                "block_id": "b0001",
                "operation": "set_table_type",
                "to": {"table_type": "data"},
                "confidence": 0.95,
            }],
        }
        # Phase A only allows "retype"
        errors = validate_patch(patch, model, ALLOWED_OPS_BY_PHASE["A"])
        assert any("set_table_type" in str(e) for e in errors)

    def test_validate_patch_empty_decisions_is_valid(self):
        model = make_minimal_model()
        patch = {"schema_version": PATCH_SCHEMA_VERSION, "phase": "A", "decisions": []}
        errors = validate_patch(patch, model, {"retype"})
        assert errors == []

    def test_validate_patch_bad_target_block_type(self):
        model = make_minimal_model([_block("b0001")])
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "A",
            "decisions": [{
                "block_id": "b0001",
                "operation": "retype",
                "to": {"block_type": "nonexistent_type"},
                "confidence": 0.85,
            }],
        }
        errors = validate_patch(patch, model, {"retype"})
        assert any("block_type" in str(e) for e in errors)


# ═════════════════════════════════════════════════════════════════════
#  apply_patch_to_model
# ═════════════════════════════════════════════════════════════════════

class TestApplyPatch:
    def test_apply_retype_body_to_list_item(self):
        blocks = [_block("b0001", "body", text="项目一")]
        model = make_minimal_model(blocks)
        report: dict = {}
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "A",
            "decisions": [{
                "block_id": "b0001",
                "operation": "retype",
                "from": {"block_type": "body"},
                "to": {"block_type": "list_item", "level": 0,
                        "list_type": "lower_letter_paren"},
                "confidence": 0.85,
                "reason": "consecutive_functional_points",
            }],
        }
        apply_patch_to_model(model, patch, report)
        block = model["document"]["blocks"][0]
        assert block["block_type"] == "list_item"
        assert block.get("level") == 0
        assert block.get("list_type") == "lower_letter_paren"
        assert len(report["llm_enhancer"]["applied"]) == 1

    def test_apply_retype_body_to_heading(self):
        blocks = [_block("b0001", "body", text="重要章节")]
        model = make_minimal_model(blocks)
        report: dict = {}
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "A",
            "decisions": [{
                "block_id": "b0001",
                "operation": "retype",
                "from": {"block_type": "body"},
                "to": {"block_type": "heading", "level": 2},
                "confidence": 0.90,
                "reason": "visual_format_suggests_heading",
            }],
        }
        apply_patch_to_model(model, patch, report)
        block = model["document"]["blocks"][0]
        assert block["block_type"] == "heading"
        assert block.get("level") == 2

    def test_apply_retype_heading_to_body(self):
        blocks = [_block("b0001", "heading", text="版本: V2.1", level=2)]
        model = make_minimal_model(blocks)
        report: dict = {}
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "B",
            "decisions": [{
                "block_id": "b0001",
                "operation": "retype",
                "from": {"block_type": "heading", "level": 2},
                "to": {"block_type": "body"},
                "confidence": 0.90,
                "reason": "cover_metadata",
            }],
        }
        apply_patch_to_model(model, patch, report)
        assert model["document"]["blocks"][0]["block_type"] == "body"

    def test_apply_adjust_level(self):
        blocks = [_block("b0001", "heading", text="章节", level=3)]
        model = make_minimal_model(blocks)
        report: dict = {}
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "B",
            "decisions": [{
                "block_id": "b0001",
                "operation": "adjust_level",
                "to": {"level": 2},
                "confidence": 0.90,
                "reason": "section_continuity",
            }],
        }
        apply_patch_to_model(model, patch, report)
        assert model["document"]["blocks"][0].get("level") == 2

    def test_apply_set_table_type(self):
        blocks = [_block("b0001", "table", table_type="unknown", rows=[])]
        model = make_minimal_model(blocks)
        report: dict = {}
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "C",
            "decisions": [{
                "block_id": "b0001",
                "operation": "set_table_type",
                "to": {"table_type": "data"},
                "confidence": 0.95,
                "reason": "data_table_headers",
            }],
        }
        apply_patch_to_model(model, patch, report)
        assert model["document"]["blocks"][0].get("table_type") == "data"

    def test_apply_set_caption_text(self):
        blocks = [_block("b0001", "caption", text="", _auto_generated=True)]
        model = make_minimal_model(blocks)
        report: dict = {}
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "C",
            "decisions": [{
                "block_id": "b0001",
                "operation": "set_caption_text",
                "to": {"text": "系统功能模块列表"},
                "confidence": 0.85,
                "reason": "caption_text_generated",
            }],
        }
        apply_patch_to_model(model, patch, report)
        assert model["document"]["blocks"][0].get("text") == "系统功能模块列表"

    def test_low_confidence_patch_is_skipped(self):
        blocks = [_block("b0001", "body")]
        model = make_minimal_model(blocks)
        report: dict = {}
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "A",
            "decisions": [{
                "block_id": "b0001",
                "operation": "retype",
                "to": {"block_type": "list_item"},
                "confidence": LOW_CONFIDENCE_THRESHOLD - 0.1,
                "reason": "test",
            }],
        }
        apply_patch_to_model(model, patch, report)
        assert not report["llm_enhancer"]["applied"]
        assert len(report["llm_enhancer"]["skipped"]) == 1
        assert report["llm_enhancer"]["skipped"][0]["skip_reason"] == "low_confidence"

    def test_unknown_block_id_is_error(self):
        blocks = [_block("b0001", "body")]
        model = make_minimal_model(blocks)
        report: dict = {}
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "A",
            "decisions": [{
                "block_id": "b9999",
                "operation": "retype",
                "to": {"block_type": "body"},
                "confidence": 0.95,
            }],
        }
        apply_patch_to_model(model, patch, report)
        assert not report["llm_enhancer"]["applied"]
        assert len(report["llm_enhancer"]["errors"]) == 1


# ═════════════════════════════════════════════════════════════════════
#  suspicion_score
# ═════════════════════════════════════════════════════════════════════

class TestSuspicionScore:
    def test_suspicion_score_zero_for_clean_document(self):
        report = {
            "source_document_model_summary": {
                "block_count": 10,
                "block_counts": {"body": 8, "heading": 2, "list_item": 0},
            },
            "parse_report": {},
        }
        assert compute_suspicion_score(report) == 0.0

    def test_suspicion_score_nonzero_for_ambiguous_document(self):
        report = {
            "source_document_model_summary": {
                "block_count": 10,
                "block_counts": {"body": 8, "heading": 2, "list_item": 0},
            },
            "parse_report": {
                "suspect_visual_headings": 4,
                "ambiguous_short_paragraphs": 3,
            },
        }
        score = compute_suspicion_score(report)
        assert 0 < score <= 1.0

    def test_suspicion_score_range_0_to_1(self):
        """Even with exaggerated signals the score must stay in [0, 1]."""
        report = {
            "source_document_model_summary": {
                "block_count": 5,
                "block_counts": {"body": 3, "heading": 2, "list_item": 0},
            },
            "parse_report": {
                "suspect_visual_headings": 100,
                "ambiguous_short_paragraphs": 100,
                "inferred_headings": 100,
                "inferred_lists": 100,
                "unstyled_paragraphs": 100,
            },
        }
        score = compute_suspicion_score(report)
        assert 0.0 <= score <= 1.0


# ═════════════════════════════════════════════════════════════════════
#  should_enhance
# ═════════════════════════════════════════════════════════════════════

class TestShouldEnhance:
    def test_should_enhance_off_returns_false(self):
        assert not should_enhance({}, "A", "off")
        assert not should_enhance({}, "B", "off")
        assert not should_enhance({}, "C", "off")

    def test_should_enhance_abc_returns_true_for_all_phases(self):
        assert should_enhance({}, "A", "abc") is True
        assert should_enhance({}, "B", "abc") is True
        assert should_enhance({}, "C", "abc") is True

    def test_should_enhance_a_returns_true_only_for_phase_a(self):
        assert should_enhance({}, "A", "a") is True
        assert should_enhance({}, "B", "a") is False
        assert should_enhance({}, "C", "a") is False

    def test_should_enhance_ab_includes_a_and_b(self):
        assert should_enhance({}, "A", "ab") is True
        assert should_enhance({}, "B", "ab") is True
        assert should_enhance({}, "C", "ab") is False

    def test_should_enhance_force_modes_work(self):
        assert should_enhance({}, "A", "force-a") is True
        assert should_enhance({}, "B", "force-a") is False
        assert should_enhance({}, "C", "force-abc") is True

    def test_should_enhance_auto_low_suspicion(self):
        report = {
            "source_document_model_summary": {
                "block_count": 50,
                "block_counts": {"body": 40, "heading": 10, "list_item": 0},
            },
            "parse_report": {},
        }
        assert not should_enhance(report, "A", "auto")

    def test_should_enhance_auto_high_suspicion(self):
        report = {
            "source_document_model_summary": {
                "block_count": 10,
                "block_counts": {"body": 8, "heading": 2, "list_item": 0},
            },
            "parse_report": {
                "suspect_visual_headings": 5,
                "ambiguous_short_paragraphs": 3,
                "inferred_headings": 2,
            },
            "template_profile": {"resolved_styles": {"heading_1": "heading 1"}},
        }
        # Suspicion score should exceed 0.15
        assert should_enhance(report, "A", "auto")


# ═════════════════════════════════════════════════════════════════════
#  enhance_document_model
# ═════════════════════════════════════════════════════════════════════

class TestEnhanceDocumentModel:
    def test_enhance_no_llm_call_returns_original(self):
        model = make_minimal_model([_block("b0001", "body")])
        report: dict = {}
        result = enhance_document_model(model, report, phase="A", llm_call=None)
        assert result is model  # Same object

    def test_enhance_unknown_phase_records_error(self):
        model = make_minimal_model()
        report: dict = {}
        result = enhance_document_model(model, report, phase="X",
                                        llm_call=fake_llm("{}"))
        assert "llm_enhancer" in report
        assert len(report["llm_enhancer"]["errors"]) >= 1
        assert result is model  # unchanged

    def test_enhance_applies_valid_patch(self):
        blocks = [
            _block("b0001", "body", text="第一项"),
            _block("b0002", "body", text="第二项"),
        ]
        model = make_minimal_model(blocks)
        report: dict = {}
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "A",
            "decisions": [{
                "block_id": "b0001",
                "operation": "retype",
                "from": {"block_type": "body"},
                "to": {"block_type": "list_item", "level": 0,
                        "list_type": "lower_letter_paren"},
                "confidence": 0.85,
                "reason": "consecutive_functional_points",
            }],
        }
        result = enhance_document_model(
            model, report, phase="A",
            llm_call=fake_llm(json.dumps(patch, ensure_ascii=False)),
        )
        assert result["document"]["blocks"][0]["block_type"] == "list_item"
        assert len(report["llm_enhancer"]["applied"]) == 1

    def test_enhance_bad_json_response_records_error(self):
        model = make_minimal_model([_block("b0001", "body")])
        report: dict = {}
        enhance_document_model(
            model, report, phase="A",
            llm_call=fake_llm("这不是 JSON"),
        )
        assert len(report["llm_enhancer"]["errors"]) >= 1

    def test_enhance_empty_llm_response_records_error(self):
        model = make_minimal_model([_block("b0001", "body")])
        report: dict = {}
        enhance_document_model(
            model, report, phase="A",
            llm_call=fake_llm(""),
        )
        assert len(report["llm_enhancer"]["errors"]) >= 1

    def test_enhance_with_hint_injected(self):
        model = make_minimal_model([_block("b0001", "body")])
        report: dict = {}
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "A",
            "decisions": [],
        }
        enhance_document_model(
            model, report, phase="A",
            llm_call=fake_llm(json.dumps(patch)),
            hint="注意功能点列表识别",
        )
        assert report["llm_enhancer"].get("original_hint") == "注意功能点列表识别"

    def test_enhance_phase_c_allows_set_caption_text(self):
        """Phase C now includes C1 capability (merged)."""
        blocks = [_block("b0001", "caption", text="", _auto_generated=True)]
        model = make_minimal_model(blocks)
        report: dict = {}
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "C",
            "decisions": [{
                "block_id": "b0001",
                "operation": "set_caption_text",
                "to": {"text": "系统功能表"},
                "confidence": 0.85,
                "reason": "caption_text_generated",
            }],
        }
        result = enhance_document_model(
            model, report, phase="C",
            llm_call=fake_llm(json.dumps(patch, ensure_ascii=False)),
        )
        assert result["document"]["blocks"][0]["text"] == "系统功能表"
        assert len(report["llm_enhancer"]["applied"]) == 1


# ═════════════════════════════════════════════════════════════════════
#  _collect_suspicious_sections
# ═════════════════════════════════════════════════════════════════════

class TestCollectSuspiciousSections:
    def test_clean_document_returns_all_as_one_section(self):
        blocks = [
            _block("b0001", "heading", text="第一章"),
            _block("b0002", "body", text="这是正常的正文段落内容。"),
        ]
        model = make_minimal_model(blocks)
        report = {"parse_report": {}}
        sections = _collect_suspicious_sections(model, report)
        assert len(sections) >= 1

    def test_short_body_paragraph_triggers_section(self):
        blocks = [
            _block("b0001", "heading", text="第一章"),
            _block("b0002", "body", text="简短歧义段"),
            _block("b0003", "body", text="这是正常的长正文段落，用于测试目标段落的筛选逻辑。"),
        ]
        model = make_minimal_model(blocks)
        report = {
            "parse_report": {
                "ambiguous_short_paragraphs": 1,
            },
        }
        sections = _collect_suspicious_sections(model, report)
        # The section with the short paragraph should be included
        assert len(sections) >= 1

    def test_inferred_heading_triggers_section(self):
        blocks = [
            _block("b0001", "heading", text="第一章"),
            _block("b0002", "body", text="疑似标题", _inferred=True),
        ]
        model = make_minimal_model(blocks)
        report = {
            "parse_report": {
                "inferred_headings": 1,
            },
        }
        sections = _collect_suspicious_sections(model, report)
        assert len(sections) >= 1
        assert sections[0].get("heading_text") == "第一章"

    def test_no_blocks_returns_empty(self):
        model = make_minimal_model([])
        report = {"parse_report": {}}
        sections = _collect_suspicious_sections(model, report)
        assert sections == []


# ═════════════════════════════════════════════════════════════════════
#  _apply_token_budget
# ═════════════════════════════════════════════════════════════════════

class TestApplyTokenBudget:
    def test_short_prompt_unchanged(self):
        prompt = "短文本"
        result = _apply_token_budget(prompt, max_chars=100)
        assert result == prompt

    def test_long_prompt_truncated(self):
        prompt = "A" * 20000
        result = _apply_token_budget(prompt, max_chars=12000)
        assert len(result) <= 12000 + 3  # Allow for "…" marker
        assert "…" in result

    def test_exact_budget_not_truncated(self):
        prompt = "B" * 100
        result = _apply_token_budget(prompt, max_chars=100)
        assert result == prompt


# ═════════════════════════════════════════════════════════════════════
#  _resolve_llm_call  (timeout)
# ═════════════════════════════════════════════════════════════════════

class TestResolveLlmCall:
    def test_none_returns_none(self):
        assert _resolve_llm_call("A", None) is None

    def test_fast_call_returns_result(self):
        def fast(prompt: str) -> str:
            return "ok"
        wrapped = _resolve_llm_call("A", fast)
        assert wrapped("test") == "ok"

    def test_slow_call_times_out(self):
        import time

        def slow(prompt: str) -> str:
            time.sleep(60)  # Must exceed the 30s timeout
            return "too late"

        wrapped = _resolve_llm_call("A", slow)
        start = time.time()
        try:
            wrapped("test")
            assert False, "Expected TimeoutError"
        except TimeoutError:
            elapsed = time.time() - start
            assert elapsed < 60, f"Timeout took too long: {elapsed:.1f}s"


# ═════════════════════════════════════════════════════════════════════
#  ALLOWED_OPS_BY_PHASE  (C1 is merged into C)
# ═════════════════════════════════════════════════════════════════════

class TestAllowedOps:
    def test_phase_c_includes_caption_text(self):
        """After C+C1 merge, Phase C should include set_caption_text."""
        assert "set_caption_text" in ALLOWED_OPS_BY_PHASE["C"]

    def test_no_c1_phase_in_allowed_ops(self):
        """C1 should no longer be a separate phase entry."""
        assert "C1" not in ALLOWED_OPS_BY_PHASE


# ═════════════════════════════════════════════════════════════════════
#  build_role_overrides_from_docx — global index mapping
# ═════════════════════════════════════════════════════════════════════

class TestBuildRoleOverrides:
    def test_build_role_overrides_uses_global_index(self):
        """Overrides must map to document-global indices, not section-local.

        Two sections each with ≥2 paragraphs; the fake LLM returns a
        decision for *local* block_id "0" in every section.  With the
        fix each section's first paragraph maps to its *global* index.
        """
        import io
        from docx import Document as _Document

        doc = _Document()
        doc.add_paragraph("Section 1", style="Heading 1")
        doc.add_paragraph("First para")    # global_index=1 (heading at 0)
        doc.add_paragraph("Second para")   # global_index=2
        doc.add_paragraph("Section 2", style="Heading 1")
        doc.add_paragraph("Third para")    # global_index=4 (heading at 3)
        doc.add_paragraph("Fourth para")   # global_index=5

        # Use doc directly — save/load may lose style information.
        src_doc = doc

        patch = {
            "schema_version": "1.0",
            "phase": "A",
            "decisions": [{
                "block_id": "0",
                "operation": "retype",
                "from": {"block_type": "body"},
                "to": {"block_type": "list_item", "level": 0,
                        "list_type": "lower_letter_paren"},
                "confidence": 0.85,
                "reason": "test",
            }],
        }
        _fake = lambda p: json.dumps(patch)

        overrides = build_role_overrides_from_docx(src_doc, True, llm_call=_fake)

        # With headings counted in global_index, first body paragraphs are at 1 and 4
        assert 1 in overrides, "Section 1 first para (global 1) missing"
        assert 4 in overrides, "Section 2 first para (global 4) missing"
        assert len(overrides) == 2, (
            f"Expected 2 overrides (global 1 and 4), got {len(overrides)} "
            f"keys={sorted(overrides.keys())}"
        )


# ═════════════════════════════════════════════════════════════════════
#  _collect_suspicious_sections — heading blocks in section block list
# ═════════════════════════════════════════════════════════════════════

class TestCollectSuspiciousSectionsHeadingBlocks:
    def test_phase_a_prompt_includes_heading_blocks(self):
        """Heading blocks must be included in section block lists."""
        blocks = [
            {"id": "b0001", "block_type": "heading", "text": "版本: V1.0",
             "level": 1},
            {"id": "b0002", "block_type": "body",
             "text": "This is a short ambiguous para"},
        ]
        model = make_minimal_model(blocks)
        report = {"parse_report": {"ambiguous_short_paragraphs": 1}}
        sections = _collect_suspicious_sections(model, report)

        assert len(sections) > 0
        found = any(
            any(b.get("id") == "b0001" for b in sec.get("blocks", []))
            for sec in sections
        )
        assert found, "Heading block b0001 should be in a section's blocks"

    def test_heading_block_id_in_suspicious_ids_when_inferred(self):
        """An inferred heading block should still be sent when its
        section is flagged."""
        blocks = [
            {"id": "b0001", "block_type": "heading", "text": "第一章",
             "_inferred": True, "level": 1},
            {"id": "b0002", "block_type": "body",
             "text": "This is a short ambiguous para"},
        ]
        model = make_minimal_model(blocks)
        report = {"parse_report": {"ambiguous_short_paragraphs": 1,
                                    "inferred_headings": 1}}
        sections = _collect_suspicious_sections(model, report)

        assert len(sections) > 0
        found = any(
            any(b.get("id") == "b0001" for b in sec.get("blocks", []))
            for sec in sections
        )
        assert found, ("Inferred heading b0001 should appear in "
                       "suspicious section blocks")


# ═════════════════════════════════════════════════════════════════════
#  should_enhance — auto-mode phase-aware gating
# ═════════════════════════════════════════════════════════════════════

class TestShouldEnhanceAutoModeGate:
    def _base_report(self) -> dict:
        return {
            "source_document_model_summary": {
                "block_count": 10,
                "block_counts": {"body": 8, "heading": 2, "list_item": 0},
            },
            "template_profile": {"path": "/fake"},
        }

    def test_auto_phase_a_requires_suspicion(self):
        """Phase A in auto mode should gate on suspicion score."""
        report = self._base_report()
        report["parse_report"] = {}
        assert not should_enhance(report, "A", "auto")

    def test_auto_skips_bc_when_phase_a_low_modification(self):
        """B/C in auto mode should skip when Phase A applied_rate < 0.05."""
        report = self._base_report()
        report["llm_enhancer"] = {"phase_a_applied_rate": 0.04}
        assert not should_enhance(report, "B", "auto")
        assert not should_enhance(report, "C", "auto")

    def test_auto_runs_bc_when_phase_a_high_modification(self):
        """B/C in auto mode should run when Phase A applied_rate >= 0.05."""
        report = self._base_report()
        report["llm_enhancer"] = {"phase_a_applied_rate": 0.05}
        assert should_enhance(report, "B", "auto")
        assert should_enhance(report, "C", "auto")

    def test_auto_phase_a_skips_short_document(self):
        """Phase A in auto mode skips documents with < 5 blocks."""
        report = self._base_report()
        report["source_document_model_summary"]["block_count"] = 3
        report["parse_report"] = {"ambiguous_short_paragraphs": 10}
        assert not should_enhance(report, "A", "auto")

    def test_auto_phase_a_skips_no_template(self):
        """Phase A in auto mode skips when no template_profile."""
        report = self._base_report()
        report.pop("template_profile", None)
        report["parse_report"] = {"ambiguous_short_paragraphs": 10}
        assert not should_enhance(report, "A", "auto")


# ═════════════════════════════════════════════════════════════════════
#  hint sanitization
# ═════════════════════════════════════════════════════════════════════

class TestHintSanitization:
    def test_hint_sanitization_removes_control_chars(self):
        """Control characters must be stripped from user hints."""
        model = make_minimal_model([_block("b0001", "body")])
        report: dict = {}
        hint = "注意\x00功能\x1b点列表\x7f识别"
        enhance_document_model(
            model, report, phase="A",
            llm_call=None,
            hint=hint,
        )
        enh = report["llm_enhancer"]
        sanitized = enh["sanitized_hint"]
        assert "功能" in sanitized
        assert "识别" in sanitized
        assert "\x00" not in sanitized
        assert "\x1b" not in sanitized
        assert "\x7f" not in sanitized
        assert enh.get("original_hint") == hint

    def test_hint_sanitization_preserves_clean_hint(self):
        """A clean hint should pass through mostly unchanged."""
        model = make_minimal_model([_block("b0001", "body")])
        report: dict = {}
        hint = "功能点列表识别 - 测试"
        enhance_document_model(
            model, report, phase="A",
            llm_call=None,
            hint=hint,
        )
        enh = report["llm_enhancer"]
        assert enh["sanitized_hint"] == hint
        assert "hint_sanitized" not in enh
        assert "hint_truncated" not in enh

    def test_hint_truncation_at_500_chars(self):
        """Hints longer than 500 chars should be truncated."""
        model = make_minimal_model([_block("b0001", "body")])
        report: dict = {}
        hint = "x" * 600
        enhance_document_model(
            model, report, phase="A",
            llm_call=None,
            hint=hint,
        )
        enh = report["llm_enhancer"]
        assert len(enh["sanitized_hint"]) <= 500
        assert enh.get("hint_truncated") is True


# ═════════════════════════════════════════════════════════════════════
#  New exports for batching, pre-validation, phase_metrics
# ═════════════════════════════════════════════════════════════════════

class TestNewExports:
    def test_imported_constants(self):
        assert PHASE_B_SECTION_BATCH_SIZE == 20
        assert PHASE_C_TABLE_BATCH_SIZE == 10

    def test_iter_batches(self):
        items = list(range(25))
        batches = _iter_batches(items, 10)
        assert len(batches) == 3
        assert batches[0] == [0, 1, 2, 3, 4, 5, 6, 7, 8, 9]
        assert batches[1] == [10, 11, 12, 13, 14, 15, 16, 17, 18, 19]
        assert batches[2] == [20, 21, 22, 23, 24]


class TestCollectPhaseBSections:
    def test_collects_sections_by_heading(self):
        blocks = [
            _block("b0001", "heading", text="摘要", level=1),
            _block("b0002", "body", text="正文内容"),
            _block("b0003", "heading", text="第一章", level=2),
            _block("b0004", "list_item", text="项目一"),
        ]
        model = make_minimal_model(blocks)
        sections = _collect_phase_b_sections(model)
        assert len(sections) == 2
        assert sections[0]["heading_text"] == "摘要"
        assert sections[1]["heading_text"] == "第一章"

    def test_no_heading_returns_one_section(self):
        blocks = [
            _block("b0001", "body", text="封面文字"),
        ]
        model = make_minimal_model(blocks)
        sections = _collect_phase_b_sections(model)
        assert len(sections) == 1
        assert sections[0]["heading_text"] == ""

    def test_empty_model_returns_empty_section(self):
        model = make_minimal_model([])
        sections = _collect_phase_b_sections(model)
        assert len(sections) >= 1


class TestCollectPhaseCTableGroups:
    def test_collects_table_with_heading_and_captions(self):
        blocks = [
            _block("b0001", "heading", text="数据表", level=2),
            _block("b0002", "caption", text="表1: 数据", caption_type="table"),
            _block("b0003", "table", table_type="data", rows=[]),
            _block("b0004", "caption", text="", _auto_generated=True),
        ]
        model = make_minimal_model(blocks)
        groups = _collect_phase_c_table_groups(model)
        assert len(groups) == 1
        assert groups[0]["table_id"] == "b0003"
        # Must include heading, preceding caption, table, following caption
        group_bids = {b.get("id") for b in groups[0]["blocks"]}
        assert "b0001" in group_bids  # heading
        assert "b0002" in group_bids  # preceding caption
        assert "b0003" in group_bids  # table
        assert "b0004" in group_bids  # following caption

    def test_no_table_returns_empty(self):
        model = make_minimal_model([_block("b0001", "body")])
        groups = _collect_phase_c_table_groups(model)
        assert groups == []


class TestPrevalidatePatchSchema:
    def test_valid_patch_passes(self):
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "B",
            "decisions": [{
                "block_id": "b0001",
                "operation": "retype",
                "to": {"block_type": "body"},
            }],
        }
        errors = _prevalidate_patch_schema(patch, "B", {"retype"})
        assert errors == []

    def test_invalid_schema_version(self):
        patch = {"schema_version": "0.9", "phase": "B", "decisions": []}
        errors = _prevalidate_patch_schema(patch, "B", {"retype"})
        assert any("schema_version" in str(e) for e in errors)

    def test_phase_mismatch(self):
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "C",
            "decisions": [],
        }
        errors = _prevalidate_patch_schema(patch, "B", {"retype"})
        assert any("phase" in str(e) for e in errors)

    def test_decisions_not_list(self):
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "B",
            "decisions": "not_a_list",
        }
        errors = _prevalidate_patch_schema(patch, "B", {"retype"})
        assert any("must be list" in str(e) for e in errors)

    def test_rejects_disallowed_operation(self):
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "B",
            "decisions": [{
                "block_id": "b0001",
                "operation": "set_table_type",
                "to": {"table_type": "data"},
            }],
        }
        errors = _prevalidate_patch_schema(patch, "B", {"retype"})
        assert any("operation not in allowed_ops" in str(e) for e in errors)

    def test_decision_not_dict(self):
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "B",
            "decisions": ["not_a_dict"],
        }
        errors = _prevalidate_patch_schema(patch, "B", {"retype"})
        assert any("must be object" in str(e) for e in errors)

    def test_empty_patch_rejected(self):
        errors = _prevalidate_patch_schema("not_a_dict", "B", {"retype"})
        assert any("must be object" in str(e) for e in errors)


class TestRecordPhaseMetric:
    def test_records_metric(self):
        enh: dict = {}
        _record_phase_metric(enh, {"phase": "B", "batch_index": 0})
        assert len(enh["phase_metrics"]) == 1
        assert enh["phase_metrics"][0]["phase"] == "B"
        assert enh["phase_metrics"][0]["batch_index"] == 0

    def test_appends_multiple_metrics(self):
        enh: dict = {}
        _record_phase_metric(enh, {"batch": 0})
        _record_phase_metric(enh, {"batch": 1})
        assert len(enh["phase_metrics"]) == 2


class TestAppendPhaseSummary:
    def test_first_call_uses_zero_prev(self):
        enh: dict = {"applied": [], "skipped": [], "errors": []}
        _append_phase_summary(enh, "B")
        assert len(enh["phase_summaries"]) == 1
        assert enh["phase_summaries"][0]["phase"] == "B"
        assert enh["_prev_applied"] == 0

    def test_second_call_diffs_correctly(self):
        enh: dict = {"applied": [{"block_id": "b1"}],
                      "skipped": [], "errors": []}
        _append_phase_summary(enh, "B")
        # Simulate a second call with more applied
        enh["applied"].append({"block_id": "b2"})
        _append_phase_summary(enh, "C")
        summaries = enh["phase_summaries"]
        assert len(summaries) == 2
        assert summaries[0]["phase"] == "B"
        assert summaries[1]["phase"] == "C"


class TestPhaseBMultiBatch:
    def test_phase_b_calls_llm_multiple_times_for_many_sections(self):
        """Phase B with >20 sections should call the LLM in multiple batches."""
        blocks = []
        for i in range(25):
            blocks.append(
                _block(f"b{i*2:04d}", "heading",
                       text=f"Section {i}", level=2)
            )
            blocks.append(
                _block(f"b{i*2+1:04d}", "list_item",
                       text=f"Item {i}")
            )
        model = make_minimal_model(blocks)
        report: dict = {}
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "B",
            "decisions": [],
        }

        call_count: list[int] = [0]

        def counting_llm(prompt: str) -> str:
            call_count[0] += 1
            return json.dumps(patch, ensure_ascii=False)

        enhance_document_model(
            model, report, phase="B",
            llm_call=counting_llm,
        )

        # 25 sections with batch size 20 → 2 batches
        assert call_count[0] == 2, (
            f"Expected 2 LLM calls for 25 sections, got {call_count[0]}"
        )


class TestPhaseCMultiBatch:
    def test_phase_c_calls_llm_multiple_times_for_many_tables(self):
        """Phase C with >10 table groups should call the LLM in multiple
        batches."""
        blocks = []
        for i in range(15):
            blocks.append(
                _block(f"b{i*3:04d}", "heading",
                       text=f"Table Section {i}", level=2)
            )
            blocks.append(
                _block(f"b{i*3+1:04d}", "table",
                       table_type="unknown", rows=[])
            )
            blocks.append(
                _block(f"b{i*3+2:04d}", "caption",
                       text="", _auto_generated=True)
            )
        model = make_minimal_model(blocks)
        report: dict = {}
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "C",
            "decisions": [],
        }

        call_count: list[int] = [0]

        def counting_llm(prompt: str) -> str:
            call_count[0] += 1
            return json.dumps(patch, ensure_ascii=False)

        enhance_document_model(
            model, report, phase="C",
            llm_call=counting_llm,
        )

        # 15 table groups with batch size 10 → 2 batches
        assert call_count[0] == 2, (
            f"Expected 2 LLM calls for 15 table groups, "
            f"got {call_count[0]}"
        )


class TestPhaseMetricsPresence:
    def test_phase_metrics_on_success(self):
        """phase_metrics should exist after a successful Phase B batch."""
        blocks = [
            _block("b0001", "heading", text="第一章", level=2),
            _block("b0002", "body", text="正文内容"),
        ]
        model = make_minimal_model(blocks)
        report: dict = {}
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "B",
            "decisions": [],
        }

        enhance_document_model(
            model, report, phase="B",
            llm_call=fake_llm(json.dumps(patch, ensure_ascii=False)),
        )

        enh = report["llm_enhancer"]
        assert "phase_metrics" in enh, "phase_metrics should exist"
        assert len(enh["phase_metrics"]) == 1
        pm = enh["phase_metrics"][0]
        assert pm["status"] == "ok"
        assert pm["batch_index"] == 0
        assert pm["batch_count"] == 1
        assert "estimated_tokens" in pm
        assert pm["estimated_tokens"] > 0
        assert "wall_time_sec" in pm
        assert pm["wall_time_sec"] >= 0

    def test_phase_metrics_on_error(self):
        """phase_metrics should exist even when LLM call fails."""
        blocks = [
            _block("b0001", "heading", text="第一章", level=2),
            _block("b0002", "body", text="正文内容"),
        ]
        model = make_minimal_model(blocks)
        report: dict = {}

        def failing_llm(prompt: str) -> str:
            raise RuntimeError("LLM failure")

        enhance_document_model(
            model, report, phase="B",
            llm_call=failing_llm,
        )

        enh = report["llm_enhancer"]
        assert "phase_metrics" in enh, (
            "phase_metrics should exist even on error"
        )
        assert len(enh["phase_metrics"]) == 1
        pm = enh["phase_metrics"][0]
        assert pm["status"] == "error"


class TestPhaseBWithBatchContext:
    def test_phase_b_batch_meta_present_in_prompts(self):
        """Phase B batch prompts should include batch meta and be split."""
        blocks = []
        for i in range(25):
            blocks.append(
                _block(f"b{i*2:04d}", "heading",
                       text=f"Section {i}", level=2)
            )
            blocks.append(
                _block(f"b{i*2+1:04d}", "list_item",
                       text=f"Item {i}")
            )
        model = make_minimal_model(blocks)
        report: dict = {}
        patch = {
            "schema_version": PATCH_SCHEMA_VERSION,
            "phase": "B",
            "decisions": [],
        }

        enhance_document_model(
            model, report, phase="B",
            llm_call=fake_llm(json.dumps(patch, ensure_ascii=False)),
        )

        enh = report["llm_enhancer"]
        prompts = enh.get("prompts", [])
        # 25 sections with batch size 20 → 2 prompts (batches)
        assert len(prompts) == 2, (
            f"Expected 2 prompt entries for 2 batches, got {len(prompts)}"
        )
        assert "batch" in prompts[0]
        assert prompts[0]["batch"] == 0
        assert prompts[1]["batch"] == 1


# ═════════════════════════════════════════════════════════════════════
#  Prompt keyword tests — first-principles refactor (TASK-FIX-004)
# ═════════════════════════════════════════════════════════════════════

class TestPhaseAPromptFirstPrinciples:
    """Phase A prompt must position the LLM as rule-result reviewer."""

    def make_model(self, text: str = "some content") -> dict:
        return {"document": {"blocks": [
            {"id": "b0001", "block_type": "body", "role": "body",
             "level": None, "list_type": None, "caption_type": None,
             "text": text},
        ]}}

    def test_task_description(self):
        """Phase A prompt should contain '规则结果审查器'."""
        model = self.make_model()
        prompt = _build_phase_a_prompt(model)
        assert "规则结果审查器" in prompt, (
            "Phase A should describe the LLM as 规则结果审查器"
        )

    def test_empty_decisions(self):
        """Phase A prompt should say empty decisions = rules are correct."""
        model = self.make_model()
        prompt = _build_phase_a_prompt(model)
        assert "空 decisions 表示规则全部正确" in prompt, (
            "Phase A should state that empty decisions = rules all correct"
        )

    def test_omit_block_means_accept(self):
        """Phase A prompt should say omitting a block = accepting rule result."""
        model = self.make_model()
        prompt = _build_phase_a_prompt(model)
        assert "省略" in prompt and "接受规则结果" in prompt, (
            "Phase A should say omitting a block = accepting rule result"
        )

    def test_default_rule_correct(self):
        """Phase A prompt should say default = rule correct."""
        model = self.make_model()
        prompt = _build_phase_a_prompt(model)
        assert "默认规则判断正确" in prompt

    def test_block_info_shows_rule_fields(self):
        """Phase A block listing should include role/level/list_type/caption_type."""
        model = self.make_model()
        prompt = _build_phase_a_prompt(model)
        # Lines should show role=... level=... list=... cap=...
        assert "role=" in prompt and "level=" in prompt
        assert "list=" in prompt and "cap=" in prompt

    def test_focus_areas_present(self):
        """Phase A should include focus areas for review."""
        model = self.make_model()
        prompt = _build_phase_a_prompt(model)
        assert "连续 body 功能点列表" in prompt
        assert "封面元信息误判" in prompt
        assert "题注" in prompt and "误判" in prompt
        assert "正文误判" in prompt


class TestPhaseBPromptFirstPrinciples:
    """Phase B prompt must position the LLM as heading-level anomaly reviewer."""

    def make_model(self) -> dict:
        return {"document": {"blocks": [
            {"id": "b0001", "block_type": "heading", "level": 1,
             "text": "文档标题"},
            {"id": "b0002", "block_type": "heading", "level": 2,
             "text": "章节一"},
        ]}}

    def test_task_description(self):
        """Phase B prompt should contain '标题层级异常审查器'."""
        model = self.make_model()
        prompt = _build_phase_b_prompt(model)
        assert "标题层级异常审查器" in prompt, (
            "Phase B should describe the LLM as 标题层级异常审查器"
        )

    def test_final_ast_emphasis(self):
        """Phase B prompt should say '规则处理后的最终 AST 摘要'."""
        model = self.make_model()
        prompt = _build_phase_b_prompt(model)
        assert "规则处理后" in prompt and "AST" in prompt

    def test_consecutive_same_level_is_legal(self):
        """Phase B should say consecutive same-level headings are legal."""
        model = self.make_model()
        prompt = _build_phase_b_prompt(model)
        assert "连续同级标题" in prompt

    def test_doc_title_to_h2_h3_legal(self):
        """Phase B should say H2/H3 after document title is legal."""
        model = self.make_model()
        prompt = _build_phase_b_prompt(model)
        assert "文档标题后直接进入" in prompt

    def test_empty_decisions_means_legal(self):
        """Phase B should say empty decisions = heading structure is legal."""
        model = self.make_model()
        prompt = _build_phase_b_prompt(model)
        assert "空 decisions 表示当前标题结构合法" in prompt

    def test_prefer_adjust_level_over_retype(self):
        """Phase B should prioritize adjust_level over retype."""
        model = self.make_model()
        prompt = _build_phase_b_prompt(model)
        assert "优先使用 adjust_level" in prompt

    def test_heading_shows_prev_level(self):
        """Phase B block listing should include previous heading level."""
        model = self.make_model()
        prompt = _build_phase_b_prompt(model)
        # Second heading (H2) should show prev=H1
        assert "prev=H1" in prompt

    def test_default_current_level_correct(self):
        """Phase B should say default current level is correct."""
        model = self.make_model()
        prompt = _build_phase_b_prompt(model)
        assert "默认当前层级正确" in prompt
