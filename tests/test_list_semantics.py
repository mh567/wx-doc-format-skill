from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SCRIPTS = Path(__file__).resolve().parents[1] / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from docx_pipeline import infer_docx_role
from list_detector import analyze_docx_lists
from model_normalization import normalize_document_model_simple
from text_utils import (
    heading_level_from_text,
    looks_like_list_item,
    strip_heading_marker,
    strip_list_marker,
)


def _add_single_level_abstract(doc: Document, abstract_id: int) -> None:
    root = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    level.append(lvl_text)
    abstract.append(level)
    root.insert(0, abstract)


def _add_num(doc: Document, num_id: int, abstract_id: int, start: int) -> None:
    root = doc.part.numbering_part.element
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    if start != 1:
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), str(start))
        override.append(start_override)
        num.append(override)
    root.append(num)


def _set_direct_num(paragraph, num_id: int) -> None:
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(num)
    paragraph._p.get_or_add_pPr().append(num_pr)


def _append_numbered(doc: Document, text: str, num_id: int, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(text, style=style)
    _set_direct_num(paragraph, num_id)


def test_hierarchical_heading_marker_is_atomic() -> None:
    for text in (
        "6.3 标准化注册时序流程",
        "6．4 指纹注册采集交互序列图",
        "6.5.1 后台入口与账号关联规则",
    ):
        assert heading_level_from_text(text) in {2, 3}
        assert looks_like_list_item(text) is False
        assert strip_list_marker(text) == text
        assert strip_heading_marker(text).startswith(("标准", "指纹", "后台"))


def test_single_level_decimal_markers_remain_lists() -> None:
    for text in ("1. 操作步骤", "2.配置管理", "3) 检查结果"):
        assert looks_like_list_item(text) is True
        assert strip_list_marker(text) in {"操作步骤", "配置管理", "检查结果"}


def test_decimal_like_body_text_is_not_promoted_to_heading() -> None:
    assert heading_level_from_text("1.2 协议用于以下场景。") is None
    assert heading_level_from_text("2026.07.19") is None
    assert heading_level_from_text("1.2%") is None


def test_ip_date_and_version_prefixes_are_not_promoted_to_heading() -> None:
    for text in (
        "192.168.1.1 地址配置为主节点",
        "10.0.0.1 IP地址",
        "1.2.3.4 地址规划",
        "2026.7.19 发布生产版本",
        "1.2.3 版本号说明如下",
    ):
        assert heading_level_from_text(text) is None

    assert heading_level_from_text("1.2.3 版本管理") == 3


def test_role_inference_prefers_hierarchical_heading_over_list_rules() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("6.3 标准化注册时序流程")
    report: dict = {}

    text, style, role = infer_docx_role(paragraph, True, report)

    assert text == "6.3 标准化注册时序流程"
    assert style == "Heading 2"
    assert role == "heading"
    assert report["inferred_headings"][0]["source"] == "docx-hierarchical-text"


def test_invalid_numbering_cannot_be_revived_by_list_style() -> None:
    doc = Document()
    style_name = "1.1一级列项-编号"
    doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    paragraph = doc.add_paragraph(
        "生物特征认证系统与IAM之间的逻辑接口如表1所示，具体接口见附录。",
        style=style_name,
    )
    _set_direct_num(paragraph, 0)
    report: dict = {}
    numbering = analyze_docx_lists(doc, report)[0]
    parse_report: dict = {}

    _, _, role = infer_docx_role(paragraph, True, parse_report, numbering)

    assert numbering["status"] == "ignored"
    assert "invalid_numbering_definition" in numbering["evidence"]
    assert role == "body"
    assert parse_report["suppressed_list_style_conflicts"]


def test_cross_num_id_instances_are_reconstructed_as_one_sequence() -> None:
    doc = Document()
    abstract_id = 100
    _add_single_level_abstract(doc, abstract_id)
    for num_id, start in ((100, 1), (101, 3), (102, 4), (103, 5)):
        _add_num(doc, num_id, abstract_id, start)
    _append_numbered(doc, "步骤1", 100)
    _append_numbered(doc, "步骤2", 100)
    _append_numbered(doc, "步骤3", 101)
    _append_numbered(doc, "步骤4", 102)
    for step in range(5, 14):
        _append_numbered(doc, f"步骤{step}", 103)
    report: dict = {}

    descriptors = analyze_docx_lists(doc, report)
    ordered = [descriptors[index] for index in sorted(descriptors)]

    assert len(ordered) == 13
    assert all(item["status"] == "detected" for item in ordered)
    assert all(item["group_size"] == 13 for item in ordered)
    assert [item["restart"] for item in ordered] == [True] + [False] * 12
    assert all(item["logical_num_ids"] == [100, 101, 102, 103] for item in ordered)
    assert report["source_lists"]["cross_num_id_group_count"] == 1


def test_cross_num_id_merge_requires_visible_sequence_continuity() -> None:
    doc = Document()
    abstract_id = 110
    _add_single_level_abstract(doc, abstract_id)
    _add_num(doc, 110, abstract_id, 1)
    _add_num(doc, 111, abstract_id, 1)
    _append_numbered(doc, "第一组步骤1", 110)
    _append_numbered(doc, "第一组步骤2", 110)
    _append_numbered(doc, "第二组步骤1", 111)
    _append_numbered(doc, "第二组步骤2", 111)
    report: dict = {}

    descriptors = analyze_docx_lists(doc, report)
    groups = [descriptors[index]["group_id"] for index in sorted(descriptors)]
    restarts = [descriptors[index]["restart"] for index in sorted(descriptors)]

    assert groups[0] == groups[1]
    assert groups[2] == groups[3]
    assert groups[0] != groups[2]
    assert restarts == [True, False, True, False]
    assert report["source_lists"]["cross_num_id_group_count"] == 0


def test_normalization_repairs_legacy_role_conflicts() -> None:
    source_model = {
        "document": {
            "blocks": [
                {
                    "id": "b1",
                    "block_type": "list_item",
                    "level": 1,
                    "list_type": "decimal_paren",
                    "text": "3 标准化注册时序流程",
                    "restart": True,
                    "source": {
                        "raw_text": "6.3 标准化注册时序流程",
                        "inferred_role": "list",
                    },
                },
                {
                    "id": "b2",
                    "block_type": "list_item",
                    "level": 0,
                    "list_type": "lower_letter_paren",
                    "text": "逻辑接口如表1所示。",
                    "restart": True,
                    "source": {
                        "raw_text": "逻辑接口如表1所示。",
                        "inferred_role": "list",
                        "numbering": {
                            "status": "ignored",
                            "evidence": ["invalid_numbering_definition"],
                        },
                    },
                },
            ]
        }
    }
    report: dict = {}

    normalized = normalize_document_model_simple(source_model, report)
    heading, body = normalized["document"]["blocks"]

    assert heading["block_type"] == "heading"
    assert heading["level"] == 2
    assert heading["text"] == "标准化注册时序流程"
    assert body["block_type"] == "body"
    assert body["text"] == "逻辑接口如表1所示。"
    repair_types = {item["type"] for item in report["model_normalization_repairs"]}
    assert "list_item_retyped_as_hierarchical_heading" in repair_types
    assert "invalid_list_style_retyped_as_body" in repair_types
