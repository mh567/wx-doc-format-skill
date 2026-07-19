from __future__ import annotations

import hashlib
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


SCRIPTS = Path(__file__).resolve().parents[1] / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from table_formatting import audit_document_tables, normalize_document_tables, normalize_table
from template_profile import load_template_profile
from reporting import add_risk_warnings


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "assets" / "wx_template.docx"


def _empty_template_document() -> Document:
    doc = Document(TEMPLATE)
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    return doc


def _xml_digest(table) -> str:
    return hashlib.sha256(table._tbl.xml.encode("utf-8")).hexdigest()


def test_table_normalization_clears_paragraph_overrides_and_is_idempotent() -> None:
    profile = load_template_profile(TEMPLATE)
    doc = _empty_template_document()
    table = doc.add_table(rows=2, cols=2)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.text = "内容"
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.left_indent = Pt(12)
    paragraph.runs[0].bold = True
    paragraph.runs[0].font.size = Pt(15)

    first = normalize_document_tables(doc, profile, 0.69, "at-least", table_roles=["data"])
    first_digest = _xml_digest(table)
    second = normalize_document_tables(doc, profile, 0.69, "at-least", table_roles=["data"])

    assert first
    assert second == []
    assert _xml_digest(table) == first_digest
    assert paragraph.style.name == "表正文"
    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraph._p.pPr.find(qn("w:spacing")) is None
    assert paragraph._p.pPr.find(qn("w:ind")) is None
    assert paragraph.runs[0]._r.rPr.find(qn("w:b")) is None
    assert table.rows[0].height_rule == WD_ROW_HEIGHT_RULE.AT_LEAST
    assert abs(table.rows[0].height.cm - 0.69) < 0.02


def test_table_normalization_handles_empty_merged_and_nested_cells() -> None:
    profile = load_template_profile(TEMPLATE)
    doc = _empty_template_document()
    table = doc.add_table(rows=2, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.paragraphs[0].paragraph_format.space_after = Pt(4)
    nested = table.cell(1, 0).add_table(rows=1, cols=1)
    nested_paragraph = nested.cell(0, 0).paragraphs[0]
    nested_paragraph.paragraph_format.space_before = Pt(3)

    normalize_table(table, profile, 0.69, "at-least", role="data")

    assert merged.paragraphs[0].style.name == "表正文"
    assert merged.paragraphs[0]._p.pPr.find(qn("w:spacing")) is None
    assert nested_paragraph.style.name == "表正文"
    assert nested_paragraph._p.pPr.find(qn("w:spacing")) is None


def test_code_sample_alignment_is_an_explicit_policy() -> None:
    profile = load_template_profile(TEMPLATE)
    doc = _empty_template_document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "GET /api"

    normalize_table(table, profile, 0.69, "at-least", role="code_sample")

    assert table.cell(0, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_table_contract_audit_rejects_direct_spacing_and_passes_after_normalization() -> None:
    profile = load_template_profile(TEMPLATE)
    doc = _empty_template_document()
    table = doc.add_table(rows=1, cols=1)
    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.text = "内容"
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "40")
    paragraph._p.get_or_add_pPr().append(spacing)

    before = audit_document_tables(doc, profile, 0.69, "at-least", table_roles=["data"])
    normalize_document_tables(doc, profile, 0.69, "at-least", table_roles=["data"])
    after = audit_document_tables(doc, profile, 0.69, "at-least", table_roles=["data"])

    assert before["passed"] is False
    assert before["paragraph_direct_format_issues"]
    assert after["passed"] is True


def test_template_profile_resolves_real_table_style_id() -> None:
    profile = load_template_profile(TEMPLATE)

    assert profile["table_style"]["name"] == "Table Grid"
    assert profile["table_style"]["style_id"] == "20"
    assert profile["table_style"]["tbl_look"] == {}


def test_failed_table_contract_becomes_a_risk_warning() -> None:
    report = {
        "non_text_objects": {},
        "media_relationships_preserved": 0,
        "risk_warnings": [],
        "audit": {
            "table_format_contract": {
                "passed": False,
                "paragraph_direct_format_issues": [{"paragraph": "1.r1.c1.p1"}],
            },
        },
    }

    add_risk_warnings(report, "at-least")

    assert any(item["type"] == "table_format_contract" for item in report["risk_warnings"])
